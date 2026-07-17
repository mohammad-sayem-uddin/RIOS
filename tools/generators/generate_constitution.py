#!/usr/bin/env python3
"""
RIOS Claude Code Constitution v1.0 Generator
Produces a professional Microsoft Word (.docx) document.
"""

import sys
sys.path.insert(0, '.')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ─── Constants ───────────────────────────────────────────────────────────────
DARK_NAVY = RGBColor(0x0B, 0x1D, 0x3A)
MEDIUM_BLUE = RGBColor(0x1B, 0x3A, 0x6B)
ACCENT_BLUE = RGBColor(0x2C, 0x5F, 0x8A)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "0B1D3A"
TABLE_ALT_BG = "F0F4F8"
TABLE_WHITE_BG = "FFFFFF"
ACCENT_GOLD = RGBColor(0xC4, 0x9A, 0x1A)
LAW_RED = RGBColor(0x8B, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x6B, 0x3F)

doc = Document()

# ─── Style Setup ─────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = DARK_GRAY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = DARK_NAVY if level <= 2 else MEDIUM_BLUE
    h.font.bold = True
    if level == 1:
        h.font.size = Pt(22)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    elif level == 3:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
    else:
        h.font.size = Pt(11)
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(4)

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ─── Helper Functions ────────────────────────────────────────────────────────

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table_row(table, cells_data, is_header=False, alt=False):
    row = table.add_row()
    for i, (text, bold) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        if is_header:
            run.font.color.rgb = WHITE
            run.bold = True
            set_cell_shading(cell, TABLE_HEADER_BG)
        else:
            run.font.color.rgb = DARK_GRAY
            run.bold = bold
            if alt:
                set_cell_shading(cell, TABLE_ALT_BG)
    return row

def create_table(headers, rows_data):
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, [(h, True) for h in headers], is_header=True)
    for i, row_data in enumerate(rows_data):
        add_table_row(table, row_data, alt=(i % 2 == 0))
    return table

def add_law(number, title, text):
    p = doc.add_paragraph()
    run = p.add_run(f"Law {number} — {title}")
    run.bold = True
    run.font.color.rgb = LAW_RED
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    p2 = doc.add_paragraph(text)
    p2.paragraph_format.left_indent = Cm(0.5)
    return p2

def add_rule(number, text):
    p = doc.add_paragraph()
    run = p.add_run(f"{number}: ")
    run.bold = True
    run.font.color.rgb = MEDIUM_BLUE
    run.font.size = Pt(10)
    p.add_run(text).font.size = Pt(10)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p

def add_numbered(text, level=0):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p

def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = LIGHT_GRAY
    run.font.size = Pt(10)
    return p

def section_break():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("━" * 70)
    run.font.color.rgb = ACCENT_BLUE
    run.font.size = Pt(6)

def add_decision_tree(question, branches):
    p = doc.add_paragraph()
    run = p.add_run(f"Decision: {question}")
    run.bold = True
    run.font.color.rgb = MEDIUM_BLUE
    for condition, outcome in branches:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1.0)
        run_c = p2.add_run(f"If {condition} → ")
        run_c.font.size = Pt(10)
        run_o = p2.add_run(outcome)
        run_o.bold = True
        run_o.font.color.rgb = GREEN if "Yes" in outcome or "Accept" in outcome or "Create" in outcome else LAW_RED
        run_o.font.size = Pt(10)

# ═════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═════════════════════════════════════════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RIOS")
run.font.size = Pt(48)
run.font.color.rgb = DARK_NAVY
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Claude Code Constitution")
run.font.size = Pt(32)
run.font.color.rgb = MEDIUM_BLUE
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Version 1.0")
run.font.size = Pt(18)
run.font.color.rgb = ACCENT_BLUE

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("The Permanent Engineering Constitution for AI Coding Agents")
run.font.size = Pt(14)
run.font.color.rgb = LIGHT_GRAY
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("and Human Engineers Working on the Research Identity Operating System")
run.font.size = Pt(14)
run.font.color.rgb = LIGHT_GRAY
run.italic = True

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Classification: Normative — Permanent Authority")
run.font.size = Pt(11)
run.font.color.rgb = DARK_NAVY
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Status: Active")
run.font.size = Pt(11)
run.font.color.rgb = DARK_NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y')}")
run.font.size = Pt(11)
run.font.color.rgb = DARK_NAVY

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# VERSION HISTORY
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("Version History", level=1)
create_table(
    ["Version", "Date", "Author", "Classification", "Description"],
    [
        [("1.0", False), (datetime.datetime.now().strftime('%Y-%m-%d'), False),
         ("RIOS Architecture Board", False), ("Normative", False),
         ("Initial release of the RIOS Claude Code Constitution", False)]
    ]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("Table of Contents", level=1)

toc_items = [
    "Version History",
    "Table of Contents",
    "Purpose and Audience",
    "Authority Hierarchy",
    "PART I — Engineering Philosophy",
    "PART II — Architectural Laws",
    "PART III — AI Engineering Behavior",
    "PART IV — Code Generation Rules",
    "PART V — Architecture Compliance Rules",
    "PART VI — Implementation Workflow",
    "PART VII — Refactoring Constitution",
    "PART VIII — Testing Constitution",
    "PART IX — Review Constitution",
    "PART X — Evolution Constitution",
    "PART XI — Forbidden Actions",
    "PART XII — Engineering Checklists",
    "PART XIII — AI Decision Trees",
    "PART XIV — Prompting Constitution",
    "PART XV — AI Self-Audit Framework",
    "Glossary",
    "Appendix A — Architecture Compliance Matrix",
    "Appendix B — Review Scorecard Template",
    "Appendix C — AI Self-Audit Template",
]

for i, item in enumerate(toc_items, 1):
    p = doc.add_paragraph(f"{i}. {item}")
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# PURPOSE AND AUDIENCE
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("Purpose and Audience", level=1)

doc.add_heading("Purpose", level=2)
doc.add_paragraph(
    "This document constitutes the RIOS Claude Code Constitution. It defines the permanent engineering philosophy, "
    "architectural laws, behavioral rules, compliance requirements, and self-audit frameworks that every AI coding agent "
    "and human engineer SHALL follow when designing, reviewing, implementing, testing, refactoring, and evolving the "
    "Research Identity Operating System (RIOS)."
)
doc.add_paragraph(
    "This Constitution is the highest engineering authority after the RIOS Architecture itself. "
    "It teaches AI agents HOW to think, not merely HOW to code."
)

doc.add_heading("Audience", level=2)
doc.add_paragraph("This document addresses:")
for audience in [
    "Claude Code (Anthropic)",
    "GitHub Copilot (Microsoft)",
    "Cursor AI (Cursor Inc.)",
    "OpenAI Codex (OpenAI)",
    "Google Gemini (Google DeepMind)",
    "Future AI coding agents",
    "Human software engineers implementing RIOS",
    "Architecture reviewers and governance boards",
    "Quality assurance engineers",
    "Technical documentation authors",
]:
    add_bullet(audience)

doc.add_heading("Authority Hierarchy", level=2)
doc.add_paragraph(
    "The following hierarchy establishes the authoritative precedence of RIOS governance documents. "
    "When conflicts arise, higher-ranked documents take precedence."
)

create_table(
    ["Rank", "Document", "Classification", "Authority"],
    [
        [("1", False), ("Master Architecture Blueprint (MAB)", True), ("Normative", False), ("Supreme Architectural Authority", False)],
        [("2", False), ("Architecture Decision Records (ADR)", True), ("Normative", False), ("Decision Authority", False)],
        [("3", False), ("Domain Dependency Matrix (DDM)", True), ("Normative", False), ("Dependency Authority", False)],
        [("4", False), ("Domain Model Specification (DMS)", True), ("Normative", False), ("Domain Authority", False)],
        [("5", False), ("Canonical Terminology Dictionary (CTD)", True), ("Normative", False), ("Semantic Authority", False)],
        [("6", False), ("Domain Ownership Matrix (DOM)", True), ("Normative", False), ("Ownership Authority", False)],
        [("7", False), ("Architecture Governance Standard (AGS)", True), ("Normative", False), ("Governance Authority", False)],
        [("8", False), ("Requirement Taxonomy Standard (RTS)", True), ("Normative", False), ("Requirement Authority", False)],
        [("9", False), ("RIOS Editorial Standard (RES)", True), ("Normative", False), ("Documentation Authority", False)],
        [("10", False), ("RIOS Claude Code Constitution", True), ("Normative", False), ("Engineering Authority", False)],
    ]
)

add_quote(
    "Architecture owns semantics. Engineering owns realization. Implementation never defines architecture. "
    "This Constitution teaches AI agents to honor that boundary absolutely."
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# PART I — ENGINEERING PHILOSOPHY
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("PART I — Engineering Philosophy", level=1)
section_break()

doc.add_heading("1.1 What RIOS Engineering Means", level=2)
doc.add_paragraph(
    "RIOS engineering is the disciplined practice of translating architectural intent into working software "
    "while preserving the semantic integrity, dependency structure, domain boundaries, and long-term evolvability "
    "of the Research Identity Operating System. RIOS engineering is not coding. It is architecture-preserving realization."
)

doc.add_heading("1.2 The Thirteen Pillars of RIOS Engineering", level=2)

pillars = [
    ("Research-First Engineering",
     "Every engineering decision originates from a research question. Implementation serves scientific inquiry, "
     "not the reverse. A feature that cannot trace its origin to a research concept SHALL NOT exist in RIOS."),
    ("Architecture-First Engineering",
     "Architecture precedes implementation. When architecture and implementation conflict, architecture ALWAYS wins. "
     "No implementation convenience justifies architectural compromise."),
    ("Semantic-First Engineering",
     "Meaning precedes mechanism. Before defining an interface, define its semantic contract. Before building a service, "
     "define its bounded context. Before writing a function, define its purpose in domain language."),
    ("Knowledge-First Engineering",
     "Knowledge is the highest intellectual output of RIOS. Every engineering decision must ultimately serve the "
     "capture, preservation, communication, or evolution of scientific knowledge."),
    ("Identity-First Engineering",
     "Research Identity is the primary organizing concept of RIOS. Every subsystem exists to strengthen or "
     "communicate Research Identity. Identity is emergent—it derives from evidence, not declaration."),
    ("Technology Independence",
     "Architecture SHALL NOT embed implementation-specific technology choices. RIOS architecture must remain "
     "valid across programming languages, frameworks, databases, and deployment platforms. Technology serves "
     "architecture; architecture does not serve technology."),
    ("Long-Term Evolution",
     "Engineering decisions are evaluated on decades-long timescales. A solution that works today but "
     "constrains tomorrow is rejected. RIOS is designed to outlive every technology choice."),
    ("Clean Engineering",
     "Clean code is not merely aesthetic—it is a prerequisite for correct reasoning about domain logic. "
     "Clean boundaries enable clean thinking. Messy code produces messy architecture."),
    ("Minimal Complexity",
     "The best engineering is the simplest engineering that correctly serves the architecture. Every unnecessary "
     "abstraction, every premature optimization, every clever trick introduces risk. Simplicity is a feature."),
    ("Maintainability",
     "Every module SHALL be maintainable by a developer who has not previously seen the code. "
     "Self-documenting structure, consistent naming, and explicit contracts replace tribal knowledge."),
    ("Readability",
     "Code is read more often than it is written. Readability serves every other quality attribute. "
     "If a human cannot understand the code in a single reading, the code is wrong."),
    ("Explicitness",
     "Implicit behavior is dangerous behavior. Every dependency, every side effect, every assumption "
     "SHALL be explicitly declared. Magic is the enemy of reliability."),
    ("Preservation Over Innovation",
     "When in doubt, preserve the existing architecture. Innovation is welcome when it strengthens the "
     "architecture. Innovation that compromises architecture is not innovation—it is destruction."),
]

for i, (title, desc) in enumerate(pillars, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"Pillar {i}: {title}")
    run.bold = True
    run.font.color.rgb = MEDIUM_BLUE
    run.font.size = Pt(12)
    doc.add_paragraph(desc)

doc.add_heading("1.3 Engineering Philosophy Matrix", level=2)
create_table(
    ["Principle", "Means", "Does NOT Mean"],
    [
        [("Research-First", False), ("Implementation serves scientific inquiry", False), ("Research determines technology", False)],
        [("Architecture-First", False), ("Architecture owns all semantic decisions", False), ("Architecture dictates code style", False)],
        [("Semantic-First", False), ("Meaning precedes mechanism", False), ("Semantics prevent optimization", False)],
        [("Knowledge-First", False), ("All work serves knowledge capture", False), ("Only knowledge features are built", False)],
        [("Identity-First", False), ("Identity is the organizing principle", False), ("Identity overrides all concerns", False)],
        [("Tech Independent", False), ("Architecture is portable", False), ("No technology is ever chosen", False)],
        [("Long-Term", False), ("Decades-long evaluation horizon", False), ("No short-term decisions allowed", False)],
        [("Clean", False), ("Clear boundaries and structure", False), ("Perfect code is required", False)],
        [("Minimal", False), ("Simplest correct solution", False), ("No abstraction is allowed", False)],
        [("Explicit", False), ("All behavior is declared", False), ("Verbosity is required", False)],
    ]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# PART II — ARCHITECTURAL LAWS
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("PART II — Architectural Laws", level=1)
section_break()
doc.add_paragraph(
    "The following laws are PERMANENT. They SHALL NOT be overridden, weakened, or suspended by any implementation "
    "decision. Violation of any architectural law constitutes a CRITICAL engineering failure."
)

doc.add_heading("2.1 Sovereignty Laws", level=2)

sovereignty_laws = [
    ("AL-001", "Architecture Owns Semantics", "The architecture exclusively defines the meaning of every concept, entity, relationship, and boundary within RIOS. Implementation SHALL NOT alter, extend, or reinterpret architectural semantics."),
    ("AL-002", "Engineering Owns Realization", "Engineering determines HOW architectural intent is realized in software. Engineering decisions are constrained by but do not alter architecture."),
    ("AL-003", "Implementation Never Defines Architecture", "Code SHALL NOT introduce concepts, boundaries, or relationships not defined in the architecture. If implementation needs something undefined, an architectural proposal SHALL be raised."),
    ("AL-004", "MAB Is Supreme", "The Master Architecture Blueprint is the single supreme authority. No document, decision, or implementation may contradict the MAB without a formal architectural revision."),
    ("AL-005", "ADRs Govern Exceptions", "Every exception to architectural rules SHALL be documented in an Architecture Decision Record with full context, alternatives, decision, and consequences."),
]
for num, title, text in sovereignty_laws:
    add_law(num, title, text)

doc.add_heading("2.2 Domain Laws", level=2)

domain_laws = [
    ("AL-006", "One Responsibility Per Domain", "Every domain SHALL own exactly one primary responsibility. No responsibility SHALL be shared between domains without explicit architectural approval documented in an ADR."),
    ("AL-007", "DDD Boundaries Are Inviolable", "Bounded Contexts, Aggregates, and domain boundaries SHALL NOT be crossed without using published interfaces. Direct internal access across domains is PROHIBITED."),
    ("AL-008", "Every Feature Maps to Architecture", "Every feature in the implementation SHALL map to at least one architectural concept, domain entity, or requirement. Unmapped features SHALL be removed."),
    ("AL-009", "Every Module Maps to a Domain", "Every software module SHALL belong to exactly one domain. Modules spanning multiple domains SHALL be decomposed."),
    ("AL-010", "Aggregates Are Consistency Boundaries", "Each Aggregate Root guarantees the consistency of the entities within its boundary. Cross-aggregate consistency uses eventual consistency through domain events."),
    ("AL-011", "Domain Events Are Immutable", "Once a domain event is recorded, it SHALL NOT be modified or deleted. Domain events represent immutable historical facts."),
    ("AL-012", "Canonical Dependency Direction", "Dependencies SHALL flow downward: Identity → Knowledge → Narrative → Publication → Visualization → Motion → Engineering → Evolution. Reverse dependencies are PROHIBITED."),
    ("AL-013", "No Circular Dependencies", "Circular architectural dependencies SHALL NOT exist. All dependency chains SHALL terminate at the Identity Domain or at a domain with no upstream dependencies."),
    ("AL-014", "Stable Domains Precede Variable Domains", "Domains representing long-lived concepts (Identity, Knowledge) SHALL NOT depend on rapidly changing domains (Motion, Visualization). Scientific meaning precedes presentation."),
    ("AL-015", "Orthogonal Concerns Remain Cross-Cutting", "Engineering and Evolution are orthogonal cross-cutting concerns. They SHALL NOT enter the strict layered hierarchy of intellectual domains."),
]
for num, title, text in domain_laws:
    add_law(num, title, text)

doc.add_heading("2.3 Semantic Laws", level=2)

semantic_laws = [
    ("AL-016", "Semantic Contracts Before APIs", "Before defining any interface, define its semantic contract: Purpose, Input, Output, Consistency, Ownership. APIs are realization of contracts."),
    ("AL-017", "Canonical Terminology Is Binding", "All terms defined in the Canonical Terminology Dictionary SHALL be used exactly as defined. Alternative interpretations are non-conforming unless documented in an ADR."),
    ("AL-018", "Single Source of Truth", "No intellectual concept SHALL be defined in more than one domain. Definitions are canonical; references are by citation."),
    ("AL-019", "Evidence Before Claims", "Every Knowledge entity SHALL reference supporting Evidence. No claim shall exist without verifiable evidence."),
    ("AL-020", "Identity Before Knowledge", "Research Identity is the primary organizing concept. Knowledge, Narrative, Publication, and all other domains ultimately serve Identity."),
    ("AL-021", "Questions Before Projects", "Research Questions define scientific direction. Projects are temporary vehicles for pursuing questions. Questions persist; projects conclude."),
    ("AL-022", "Knowledge Before Documents", "Understanding precedes expression. Knowledge SHALL be captured before it is published. Publication without knowledge is noise."),
    ("AL-023", "Motion Serves Cognition", "Motion design serves cognitive understanding, not aesthetic preference. Every animation, transition, and interaction SHALL have a cognitive justification."),
    ("AL-024", "Visualization Serves Meaning", "Visualization communicates scientific reasoning. It does not decorate. Every visual element SHALL serve semantic communication."),
]
for num, title, text in semantic_laws:
    add_law(num, title, text)

doc.add_heading("2.4 CQRS Laws", level=2)

cqrs_laws = [
    ("AL-025", "CQRS Is Mandatory", "RIOS mandates strict Command Query Responsibility Segregation. Commands mutate state. Queries read state. These SHALL NOT be mixed."),
    ("AL-026", "Identity Is Read-Only Projection", "Research Identity is an emergent, read-only projection derived from the event stream. Identity cannot be directly edited—it can only be inferred from evidence."),
    ("AL-027", "Commands Produce Events", "Every state change SHALL produce a verifiable domain event. No aggregate root may change state without a corresponding domain event."),
    ("AL-028", "Event Stream Is Append-Only", "The academic event stream is append-only. History SHALL NOT be updated or deleted, preserving absolute traceability."),
    ("AL-029", "Deterministic Replay", "The system SHALL be capable of reconstructing any projection from event zero, ensuring that every read model is a pure function of verifiable history."),
]
for num, title, text in cqrs_laws:
    add_law(num, title, text)

doc.add_heading("2.5 Quality Laws", level=2)

quality_laws = [
    ("AL-030", "Explicit Cardinality", "All relationships SHALL explicitly declare structural cardinality (1:1, 1:N, M:N) to ensure precise semantic bounds."),
    ("AL-031", "Technology Independence", "Architecture documents SHALL NOT contain implementation-specific technology choices. REST, JSON, gRPC, SQL, and other technologies belong to engineering, not architecture."),
    ("AL-032", "Normative Language Precision", "SHALL/MUST denotes absolute requirements. SHOULD denotes strong recommendations. MAY denotes optional paths. Normative language SHALL NOT be used for philosophical statements."),
    ("AL-033", "Document Structure Consistency", "All domain volumes SHALL follow the canonical document structure: Purpose → Ontology → Entities → Relationships → Rules → Semantic Contracts → Verification."),
    ("AL-034", "Traceability Completeness", "Every requirement, entity, relationship, and decision SHALL be traceable to its architectural origin. Untraceable artifacts are non-conforming."),
    ("AL-035", "Knowledge Preservation", "Architectural reasoning SHALL be preserved alongside architectural outcomes. Future contributors SHALL understand WHY a decision was made, not only WHAT was decided."),
    ("AL-036", "Versioning Discipline", "Architecture documents SHALL follow Semantic Versioning. Major versions indicate architectural changes. Minor versions indicate new capabilities. Patch versions indicate editorial improvements."),
]
for num, title, text in quality_laws:
    add_law(num, title, text)

doc.add_page_break()

print("Parts I-II complete. Continuing...")

# ═════════════════════════════════════════════════════════════════════════════
# PART III — AI ENGINEERING BEHAVIOR
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("PART III — AI Engineering Behavior", level=1)
section_break()

doc.add_heading("3.1 Identity of the AI Agent", level=2)
doc.add_paragraph(
    "The AI coding agent operating on RIOS is NOT a code generator. It is NOT an autocomplete tool. "
    "The AI agent is an Architect, an Engineer, a Reviewer, and a Guardian of RIOS. Its primary responsibility "
    "is preserving the architecture. If architecture and implementation conflict, architecture ALWAYS wins."
)

doc.add_heading("3.2 Mandatory Pre-Generation Reasoning Workflow", level=2)
doc.add_paragraph(
    "Before generating ANY code, the AI agent SHALL execute the following reasoning workflow completely. "
    "No step may be skipped, abbreviated, or assumed."
)

workflow_steps = [
    ("Step 1: Understand the Request",
     "Parse the user's request completely. Identify what is being asked: new feature, bug fix, refactoring, "
     "architecture change, documentation update, or review. Classify the request's scope and impact."),
    ("Step 2: Locate the Architecture",
     "Identify which architectural domains are affected. Read the relevant domain specifications, "
     "semantic contracts, and invariant rules. Do NOT proceed without understanding the architecture."),
    ("Step 3: Read Relevant ADRs",
     "Search the Architecture Decision Records for decisions affecting the current task. "
     "Understand the context, alternatives considered, and consequences of past decisions."),
    ("Step 4: Locate Traceability",
     "Find the traceability entries for the concepts involved. Identify parent requirements, "
     "dependent requirements, and affected domains. No implementation begins without traceability."),
    ("Step 5: Understand Ownership",
     "Determine the Primary Owner of every concept involved. Verify that the task does not "
     "violate ownership boundaries. Use published interfaces for cross-domain interaction."),
    ("Step 6: Determine the Domain",
     "Assign every new concept, entity, and service to exactly one domain. "
     "If a concept spans domains, propose a decomposition or raise an architectural question."),
    ("Step 7: Determine the Aggregate",
     "Identify which Aggregate Root governs the entities involved. "
     "Ensure consistency boundaries are respected. Cross-aggregate operations use eventual consistency."),
    ("Step 8: Determine the Semantic Contract",
     "For any interface being created or modified, define its semantic contract: "
     "Purpose, Input, Output, Consistency guarantee, and Ownership."),
    ("Step 9: Determine the Interfaces",
     "Define the technology-independent interfaces that expose the domain's capabilities. "
     "Interfaces SHALL NOT expose internal implementation details."),
    ("Step 10: Generate Implementation Plan",
     "Produce a step-by-step implementation plan. The plan SHALL include: "
     "modules affected, new files created, modifications to existing files, test strategy, and compliance notes."),
    ("Step 11: Review the Plan",
     "Self-audit the implementation plan against: architectural laws, domain rules, "
     "dependency direction, ownership boundaries, semantic contracts, and CQRS patterns."),
    ("Step 12: Generate Code",
     "Only after the plan passes self-audit, generate the implementation code following "
     "all rules in Part IV (Code Generation Rules)."),
    ("Step 13: Review the Code",
     "Self-audit the generated code against: naming conventions, structural patterns, "
     "error handling, type safety, immutability, documentation, and test coverage."),
    ("Step 14: Run Self-Audit",
     "Execute the complete AI Self-Audit framework from Part XV. "
     "Do not present results until the self-audit passes."),
    ("Step 15: Produce Implementation",
     "Present the final implementation with: code, documentation, test specifications, "
     "compliance report, and traceability updates."),
]

for title, desc in workflow_steps:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = MEDIUM_BLUE
    run.font.size = Pt(12)
    doc.add_paragraph(desc)

doc.add_heading("3.3 AI Behavioral Invariants", level=2)

invariants = [
    ("INV-001: Never Skip Architecture Reading", "The AI SHALL NOT generate code without first reading the relevant architectural specifications."),
    ("INV-002: Never Assume Ownership", "The AI SHALL NOT assume which domain owns a concept. It SHALL verify ownership from the Domain Ownership Matrix."),
    ("INV-003: Never Invent Architecture", "The AI SHALL NOT create architectural concepts, boundaries, or relationships not defined in the MAB."),
    ("INV-004: Never Bypass ADRs", "The AI SHALL NOT make architectural decisions without checking existing ADRs and creating new ones when needed."),
    ("INV-005: Never Violate Semantic Contracts", "The AI SHALL NOT produce code that violates the semantic contracts defined in domain specifications."),
    ("INV-006: Never Suppress Warnings", "If the AI detects an architectural conflict, it SHALL report it explicitly, never suppress it."),
    ("INV-007: Always Trace to Architecture", "Every line of code the AI generates SHALL trace back to an architectural concept or requirement."),
    ("INV-008: Always Preserve CQRS", "The AI SHALL maintain strict separation between Commands and Queries. Never mix read and write models."),
    ("INV-009: Always Use Canonical Terms", "The AI SHALL use terms exactly as defined in the Canonical Terminology Dictionary."),
    ("INV-010: Always Document Decisions", "The AI SHALL document the reasoning behind every implementation choice, referencing architectural justification."),
]

for inv in invariants:
    p = doc.add_paragraph()
    run = p.add_run(f"{inv[0]}: ")
    run.bold = True
    run.font.color.rgb = LAW_RED
    run.font.size = Pt(10)
    p.add_run(inv[1]).font.size = Pt(10)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# PART IV — CODE GENERATION RULES
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("PART IV — Code Generation Rules", level=1)
section_break()

doc.add_heading("4.1 Naming Rules", level=2)
naming_rules = [
    ("CGR-N-001", "All class names SHALL use PascalCase and reflect domain language."),
    ("CGR-N-002", "All method names SHALL use camelCase and express intent or query."),
    ("CGR-N-003", "All constants SHALL use UPPER_SNAKE_CASE."),
    ("CGR-N-004", "All domain entities SHALL be named using canonical terminology from the CTD."),
    ("CGR-N-005", "All value objects SHALL include their semantic category in the name (e.g., EmailAddress, ResearchQuestion)."),
    ("CGR-N-006", "All domain events SHALL be past-tense verbs (e.g., EvidencePublished, AgendaAdopted)."),
    ("CGR-N-007", "All commands SHALL be imperative verbs (e.g., PublishEvidence, AdoptAgenda)."),
    ("CGR-N-008", "All queries SHALL be noun-based (e.g., ResearchIdentityQuery, PublicationSearch)."),
    ("CGR-N-009", "Factory names SHALL end with 'Factory' (e.g., PublicationFactory)."),
    ("CGR-N-010", "Repository names SHALL end with 'Repository' (e.g., EvidenceRepository)."),
    ("CGR-N-011", "Service names SHALL end with 'Service' (e.g., KnowledgeExtractionService)."),
    ("CGR-N-012", "Policy names SHALL end with 'Policy' (e.g., EvidenceVerificationPolicy)."),
    ("CGR-N-013", "Interface names SHALL begin with 'I' or use descriptive nouns depending on language convention."),
    ("CGR-N-014", "Boolean variables and methods SHALL be prefixed with is, has, can, or should."),
    ("CGR-N-015", "Collection variables SHALL use plural nouns (e.g., publications, researchAreas)."),
    ("CGR-N-016", "Names SHALL NOT contain abbreviations unless they are canonical domain terms."),
    ("CGR-N-017", "Names SHALL NOT contain implementation details (e.g., no SqlRepository, RestController)."),
    ("CGR-N-018", "Names SHALL be pronounceable and self-documenting."),
    ("CGR-N-019", "Package/namespace structure SHALL mirror domain boundaries."),
    ("CGR-N-020", "No name SHALL contain the word 'Helper', 'Util', or 'Manager' unless architecturally justified."),
]
for num, text in naming_rules:
    add_rule(num, text)

doc.add_heading("4.2 Structural Rules", level=2)
structural_rules = [
    ("CGR-S-001", "Each domain SHALL be organized into a dedicated directory/package."),
    ("CGR-S-002", "Domain directories SHALL contain subdirectories for: entities, value_objects, aggregates, events, commands, queries, services, repositories, factories, policies."),
    ("CGR-S-003", "Infrastructure code SHALL be in a separate directory from domain code."),
    ("CGR-S-004", "Application services SHALL orchestrate domain objects. They SHALL NOT contain domain logic."),
    ("CGR-S-005", "Domain services SHALL contain logic that does not naturally belong to a single entity."),
    ("CGR-S-006", "Each Aggregate Root SHALL be in its own file."),
    ("CGR-S-007", "Each Entity SHALL be in its own file unless closely related to its Aggregate Root."),
    ("CGR-S-008", "Each Value Object SHALL be in its own file."),
    ("CGR-S-009", "Each Domain Event SHALL be in its own file."),
    ("CGR-S-010", "Shared kernel code SHALL be explicitly labeled and minimized."),
]
for num, text in structural_rules:
    add_rule(num, text)

doc.add_heading("4.3 DDD Rules", level=2)
ddd_rules = [
    ("CGR-D-001", "Entities SHALL have a unique identity that persists through state changes."),
    ("CGR-D-002", "Value Objects SHALL be immutable. They are defined by their attributes, not identity."),
    ("CGR-D-003", "Aggregates SHALL protect their invariants. External code SHALL NOT modify aggregate internals."),
    ("CGR-D-004", "Aggregates SHALL reference other aggregates only by identity, not by direct object reference."),
    ("CGR-D-005", "Domain Events SHALL carry all information needed to understand the state change."),
    ("CGR-D-006", "Domain Events SHALL be immutable once created."),
    ("CGR-D-007", "Repositories SHALL be defined in the domain layer. Implementations belong to infrastructure."),
    ("CGR-D-008", "Factories SHALL encapsulate complex object creation logic."),
    ("CGR-D-009", "Policies SHALL encapsulate business rules that span multiple entities or aggregates."),
    ("CGR-D-010", "Anti-corruption layers SHALL translate between bounded contexts."),
    ("CGR-D-011", "Domain logic SHALL NOT exist in application services, controllers, or infrastructure."),
    ("CGR-D-012", "Bounded contexts SHALL communicate through events or published interfaces, not direct coupling."),
    ("CGR-D-013", "Ubiquitous language SHALL be used consistently within each bounded context."),
    ("CGR-D-014", "Anemic domain models (entities with only getters/setters) are PROHIBITED."),
    ("CGR-D-015", "Domain objects SHALL enforce their own invariants through encapsulation."),
]
for num, text in ddd_rules:
    add_rule(num, text)

doc.add_heading("4.4 Dependency Rules", level=2)
dep_rules = [
    ("CGR-DEP-001", "Dependencies SHALL flow downward only: Identity → Knowledge → Narrative → Publication → Visualization → Motion."),
    ("CGR-DEP-002", "Domain code SHALL NOT import from infrastructure layers."),
    ("CGR-DEP-003", "Domain code SHALL NOT import from application layers."),
    ("CGR-DEP-004", "Infrastructure SHALL depend on domain interfaces, not domain implementations."),
    ("CGR-DEP-005", "Application services SHALL depend on domain interfaces."),
    ("CGR-DEP-006", "The Dependency Inversion Principle SHALL be used to reverse dependencies where needed."),
    ("CGR-DEP-007", "Circular dependencies between modules are PROHIBITED."),
    ("CGR-DEP-008", "Each domain SHALL expose a clean public API. Internal details SHALL NOT be exported."),
    ("CGR-DEP-009", "Shared libraries SHALL be versioned and treated as external dependencies."),
    ("CGR-DEP-010", "Engineering and Evolution services are accessible by all domains as orthogonal concerns."),
]
for num, text in dep_rules:
    add_rule(num, text)

doc.add_heading("4.5 Layering Rules", level=2)
layer_rules = [
    ("CGR-L-001", "Presentation layer SHALL NOT contain business logic."),
    ("CGR-L-002", "Application layer SHALL orchestrate, not implement domain logic."),
    ("CGR-L-003", "Domain layer SHALL be self-contained and infrastructure-independent."),
    ("CGR-L-004", "Infrastructure layer SHALL implement domain interfaces."),
    ("CGR-L-005", "Each layer SHALL only depend on layers below it or on abstractions it owns."),
    ("CGR-L-006", "Cross-cutting concerns (logging, security, caching) SHALL use decorators or middleware, not inline code."),
]
for num, text in layer_rules:
    add_rule(num, text)

doc.add_heading("4.6 Error Handling Rules", level=2)
error_rules = [
    ("CGR-E-001", "Domain errors SHALL be represented as domain-specific exception types."),
    ("CGR-E-002", "Exceptions SHALL carry enough context for diagnosis."),
    ("CGR-E-003", "Infrastructure errors SHALL be translated to domain errors at boundary."),
    ("CGR-E-004", "Error messages SHALL use canonical terminology."),
    ("CGR-E-005", "Errors SHALL be logged with sufficient context for investigation."),
    ("CGR-E-006", "Unhandled exceptions SHALL NEVER propagate to the user interface."),
    ("CGR-E-007", "Invariant violations SHALL throw domain-specific exceptions."),
    ("CGR-E-008", "Retry policies SHALL be defined for transient infrastructure failures."),
]
for num, text in error_rules:
    add_rule(num, text)

doc.add_heading("4.7 Type Safety and Immutability Rules", level=2)
type_rules = [
    ("CGR-T-001", "All public APIs SHALL use strong typing. No untyped parameters."),
    ("CGR-T-002", "Value Objects SHALL be immutable. Once created, their state SHALL NOT change."),
    ("CGR-T-003", "Domain Events SHALL be immutable."),
    ("CGR-T-004", "Collections returned from entities SHALL be defensive copies."),
    ("CGR-T-005", "Nullable types SHALL be explicitly declared."),
    ("CGR-T-006", "Type assertions SHALL validate input at domain boundaries."),
    ("CGR-T-007", "Generic types SHALL carry meaningful constraints."),
    ("CGR-T-008", "State mutation SHALL only occur through aggregate methods."),
    ("CGR-T-009", "Entity state SHALL be validated before persistence."),
    ("CGR-T-010", "Immutable data structures SHALL be preferred over mutable ones."),
]
for num, text in type_rules:
    add_rule(num, text)

doc.add_heading("4.8 Documentation Rules", level=2)
doc_rules = [
    ("CGR-DOC-001", "Every public API SHALL have documentation explaining purpose, parameters, return values, and exceptions."),
    ("CGR-DOC-002", "Every domain entity SHALL document its invariants."),
    ("CGR-DOC-003", "Every aggregate SHALL document its consistency boundaries."),
    ("CGR-DOC-004", "Every domain event SHALL document what it represents and when it is raised."),
    ("CGR-DOC-005", "Every semantic contract SHALL be documented with Purpose, Input, Output, Consistency, and Ownership."),
    ("CGR-DOC-006", "Complex algorithms SHALL include explanatory comments."),
    ("CGR-DOC-007", "Code comments SHALL explain WHY, not WHAT."),
    ("CGR-DOC-008", "Architecture Decision Records SHALL be referenced in code comments where relevant."),
    ("CGR-DOC-009", "README files SHALL exist at domain and module levels."),
    ("CGR-DOC-010", "Generated documentation SHALL be kept in sync with code changes."),
]
for num, text in doc_rules:
    add_rule(num, text)

doc.add_heading("4.9 Validation Rules", level=2)
val_rules = [
    ("CGR-V-001", "Input validation SHALL occur at domain boundaries."),
    ("CGR-V-002", "Domain invariants SHALL be validated within aggregates."),
    ("CGR-V-003", "Cross-aggregate validation SHALL use domain events and policies."),
    ("CGR-V-004", "External input SHALL be validated before entering the domain layer."),
    ("CGR-V-005", "Validation errors SHALL use domain-specific error types."),
    ("CGR-V-006", "Validation rules SHALL be co-located with the entities they protect."),
]
for num, text in val_rules:
    add_rule(num, text)

doc.add_heading("4.10 Configuration and DI Rules", level=2)
config_rules = [
    ("CGR-C-001", "Configuration SHALL be externalized, not embedded in code."),
    ("CGR-C-002", "Dependency injection SHALL wire infrastructure to domain interfaces."),
    ("CGR-C-003", "Domain objects SHALL NOT be aware of their injection container."),
    ("CGR-C-004", "Configuration values SHALL be validated at startup."),
    ("CGR-C-005", "Environment-specific configuration SHALL be separated from architectural configuration."),
]
for num, text in config_rules:
    add_rule(num, text)

doc.add_page_break()

print("Parts III-IV complete. Continuing...")

# ═════════════════════════════════════════════════════════════════════════════
# PART V — ARCHITECTURE COMPLIANCE RULES
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("PART V — Architecture Compliance Rules", level=1)
section_break()

doc.add_heading("5.1 Compliance Categories", level=2)
create_table(
    ["Category", "Scope", "Severity of Violation"],
    [
        [("Architecture Violations", True), ("MAB, ADRs, architectural invariants", False), ("CRITICAL", False)],
        [("DDD Violations", True), ("Bounded contexts, aggregates, entities, value objects", False), ("HIGH", False)],
        [("Dependency Violations", True), ("Dependency direction, circular references", False), ("CRITICAL", False)],
        [("Naming Violations", True), ("Canonical terminology, naming conventions", False), ("MEDIUM", False)],
        [("Ownership Violations", True), ("Domain ownership, responsibility boundaries", False), ("HIGH", False)],
        [("Boundary Violations", True), ("Domain boundaries, layer boundaries", False), ("HIGH", False)],
        [("Semantic Violations", True), ("Semantic contracts, terminology consistency", False), ("CRITICAL", False)],
        [("Technology Leakage", True), ("Implementation details in architecture", False), ("HIGH", False)],
        [("Layer Violations", True), ("Dependency direction between layers", False), ("HIGH", False)],
        [("Circular Dependencies", True), ("Module or domain cycles", False), ("CRITICAL", False)],
        [("Duplicate Concepts", True), ("Same concept defined in multiple places", False), ("HIGH", False)],
    ]
)

doc.add_heading("5.2 Architecture Violation Detection", level=2)
arch_violations = [
    ("ACR-001", "Verify every feature maps to an architectural concept."),
    ("ACR-002", "Verify every module maps to a domain."),
    ("ACR-003", "Verify no implementation defines architectural concepts."),
    ("ACR-004", "Verify MAB compliance for all new features."),
    ("ACR-005", "Verify ADR compliance for all exceptions."),
    ("ACR-006", "Verify CQRS separation is maintained."),
    ("ACR-007", "Verify event stream remains append-only."),
    ("ACR-008", "Verify identity remains read-only projection."),
]
for num, text in arch_violations:
    add_rule(num, text)

doc.add_heading("5.3 DDD Violation Detection", level=2)
ddd_violations = [
    ("DCR-001", "Verify aggregates protect their invariants."),
    ("DCR-002", "Verify entities have immutable identity."),
    ("DCR-003", "Verify value objects are immutable."),
    ("DCR-004", "Verify domain events are immutable."),
    ("DCR-005", "Verify repositories are defined in domain layer."),
    ("DCR-006", "Verify domain logic is not in services/controllers/infrastructure."),
    ("DCR-007", "Verify anti-corruption layers exist between bounded contexts."),
    ("DCR-008", "Verify anemic domain models do not exist."),
    ("DCR-009", "Verify aggregates reference other aggregates by identity only."),
    ("DCR-010", "Verify factories encapsulate complex creation logic."),
]
for num, text in ddd_violations:
    add_rule(num, text)

doc.add_heading("5.4 Dependency Violation Detection", level=2)
dep_violations = [
    ("DVR-001", "Verify no domain depends on domains below it in the canonical graph."),
    ("DVR-002", "Verify no circular dependencies exist."),
    ("DVR-003", "Verify domain code does not import infrastructure."),
    ("DVR-004", "Verify stable domains do not depend on variable domains."),
    ("DVR-005", "Verify dependencies satisfy semantic necessity."),
    ("DVR-006", "Verify dependencies satisfy architectural justification."),
    ("DVR-007", "Verify dependencies satisfy long-term stability."),
    ("DVR-008", "Verify interface-first collaboration between domains."),
]
for num, text in dep_violations:
    add_rule(num, text)

doc.add_heading("5.5 Semantic Violation Detection", level=2)
sem_violations = [
    ("SVR-001", "Verify all terms match CTD definitions."),
    ("SVR-002", "Verify semantic contracts are complete (Purpose, Input, Output, Consistency, Ownership)."),
    ("SVR-003", "Verify no concept is defined in multiple domains."),
    ("SVR-004", "Verify evidence precedes knowledge claims."),
    ("SVR-005", "Verify identity derives from evidence, not declaration."),
    ("SVR-006", "Verify canonical document structure is followed."),
]
for num, text in sem_violations:
    add_rule(num, text)

doc.add_heading("5.6 Full Compliance Checklist", level=2)
compliance_items = [
    "All features map to architecture",
    "All modules map to domains",
    "No circular dependencies",
    "Dependencies flow downward only",
    "Aggregates protect invariants",
    "Value objects are immutable",
    "Domain events are immutable",
    "CQRS separation maintained",
    "Event stream is append-only",
    "Identity is read-only projection",
    "Canonical terminology used consistently",
    "Semantic contracts are complete",
    "No technology in architecture docs",
    "No domain logic in infrastructure",
    "No duplicate concepts",
    "Traceability is complete",
    "ADRs exist for all exceptions",
    "Domain boundaries respected",
    "Ownership boundaries respected",
    "Documentation is complete and current",
]
for item in compliance_items:
    add_bullet(f"☐ {item}")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# PART VI — IMPLEMENTATION WORKFLOW
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("PART VI — Implementation Workflow", level=1)
section_break()

doc.add_heading("6.1 The 15-Step Implementation Workflow", level=2)
doc.add_paragraph(
    "Every implementation task in RIOS SHALL follow this 15-step workflow. No step may be skipped. "
    "The AI agent SHALL document its progress through each step."
)

workflow_table = [
    [("1", False), ("Understand Request", False), ("Parse and classify the task", False), ("Mandatory", False)],
    [("2", False), ("Locate Architecture", False), ("Read relevant domain specs", False), ("Mandatory", False)],
    [("3", False), ("Read ADRs", False), ("Check for relevant ADRs", False), ("Mandatory", False)],
    [("4", False), ("Locate Traceability", False), ("Find traceability entries", False), ("Mandatory", False)],
    [("5", False), ("Understand Ownership", False), ("Verify domain ownership", False), ("Mandatory", False)],
    [("6", False), ("Determine Domain", False), ("Assign to correct domain", False), ("Mandatory", False)],
    [("7", False), ("Determine Aggregate", False), ("Identify governing aggregate", False), ("Mandatory", False)],
    [("8", False), ("Semantic Contract", False), ("Define or verify contract", False), ("Mandatory", False)],
    [("9", False), ("Determine Interfaces", False), ("Define tech-independent interfaces", False), ("Mandatory", False)],
    [("10", False), ("Implementation Plan", False), ("Generate step-by-step plan", False), ("Mandatory", False)],
    [("11", False), ("Review Plan", False), ("Self-audit the plan", False), ("Mandatory", False)],
    [("12", False), ("Generate Code", False), ("Write implementation code", False), ("Mandatory", False)],
    [("13", False), ("Review Code", False), ("Self-audit the code", False), ("Mandatory", False)],
    [("14", False), ("Self-Audit", False), ("Run complete self-audit", False), ("Mandatory", False)],
    [("15", False), ("Produce Output", False), ("Present final implementation", False), ("Mandatory", False)],
]
create_table(["Step", "Name", "Description", "Requirement"], workflow_table)

doc.add_heading("6.2 Step Detail: Understanding the Request", level=2)
doc.add_paragraph(
    "The AI SHALL classify every request into one of these categories: "
    "New Feature, Bug Fix, Refactoring, Architecture Change, Documentation Update, Code Review, or Research Task. "
    "The classification determines the depth of architectural analysis required."
)

doc.add_heading("6.3 Step Detail: Locating the Architecture", level=2)
doc.add_paragraph(
    "The AI SHALL identify: affected domains, affected entities, affected aggregates, affected semantic contracts, "
    "affected invariants, and affected domain events. For each affected element, the AI SHALL read the complete "
    "specification before proceeding."
)

doc.add_heading("6.4 Step Detail: Understanding Ownership", level=2)
doc.add_paragraph(
    "Every concept in RIOS has a Primary Owner. The AI SHALL verify: "
    "Who owns the concept? Is the current task within the owner's scope? "
    "If cross-domain interaction is needed, which published interface SHALL be used? "
    "Are there any ownership conflicts?"
)

doc.add_heading("6.5 Step Detail: Semantic Contract Definition", level=2)
doc.add_paragraph("Every interface SHALL have a complete semantic contract:")
create_table(
    ["Component", "Description", "Example"],
    [
        [("Purpose", False), ("Why the interface exists", False), ("Retrieve researcher identity", False)],
        [("Input", False), ("What the interface requires", False), ("Researcher Identifier", False)],
        [("Output", False), ("What the interface returns", False), ("ResearchIdentity Projection", False)],
        [("Consistency", False), ("Consistency guarantee", False), ("Eventually Consistent", False)],
        [("Ownership", False), ("Which domain owns it", False), ("Identity Domain", False)],
    ]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# PART VII — REFACTORING CONSTITUTION
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("PART VII — Refactoring Constitution", level=1)
section_break()

doc.add_heading("7.1 When Refactoring Is Allowed", level=2)
allowed = [
    "Internal implementation improvement that does not change external behavior",
    "Code simplification that preserves all semantic contracts",
    "Naming improvement that aligns better with canonical terminology",
    "Structural reorganization within a single domain boundary",
    "Performance optimization that preserves correctness",
    "Test improvement that does not change production behavior",
    "Documentation improvement that clarifies existing architecture",
    "Dependency cleanup that maintains correct direction",
]
for item in allowed:
    add_bullet(f"✓ ALLOWED: {item}")

doc.add_heading("7.2 When Refactoring Is Forbidden", level=2)
forbidden_refactor = [
    "Changing domain boundaries without an ADR",
    "Moving ownership between domains without architectural approval",
    "Altering semantic contracts without impact analysis",
    "Changing aggregate boundaries without architectural review",
    "Modifying domain events that other domains depend on",
    "Introducing new dependencies that violate the dependency graph",
    "Removing invariants without architectural justification",
    "Changing CQRS patterns without an ADR",
]
for item in forbidden_refactor:
    add_bullet(f"✗ FORBIDDEN: {item}")

doc.add_heading("7.3 Refactoring Safety Classification", level=2)
create_table(
    ["Safety Level", "Scope", "Approval Required"],
    [
        [("Safe", True), ("Method-level changes within an entity", False), ("None (self-audit only)", False)],
        [("Low Risk", True), ("Class-level changes within a module", False), ("Peer review", False)],
        [("Medium Risk", True), ("Module-level changes within a domain", False), ("Domain owner review", False)],
        [("High Risk", True), ("Cross-module changes within a domain", False), ("Architecture review", False)],
        [("Critical", True), ("Cross-domain changes", False), ("ADR required", False)],
        [("Foundational", True), ("Architecture-level changes", False), ("Full governance review", False)],
    ]
)

doc.add_heading("7.4 Refactoring Rules", level=2)
refactor_rules = [
    ("REF-001", "All refactoring SHALL preserve existing semantic contracts."),
    ("REF-002", "All refactoring SHALL preserve existing tests. If tests break, the refactoring is wrong."),
    ("REF-003", "All refactoring SHALL be done incrementally. Large-batch refactoring is PROHIBITED."),
    ("REF-004", "Refactoring SHALL NOT mix with feature changes. Refactoring and features are separate commits."),
    ("REF-005", "Every refactoring SHALL have a clear architectural motivation."),
    ("REF-006", "Refactoring that increases complexity is PROHIBITED unless justified by an ADR."),
    ("REF-007", "Refactoring SHALL NOT change public API signatures without migration plan."),
    ("REF-008", "Domain event schemas SHALL NOT be refactored. Events are immutable history."),
]
for num, text in refactor_rules:
    add_rule(num, text)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# PART VIII — TESTING CONSTITUTION
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("PART VIII — Testing Constitution", level=1)
section_break()

doc.add_heading("8.1 Testing Philosophy", level=2)
doc.add_paragraph(
    "Testing in RIOS is not an afterthought—it is an architectural verification mechanism. "
    "Tests prove that the implementation correctly realizes the architecture. "
    "A system without comprehensive tests cannot be proven to be architecturally compliant."
)

doc.add_heading("8.2 Test Categories", level=2)

test_categories = [
    ("Architecture Tests", "Verify that the implementation structure matches the architecture. "
     "Test: module boundaries, dependency direction, domain ownership, layer separation."),
    ("DDD Tests", "Verify DDD pattern compliance. "
     "Test: aggregate invariants, entity identity, value object immutability, event immutability."),
    ("Aggregate Tests", "Verify aggregate behavior. "
     "Test: state transitions, invariant enforcement, event publication, command handling."),
    ("Domain Tests", "Verify domain logic. "
     "Test: business rules, domain services, policies, factories."),
    ("Integration Tests", "Verify cross-domain interaction through published interfaces. "
     "Test: semantic contract compliance, event handling, cross-domain queries."),
    ("Contract Tests", "Verify semantic contract compliance. "
     "Test: input/output types, consistency guarantees, error behavior."),
    ("Mutation Tests", "Verify test quality by introducing mutations. "
     "Test: test suite catches intentional defects."),
    ("Property Tests", "Verify domain properties hold for arbitrary inputs. "
     "Test: invariants, consistency, boundary conditions."),
    ("Regression Tests", "Verify that changes do not break existing behavior. "
     "Test: previously fixed bugs, existing features, performance baselines."),
    ("Performance Tests", "Verify system meets performance requirements. "
     "Test: query latency, event processing throughput, projection rebuild time."),
    ("Security Tests", "Verify security properties. "
     "Test: authentication, authorization, data integrity, audit logging."),
    ("Accessibility Tests", "Verify accessibility compliance. "
     "Test: WCAG compliance, screen reader support, keyboard navigation."),
    ("Research Integrity Tests", "Verify scientific integrity properties. "
     "Test: evidence chain completeness, provenance tracking, reproducibility."),
    ("Evidence Integrity Tests", "Verify evidence immutability and traceability. "
     "Test: event immutability, evidence linking, audit trail completeness."),
]

for name, desc in test_categories:
    p = doc.add_paragraph()
    run = p.add_run(f"{name}: ")
    run.bold = True
    run.font.color.rgb = MEDIUM_BLUE
    p.add_run(desc)

doc.add_heading("8.3 Testing Rules", level=2)
test_rules = [
    ("TST-001", "Every aggregate SHALL have unit tests covering all command handlers and invariant enforcement."),
    ("TST-002", "Every domain event SHALL have tests verifying it is immutable."),
    ("TST-003", "Every value object SHALL have tests verifying it is immutable and equality is by value."),
    ("TST-004", "Every repository SHALL have integration tests with the actual persistence mechanism."),
    ("TST-005", "Every semantic contract SHALL have contract tests."),
    ("TST-006", "Cross-domain interactions SHALL have integration tests."),
    ("TST-007", "Tests SHALL NOT depend on external services without explicit configuration."),
    ("TST-008", "Test data SHALL use domain-specific builders, not raw data."),
    ("TST-009", "Test names SHALL describe the scenario and expected outcome."),
    ("TST-010", "Architecture tests SHALL verify dependency direction and domain boundaries."),
]
for num, text in test_rules:
    add_rule(num, text)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# PART IX — REVIEW CONSTITUTION
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("PART IX — Review Constitution", level=1)
section_break()

doc.add_heading("9.1 Review Dimensions", level=2)
doc.add_paragraph(
    "Every code review in RIOS SHALL evaluate ALL of the following dimensions. "
    "Failure in any dimension SHALL block the review."
)

review_dimensions = [
    ("Architecture Review", "Does the code correctly implement the architecture? Does it preserve domain boundaries, "
     "dependency direction, and architectural invariants?"),
    ("DDD Review", "Does the code correctly implement DDD patterns? Are aggregates, entities, value objects, "
     "and events properly implemented?"),
    ("Naming Review", "Does the code use canonical terminology? Are names self-documenting? "
     "Do names reflect domain language?"),
    ("Ownership Review", "Does the code respect domain ownership? Are cross-domain interactions through published interfaces?"),
    ("Security Review", "Does the code handle authentication, authorization, data integrity, and audit logging correctly?"),
    ("Performance Review", "Does the code meet performance requirements? Are there N+1 queries, memory leaks, or inefficient algorithms?"),
    ("Maintainability Review", "Is the code maintainable? Is it readable? Does it follow the DRY principle? "
     "Is complexity minimized?"),
    ("Documentation Review", "Is the code well-documented? Are semantic contracts documented? "
     "Are invariants documented? Are ADRs referenced where relevant?"),
    ("Research Integrity Review", "Does the code preserve scientific integrity? Is evidence chain complete? "
     "Is provenance tracked? Is reproducibility supported?"),
    ("Semantic Review", "Does the code use terms consistently with the CTD? "
     "Do interfaces match their semantic contracts?"),
]

for name, desc in review_dimensions:
    p = doc.add_paragraph()
    run = p.add_run(f"{name}: ")
    run.bold = True
    run.font.color.rgb = MEDIUM_BLUE
    p.add_run(desc)

doc.add_heading("9.2 Review Scorecard", level=2)
create_table(
    ["Dimension", "Weight", "Pass Criteria", "Score (1-5)"],
    [
        [("Architecture Compliance", True), ("25%", False), ("Zero architecture violations", False), ("____", False)],
        [("DDD Compliance", True), ("20%", False), ("All patterns correctly applied", False), ("____", False)],
        [("Semantic Correctness", True), ("15%", False), ("Canonical terminology used", False), ("____", False)],
        [("Dependency Compliance", True), ("10%", False), ("No dependency violations", False), ("____", False)],
        [("Code Quality", True), ("10%", False), ("Clean, readable, maintainable", False), ("____", False)],
        [("Test Coverage", True), ("10%", False), ("All critical paths tested", False), ("____", False)],
        [("Documentation", True), ("5%", False), ("Complete and accurate", False), ("____", False)],
        [("Security", True), ("5%", False), ("No security vulnerabilities", False), ("____", False)],
    ]
)

doc.add_heading("9.3 Review Rules", level=2)
review_rules = [
    ("REV-001", "No code SHALL be merged without passing all review dimensions."),
    ("REV-002", "Architecture violations SHALL be rejected immediately, not deferred."),
    ("REV-003", "Reviewers SHALL verify semantic contract compliance."),
    ("REV-004", "Reviewers SHALL verify test coverage for changed code."),
    ("REV-005", "Reviewers SHALL verify documentation is updated."),
    ("REV-006", "Review score below 3.5 average SHALL require revision."),
    ("REV-007", "Any dimension scoring 1 SHALL block the review regardless of total score."),
    ("REV-008", "Review findings SHALL be documented for future reference."),
]
for num, text in review_rules:
    add_rule(num, text)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# PART X — EVOLUTION CONSTITUTION
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("PART X — Evolution Constitution", level=1)
section_break()

doc.add_heading("10.1 Evolution Philosophy", level=2)
doc.add_paragraph(
    "RIOS architecture evolves intentionally, never accidentally. Evolution is governed by the Architecture "
    "Governance Standard. Every evolution step SHALL preserve semantic compatibility, traceability, and the "
    "fundamental architectural invariants."
)

doc.add_heading("10.2 Evolution Laws", level=2)
evolution_laws = [
    ("EVO-001", "Never Break Semantic Compatibility", "New versions of RIOS SHALL NOT break the meaning of existing concepts. "
     "Deprecated concepts SHALL be maintained alongside replacements during transition periods."),
    ("EVO-002", "Never Change Architecture Without ADR", "Every architectural change SHALL be documented in an Architecture Decision Record "
     "before implementation. Undocumented changes are non-conforming."),
    ("EVO-003", "Never Violate Traceability", "Evolution SHALL maintain complete traceability. Every new concept SHALL be traceable to "
     "its architectural origin. Removed concepts SHALL be documented as retired, not deleted."),
    ("EVO-004", "Always Update Atlas", "The ATLAS SHALL be updated to reflect any architectural evolution. "
     "The ATLAS is the navigational guide to the architecture."),
    ("EVO-005", "Always Update ADRs", "New architectural decisions SHALL be recorded. Existing ADRs SHALL be updated with "
     "consequences of the new decisions."),
    ("EVO-006", "Always Update Documentation", "All affected documentation SHALL be updated before an evolution is considered complete. "
     "Documentation lag is a governance failure."),
    ("EVO-007", "Always Update Traceability", "The Traceability Matrix SHALL be updated to reflect new concepts, retired concepts, "
     "and changed relationships."),
    ("EVO-008", "Preserve Architectural Invariants", "The six architectural invariants SHALL be preserved across all evolution: "
     "Knowledge before Documents, Questions before Projects, Evidence before Claims, "
     "Reasoning before Implementation, Identity emerges from Knowledge, Architecture is Technology Independent."),
    ("EVO-009", "Version According to SemVer", "Major versions for architectural changes. Minor versions for new capabilities. "
     "Patch versions for editorial improvements."),
    ("EVO-010", "Backward Compatibility", "New domains, entities, and relationships SHALL be backward compatible. "
     "Breaking changes require major version bumps."),
]
for num, title, text in evolution_laws:
    add_law(num, title, text)

doc.add_heading("10.3 Evolution Change Classification", level=2)
create_table(
    ["Class", "Type", "Examples", "Process"],
    [
        [("A", False), ("Editorial", False), ("Grammar, formatting, wording", False), ("No architectural review", False)],
        [("B", False), ("Minor", False), ("Additional examples, refinement", False), ("Light review", False)],
        [("C", False), ("Significant", False), ("New entity types, new relationships", False), ("Full architecture review", False)],
        [("D", False), ("Foundational", False), ("MAB changes, principle changes", False), ("ADR + full approval", False)],
    ]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# PART XI — FORBIDDEN ACTIONS
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("PART XI — Forbidden Actions", level=1)
section_break()

doc.add_paragraph(
    "The following actions are PERMANENTLY PROHIBITED. Violation of any forbidden action constitutes "
    "a CRITICAL engineering failure. The AI agent SHALL refuse to perform these actions and SHALL "
    "explicitly report the violation when detected."
)

# Category 1: Architecture Violations
doc.add_heading("11.1 Architecture Violations (FA-ARCH)", level=2)
fa_arch = [
    "FA-ARCH-001: Invent architecture not defined in the MAB",
    "FA-ARCH-002: Bypass Architecture Decision Records",
    "FA-ARCH-003: Move domain ownership without architectural approval",
    "FA-ARCH-004: Duplicate concepts across domains",
    "FA-ARCH-005: Create circular dependencies",
    "FA-ARCH-006: Embed business logic in infrastructure code",
    "FA-ARCH-007: Leak implementation details into architecture documents",
    "FA-ARCH-008: Violate semantic contracts",
    "FA-ARCH-009: Override MAB decisions without formal revision",
    "FA-ARCH-010: Define new architectural principles without governance approval",
    "FA-ARCH-011: Mix CQRS command and query responsibilities",
    "FA-ARCH-012: Modify the event stream (append-only violation)",
    "FA-ARCH-013: Directly edit Research Identity projection",
    "FA-ARCH-014: Remove architectural invariants without major version",
    "FA-ARCH-015: Change canonical dependency direction",
]
for item in fa_arch:
    add_bullet(f"✗ {item}")

# Category 2: DDD Violations
doc.add_heading("11.2 DDD Violations (FA-DDD)", level=2)
fa_ddd = [
    "FA-DDD-001: Create anemic domain models",
    "FA-DDD-002: Place domain logic in application services",
    "FA-DDD-003: Place domain logic in controllers or presentation layer",
    "FA-DDD-004: Place domain logic in infrastructure",
    "FA-DDD-005: Allow external code to modify aggregate internals",
    "FA-DDD-006: Reference aggregates by direct object instead of identity",
    "FA-DDD-007: Make value objects mutable",
    "FA-DDD-008: Make domain events mutable",
    "FA-DDD-009: Expose repository implementations in domain layer",
    "FA-DDD-010: Skip anti-corruption layers between bounded contexts",
    "FA-DDD-011: Create entities without identity",
    "FA-DDD-012: Create aggregates without invariant enforcement",
    "FA-DDD-013: Use non-domain terminology in bounded contexts",
    "FA-DDD-014: Share aggregate internals between domains",
    "FA-DDD-015: Create God Aggregates that span multiple responsibilities",
]
for item in fa_ddd:
    add_bullet(f"✗ {item}")

# Category 3: Dependency Violations
doc.add_heading("11.3 Dependency Violations (FA-DEP)", level=2)
fa_dep = [
    "FA-DEP-001: Create reverse dependencies in the canonical graph",
    "FA-DEP-002: Create circular module dependencies",
    "FA-DEP-003: Import infrastructure from domain code",
    "FA-DEP-004: Import application layer from domain code",
    "FA-DEP-005: Make stable domains depend on variable domains",
    "FA-DEP-006: Introduce dependencies without semantic necessity",
    "FA-DEP-007: Introduce dependencies without architectural justification",
    "FA-DEP-008: Introduce dependencies that lack long-term stability",
    "FA-DEP-009: Access internal structures across domain boundaries",
    "FA-DEP-010: Share mutable state between domains",
]
for item in fa_dep:
    add_bullet(f"✗ {item}")

# Category 4: Semantic Violations
doc.add_heading("11.4 Semantic Violations (FA-SEM)", level=2)
fa_sem = [
    "FA-SEM-001: Use terms inconsistently with the CTD",
    "FA-SEM-002: Define a concept in multiple locations",
    "FA-SEM-003: Remove traceability without documentation",
    "FA-SEM-004: Publish knowledge without supporting evidence",
    "FA-SEM-005: Make identity claims without evidence",
    "FA-SEM-006: Skip semantic contract definition for new interfaces",
    "FA-SEM-007: Change semantic contract without impact analysis",
    "FA-SEM-008: Use implementation-specific language in architecture docs",
    "FA-SEM-009: Mix abstraction levels in a single document",
    "FA-SEM-010: Remove canonical terms from the CTD without replacement",
]
for item in fa_sem:
    add_bullet(f"✗ {item}")

# Category 5: Implementation Violations
doc.add_heading("11.5 Implementation Violations (FA-IMP)", level=2)
fa_imp = [
    "FA-IMP-001: Create shared mutable state",
    "FA-IMP-002: Use singletons for domain objects",
    "FA-IMP-003: Hard-code configuration values",
    "FA-IMP-004: Suppress or swallow exceptions silently",
    "FA-IMP-005: Skip input validation at domain boundaries",
    "FA-IMP-006: Expose internal implementation through public API",
    "FA-IMP-007: Use technology-specific types in domain models",
    "FA-IMP-008: Mix test code with production code",
    "FA-IMP-009: Commit code without tests",
    "FA-IMP-010: Deploy code without documentation",
    "FA-IMP-011: Use magic numbers or strings",
    "FA-IMP-012: Create methods longer than 30 lines without justification",
    "FA-IMP-013: Create classes with more than 10 public methods without justification",
    "FA-IMP-014: Use global state",
    "FA-IMP-015: Implement premature optimization without evidence",
]
for item in fa_imp:
    add_bullet(f"✗ {item}")

# Category 6: Governance Violations
doc.add_heading("11.6 Governance Violations (FA-GOV)", level=2)
fa_gov = [
    "FA-GOV-001: Make architectural decisions without documentation",
    "FA-GOV-002: Skip the architecture review process",
    "FA-GOV-003: Approve changes that break the canonical dependency graph",
    "FA-GOV-004: Modify architecture without updating traceability",
    "FA-GOV-005: Release versions without updating documentation",
    "FA-GOV-006: Introduce Class D changes without full governance approval",
    "FA-GOV-007: Skip quality criteria evaluation",
    "FA-GOV-008: Override architectural invariants without major version",
    "FA-GOV-009: Make governance decisions without recording rationale",
    "FA-GOV-010: Allow architectural drift to go undetected",
]
for item in fa_gov:
    add_bullet(f"✗ {item}")

# Category 7: Documentation Violations
doc.add_heading("11.7 Documentation Violations (FA-DOC)", level=2)
fa_doc = [
    "FA-DOC-001: Create documentation without following RES structure",
    "FA-DOC-002: Use normative language for philosophical statements",
    "FA-DOC-003: Duplicate definitions across documents",
    "FA-DOC-004: Remove traceability references",
    "FA-DOC-005: Leave documentation stale after code changes",
    "FA-DOC-006: Use implementation-specific terminology in architecture docs",
    "FA-DOC-007: Skip Architecture Decision Summaries for significant choices",
    "FA-DOC-008: Create documents that cannot stand alone",
    "FA-DOC-009: Use inconsistent formatting across documents",
    "FA-DOC-010: Omit acceptance criteria from requirements",
]
for item in fa_doc:
    add_bullet(f"✗ {item}")

# Category 8: Testing Violations
doc.add_heading("11.8 Testing Violations (FA-TST)", level=2)
fa_tst = [
    "FA-TST-001: Skip aggregate invariant tests",
    "FA-TST-002: Skip domain event immutability tests",
    "FA-TST-003: Skip value object immutability tests",
    "FA-TST-004: Modify test assertions to match new (incorrect) behavior",
    "FA-TST-005: Ship code without test coverage for critical paths",
    "FA-TST-006: Skip architecture compliance tests",
    "FA-TST-007: Use production data in tests without anonymization",
    "FA-TST-008: Create tests that depend on execution order",
    "FA-TST-009: Skip regression tests after refactoring",
    "FA-TST-010: Ignore test failures in CI/CD pipeline",
]
for item in fa_tst:
    add_bullet(f"✗ {item}")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# PART XII — ENGINEERING CHECKLISTS
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("PART XII — Engineering Checklists", level=1)
section_break()

checklists = [
    ("12.1 Before Coding", [
        "☐ Read the relevant domain specification completely",
        "☐ Read the relevant ADRs",
        "☐ Understand the semantic contract",
        "☐ Verify domain ownership",
        "☐ Identify the aggregate root",
        "☐ Understand invariants",
        "☐ Map the implementation plan to architecture",
        "☐ Verify dependency direction",
        "☐ Confirm CQRS classification (Command or Query)",
        "☐ Identify test strategy",
    ]),
    ("12.2 Before Refactoring", [
        "☐ Classify refactoring safety level",
        "☐ Verify semantic contracts are preserved",
        "☐ Verify existing tests pass",
        "☐ Verify domain boundaries are respected",
        "☐ Verify dependency direction is maintained",
        "☐ Confirm no architectural concepts are affected",
        "☐ Document refactoring motivation",
        "☐ Plan incremental steps",
        "☐ Separate refactoring from feature changes",
    ]),
    ("12.3 Before Merging", [
        "☐ All tests pass",
        "☐ Architecture review complete",
        "☐ DDD review complete",
        "☐ Naming review complete",
        "☐ Ownership review complete",
        "☐ Security review complete",
        "☐ Performance review acceptable",
        "☐ Documentation updated",
        "☐ Traceability updated",
        "☐ Review scorecard passes (≥ 3.5 average)",
    ]),
    ("12.4 Before Architecture Changes", [
        "☐ ADR drafted with full context",
        "☐ Alternatives considered",
        "☐ Impact analysis complete",
        "☐ MAB compliance verified",
        "☐ DDM compliance verified",
        "☐ CTD impact assessed",
        "☐ DOM ownership confirmed",
        "☐ Traceability matrix updated",
        "☐ Classification determined (A/B/C/D)",
        "☐ Full approval obtained for Class C/D changes",
    ]),
    ("12.5 Before Adding Features", [
        "☐ Feature maps to architectural concept",
        "☐ Domain assignment confirmed",
        "☐ Aggregate identified",
        "☐ Semantic contract defined",
        "☐ Dependencies verified",
        "☐ CQRS classification confirmed",
        "☐ Events identified",
        "☐ Invariants identified",
        "☐ Test plan created",
        "☐ Documentation plan created",
    ]),
    ("12.6 Before Adding Domains", [
        "☐ Unique responsibility identified",
        "☐ No overlap with existing domains",
        "☐ Position in dependency graph determined",
        "☐ Canonical graph updated",
        "☐ Domain specification drafted",
        "☐ Semantic contracts defined",
        "☐ Interfaces defined",
        "☐ ADR created",
        "☐ Traceability matrix updated",
        "☐ Atlas updated",
    ]),
    ("12.7 Before Publishing", [
        "☐ Knowledge captured and verified",
        "☐ Evidence linked and accessible",
        "☐ Provenance documented",
        "☐ Version history complete",
        "☐ Semantic contracts satisfied",
        "☐ Research integrity verified",
        "☐ Reproducibility supported",
        "☐ Citations and references complete",
    ]),
    ("12.8 Before Deployment", [
        "☐ All tests pass",
        "☐ Security scan complete",
        "☐ Performance benchmarks acceptable",
        "☐ Documentation complete",
        "☐ Configuration externalized",
        "☐ Monitoring configured",
        "☐ Rollback plan documented",
        "☐ Architecture compliance verified",
    ]),
    ("12.9 Before Version Release", [
        "☐ Semantic versioning applied correctly",
        "☐ Changelog complete",
        "☐ Migration guide written (if breaking)",
        "☐ Documentation updated",
        "☐ Atlas updated",
        "☐ Traceability matrix updated",
        "☐ ADRs updated",
        "☐ Backward compatibility verified",
    ]),
    ("12.10 Before Documentation Updates", [
        "☐ Changes align with RES",
        "☐ Canonical terminology used",
        "☐ No duplicate definitions created",
        "☐ Traceability maintained",
        "☐ Architecture precedence respected",
        "☐ Consistent formatting",
        "☐ Acceptance criteria included",
        "☐ Cross-references updated",
    ]),
]

for title, items in checklists:
    doc.add_heading(title, level=2)
    for item in items:
        add_bullet(item)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# PART XIII — AI DECISION TREES
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("PART XIII — AI Decision Trees", level=1)
section_break()

doc.add_paragraph(
    "The following decision trees guide the AI agent through complex architectural decisions. "
    "Each tree represents a deterministic reasoning process."
)

decision_trees = [
    ("Should I create a new Entity?", [
        ("The concept has a unique identity that persists through state changes", "YES — Create Entity"),
        ("The concept is defined only by its attributes", "NO — Use Value Object instead"),
        ("The concept is part of an aggregate boundary", "YES — Create Entity within Aggregate"),
        ("The concept needs to be tracked independently", "YES — Create Entity"),
        ("The concept is immutable", "NO — Use Value Object"),
    ]),
    ("Should I create a new Aggregate?", [
        ("The concept is a consistency boundary", "YES — Create Aggregate"),
        ("Multiple entities need transactional consistency", "YES — Create Aggregate Root"),
        ("The concept needs to be accessed independently", "YES — Create Aggregate"),
        ("The concept is a child of an existing aggregate", "NO — Add to existing Aggregate"),
        ("The concept spans multiple domains", "NO — Decompose into per-domain Aggregates"),
    ]),
    ("Should I create a new Domain?", [
        ("The concept has a unique responsibility not owned by any existing domain", "YES — Create Domain"),
        ("The concept overlaps with an existing domain", "NO — Extend existing Domain"),
        ("The concept can be assigned to an existing domain", "NO — Use existing Domain"),
        ("The concept requires its own bounded context", "YES — Create Domain"),
        ("The concept is a cross-cutting concern", "NO — Add to Engineering or Evolution"),
    ]),
    ("Should I modify the Architecture?", [
        ("The change affects the MAB", "YES — Requires Class D change + ADR"),
        ("The change introduces new entities", "YES — Requires Class C change + review"),
        ("The change adds examples or refinements", "YES — Class B change, light review"),
        ("The change is grammatical or formatting", "YES — Class A change, no review"),
        ("The change violates an invariant", "NO — Reject the change"),
    ]),
    ("Should I create an ADR?", [
        ("The decision affects architectural structure", "YES — Create ADR"),
        ("The decision introduces new domain concepts", "YES — Create ADR"),
        ("The decision changes dependency relationships", "YES — Create ADR"),
        ("The decision modifies semantic contracts", "YES — Create ADR"),
        ("The decision is a routine implementation choice", "NO — No ADR needed"),
    ]),
    ("Should I split a Module?", [
        ("The module has grown to contain multiple responsibilities", "YES — Split by responsibility"),
        ("The module violates single responsibility principle", "YES — Split"),
        ("The module is cohesive and focused", "NO — Keep as single module"),
        ("The module spans multiple bounded contexts", "YES — Split by bounded context"),
        ("The module has more than 10 classes", "CONSIDER — Evaluate cohesion"),
    ]),
    ("Should I add a Domain Service?", [
        ("The operation does not naturally belong to a single entity", "YES — Create Domain Service"),
        ("The operation coordinates multiple aggregates", "YES — Create Domain Service"),
        ("The operation is a pure business operation", "YES — Create Domain Service"),
        ("The operation belongs to an entity", "NO — Add to Entity"),
        ("The operation is infrastructure-related", "NO — Create Infrastructure Service"),
    ]),
    ("Should I introduce CQRS?", [
        ("RIOS mandates CQRS", "YES — Always apply CQRS"),
        ("The feature reads identity data", "YES — Use Query (Read Model)"),
        ("The feature produces evidence", "YES — Use Command (Write Model)"),
        ("The feature needs real-time consistency", "COMMAND — Event Store guarantees consistency"),
        ("The feature needs eventual consistency", "QUERY — Projection from Event Stream"),
    ]),
    ("Should I create a Value Object?", [
        ("The concept is immutable and defined by attributes", "YES — Create Value Object"),
        ("The concept needs no identity", "YES — Create Value Object"),
        ("The concept represents a measurement or amount", "YES — Create Value Object"),
        ("The concept has a unique identity", "NO — Use Entity instead"),
        ("The concept changes state over time", "NO — Use Entity instead"),
    ]),
    ("Should I create a Domain Event?", [
        ("A state change occurred in an aggregate", "YES — Create Domain Event"),
        ("Something significant happened in the domain", "YES — Create Domain Event"),
        ("Other domains need to react to a change", "YES — Create Domain Event + Integration Event"),
        ("A read operation was performed", "NO — Queries do not produce events"),
        ("An internal implementation detail changed", "NO — Only domain-significant changes produce events"),
    ]),
    ("Should I create a Repository?", [
        ("An aggregate needs persistence", "YES — Create Repository"),
        ("Multiple aggregates share a persistence mechanism", "YES — One Repository per Aggregate"),
        ("The concept is not an aggregate", "NO — Only Aggregates have Repositories"),
        ("The concept is a value object", "NO — Value Objects are persisted with their parent Entity"),
        ("The concept is an event", "NO — Events use Event Store"),
    ]),
    ("Should I create a Factory?", [
        ("Object creation involves complex logic", "YES — Create Factory"),
        ("Object creation requires multiple steps", "YES — Create Factory"),
        ("Object creation enforces invariants", "YES — Create Factory"),
        ("Object creation is trivial", "NO — Use constructor directly"),
        ("Object creation requires external data", "YES — Create Factory with injected dependencies"),
    ]),
    ("Should I create a Policy?", [
        ("A business rule spans multiple entities", "YES — Create Policy"),
        ("A business rule depends on external conditions", "YES — Create Policy"),
        ("A business rule determines event routing", "YES — Create Policy"),
        ("A business rule belongs to a single entity", "NO — Add to Entity"),
        ("A business rule is infrastructure-related", "NO — Create Infrastructure Rule"),
    ]),
    ("Is this a Cross-Domain Operation?", [
        ("The operation involves entities from different domains", "YES — Use published interfaces or events"),
        ("The operation modifies data in multiple domains", "YES — Use Saga or Domain Events"),
        ("The operation reads data from multiple domains", "YES — Use API Composition or separate queries"),
        ("The operation is within a single domain", "NO — Use direct aggregate interaction"),
        ("The operation crosses a bounded context", "YES — Use Anti-Corruption Layer"),
    ]),
    ("Is this Dependency Valid?", [
        ("The dependency flows downward in the canonical graph", "YES — Valid dependency"),
        ("The dependency is semantically necessary", "YES — Valid dependency"),
        ("The dependency has architectural justification", "YES — Valid dependency"),
        ("The dependency is stable long-term", "YES — Valid dependency"),
        ("The dependency creates a cycle", "NO — INVALID, refactor"),
        ("The dependency points upward", "NO — INVALID, use dependency inversion"),
    ]),
]

for question, branches in decision_trees:
    add_decision_tree(question, branches)
    doc.add_paragraph()

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# PART XIV — PROMPTING CONSTITUTION
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("PART XIV — Prompting Constitution", level=1)
section_break()

doc.add_heading("14.1 How Prompts Shall Be Interpreted", level=2)
prompt_rules = [
    ("PROMPT-001", "Every prompt SHALL be interpreted in the context of the RIOS architecture. Generic software engineering interpretations are insufficient."),
    ("PROMPT-002", "When a prompt mentions a domain concept, the AI SHALL verify the concept's definition in the CTD before proceeding."),
    ("PROMPT-003", "When a prompt requests a feature, the AI SHALL map it to architectural concepts before implementation."),
    ("PROMPT-004", "When a prompt is ambiguous, the AI SHALL resolve ambiguity by consulting the architecture, not by making assumptions."),
    ("PROMPT-005", "When a prompt conflicts with the architecture, the AI SHALL report the conflict and recommend the architecturally correct approach."),
]
for num, text in prompt_rules:
    add_rule(num, text)

doc.add_heading("14.2 How Ambiguity Shall Be Resolved", level=2)
ambiguity_rules = [
    ("AMB-001", "When a term could refer to multiple concepts, consult the CTD first."),
    ("AMB-002", "When a domain is not specified, determine the domain from the context and the Domain Ownership Matrix."),
    ("AMB-003", "When a pattern is not specified, follow the DDD pattern used in existing RIOS code."),
    ("AMB-004", "When a naming convention is unclear, follow the naming rules in Part IV."),
    ("AMB-005", "When the scope is unclear, default to the narrowest scope that satisfies the request."),
    ("AMB-006", "When multiple valid approaches exist, list all approaches and recommend the one that best preserves architecture."),
]
for num, text in ambiguity_rules:
    add_rule(num, text)

doc.add_heading("14.3 How Missing Information Shall Be Handled", level=2)
missing_rules = [
    ("MIS-001", "When architectural specifications are missing for a requested feature, the AI SHALL raise the gap explicitly."),
    ("MIS-002", "When ownership is unclear, the AI SHALL ask the user to clarify before proceeding."),
    ("MIS-003", "When a semantic contract is undefined, the AI SHALL propose one for review."),
    ("MIS-004", "When tests are missing for existing code, the AI SHALL note the gap."),
    ("MIS-005", "When documentation is incomplete, the AI SHALL generate documentation as part of the task."),
    ("MIS-006", "NEVER assume architectural intent. If the architecture is silent, raise the question."),
]
for num, text in missing_rules:
    add_rule(num, text)

doc.add_heading("14.4 How Assumptions Shall Be Documented", level=2)
assume_rules = [
    ("ASM-001", "Every assumption SHALL be explicitly listed before implementation."),
    ("ASM-002", "Every assumption SHALL be validated against the architecture."),
    ("ASM-003", "Invalid assumptions SHALL be rejected, not silently corrected."),
    ("ASM-004", "Assumptions SHALL be documented in the implementation report."),
    ("ASM-005", "Assumptions that affect architecture SHALL trigger an ADR proposal."),
]
for num, text in assume_rules:
    add_rule(num, text)

doc.add_heading("14.5 How Architecture Conflicts Shall Be Reported", level=2)
conflict_rules = [
    ("CON-001", "When a request conflicts with the architecture, the AI SHALL report the conflict explicitly."),
    ("CON-002", "The report SHALL identify the specific architectural law, rule, or contract being violated."),
    ("CON-003", "The report SHALL propose the architecturally correct alternative."),
    ("CON-004", "The AI SHALL NOT implement the conflicting request without explicit user override."),
    ("CON-005", "User overrides SHALL be documented in an ADR."),
]
for num, text in conflict_rules:
    add_rule(num, text)

doc.add_heading("14.6 How Implementation Uncertainty Shall Be Managed", level=2)
uncertainty_rules = [
    ("UNC-001", "When uncertain about the correct implementation approach, the AI SHALL list alternatives with pros and cons."),
    ("UNC-002", "When uncertain about domain boundaries, the AI SHALL consult the Domain Ownership Matrix."),
    ("UNC-003", "When uncertain about aggregate boundaries, the AI SHALL consult the Domain Model Specification."),
    ("UNC-004", "When uncertain about event design, the AI SHALL consult existing domain events for patterns."),
    ("UNC-005", "When uncertainty cannot be resolved, the AI SHALL ask the user for guidance before proceeding."),
]
for num, text in uncertainty_rules:
    add_rule(num, text)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# PART XV — AI SELF-AUDIT FRAMEWORK
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("PART XV — AI Self-Audit Framework", level=1)
section_break()

doc.add_heading("15.1 Self-Audit Principle", level=2)
doc.add_paragraph(
    "Before completing ANY task, the AI agent SHALL execute a comprehensive self-audit. "
    "The self-audit verifies that every aspect of the implementation preserves the architectural "
    "integrity of RIOS. No task is complete until the self-audit passes."
)

doc.add_heading("15.2 Self-Audit Questions", level=2)
doc.add_paragraph("The AI SHALL ask itself the following questions before presenting any result:")

audit_categories = [
    ("Architecture Preservation", [
        "Did I preserve the Master Architecture Blueprint?",
        "Did I preserve all architectural invariants?",
        "Did I preserve the canonical dependency graph?",
        "Did I follow the CQRS mandate?",
        "Did I preserve the append-only event stream?",
        "Did I preserve the read-only identity projection?",
    ]),
    ("DDD Preservation", [
        "Did I preserve aggregate boundaries?",
        "Did I preserve entity identity?",
        "Did I preserve value object immutability?",
        "Did I preserve domain event immutability?",
        "Did I preserve repository patterns?",
        "Did I use factories for complex creation?",
        "Did I keep domain logic in the domain layer?",
    ]),
    ("Ownership Preservation", [
        "Did I respect domain ownership?",
        "Did I use published interfaces for cross-domain interaction?",
        "Did I verify every concept's owner?",
        "Did I avoid modifying code owned by other domains?",
    ]),
    ("Semantic Preservation", [
        "Did I use canonical terminology consistently?",
        "Did I preserve semantic contracts?",
        "Did I avoid defining concepts in multiple locations?",
        "Did I ensure evidence precedes claims?",
    ]),
    ("Traceability Preservation", [
        "Did I maintain traceability to architecture?",
        "Did I reference relevant ADRs?",
        "Did I update traceability for new concepts?",
        "Did I avoid removing traceability links?",
    ]),
    ("Technology Independence", [
        "Did I avoid embedding implementation details in architecture?",
        "Did I keep technology choices in the engineering layer?",
        "Did I preserve semantic contracts independent of technology?",
    ]),
    ("Coupling Analysis", [
        "Did I introduce unwanted coupling?",
        "Did I create direct dependencies between unrelated domains?",
        "Did I share mutable state between modules?",
        "Did I violate the dependency direction?",
    ]),
    ("Duplication Analysis", [
        "Did I introduce duplicate concepts?",
        "Did I duplicate logic across domains?",
        "Did I duplicate definitions across documents?",
    ]),
    ("Dependency Analysis", [
        "Did I introduce circular dependencies?",
        "Did I create reverse dependencies?",
        "Did I make stable domains depend on variable domains?",
    ]),
    ("Complexity Analysis", [
        "Did I simplify unnecessarily (loss of capability)?",
        "Did I overengineer (add unnecessary complexity)?",
        "Did I add premature optimization?",
        "Did I add unnecessary abstractions?",
    ]),
]

for category, questions in audit_categories:
    p = doc.add_paragraph()
    run = p.add_run(f"Category: {category}")
    run.bold = True
    run.font.color.rgb = MEDIUM_BLUE
    run.font.size = Pt(12)
    for q in questions:
        add_bullet(f"☐ {q}")

doc.add_heading("15.3 Self-Audit Scoring", level=2)
create_table(
    ["Result", "Meaning", "Action"],
    [
        [("ALL PASS", True), ("No violations detected", False), ("Proceed with task completion", False)],
        [("WARNING", True), ("Potential issue detected", False), ("Document the concern and proceed with caution", False)],
        [("FAIL", True), ("Violation detected", False), ("Fix the violation before presenting results", False)],
        [("CRITICAL FAIL", True), ("Architecture violation detected", False), ("STOP. Do not present. Fix immediately.", False)],
    ]
)

doc.add_heading("15.4 Self-Audit Reporting", level=2)
doc.add_paragraph(
    "The AI SHALL include a self-audit summary with every implementation output. "
    "The summary SHALL include: audit questions answered, results (Pass/Warning/Fail), "
    "and any warnings or concerns. This summary becomes part of the implementation record."
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# GLOSSARY
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("Glossary", level=1)
section_break()

glossary_terms = [
    ("ADR", "Architecture Decision Record — A document recording a significant architectural decision with context, alternatives, decision, and consequences."),
    ("Aggregate", "A cluster of domain objects treated as a single unit for data changes, guaranteeing semantic consistency through an Aggregate Root."),
    ("Aggregate Root", "The entry point to an Aggregate, responsible for enforcing invariants and providing a consistent interface."),
    ("Anti-Corruption Layer", "A translation layer that prevents bounded contexts from directly coupling to each other's internal models."),
    ("ATLAS", "The navigational guide to the RIOS architecture, providing a map of all domains, concepts, and relationships."),
    ("Bounded Context", "A clear boundary within which a particular domain model is defined and applicable."),
    ("Canonical Terminology Dictionary (CTD)", "The authoritative vocabulary of RIOS defining all domain terms."),
    ("Command", "In CQRS, a request to change system state. Commands are processed by the Write Model."),
    ("CQRS", "Command Query Responsibility Segregation — The pattern separating state-changing operations (Commands) from state-reading operations (Queries)."),
    ("Domain", "A bounded area of knowledge and responsibility within RIOS."),
    ("Domain Event", "An immutable, timestamped record of a significant state change in the domain."),
    ("Domain Model Specification (DMS)", "The document defining the detailed domain model for each RIOS domain."),
    ("Domain Ownership Matrix (DOM)", "The document defining which domain owns which concepts and responsibilities."),
    ("Domain Dependency Matrix (DDM)", "The document defining the canonical dependency graph between domains."),
    ("Entity", "A domain object with a unique identity that persists through state changes."),
    ("Evidence", "Observable, verifiable information used to support, challenge, or refine scientific claims."),
    ("Event Stream", "The append-only sequence of all domain events, representing the complete history of the system."),
    ("Factory", "A DDD pattern encapsulating complex object creation logic."),
    ("Identity Projection", "The synthesized, point-in-time state of a researcher's career, built from the event stream."),
    ("Invariant", "A business rule that must always be true within an aggregate boundary."),
    ("Knowledge", "Generalized understanding derived from interpreted evidence."),
    ("Master Architecture Blueprint (MAB)", "The supreme architectural authority document for RIOS."),
    ("Policy", "A DDD pattern encapsulating business rules that span multiple entities or aggregates."),
    ("Query", "In CQRS, a request to read system state. Queries are processed by the Read Model."),
    ("Repository", "A DDD pattern providing the illusion of an in-memory collection of aggregates."),
    ("Research Identity", "The persistent, emergent representation of a researcher's intellectual direction derived from evidence."),
    ("Research Object", "A versioned, inspectable artifact that communicates or preserves scientific knowledge."),
    ("Semantic Contract", "A technology-independent definition of an interface's purpose, input, output, consistency, and ownership."),
    ("Traceability Matrix", "The document linking architectural concepts to requirements, implementations, and ADRs."),
    ("Value Object", "An immutable domain object defined by its attributes rather than identity."),
]

create_table(
    ["Term", "Definition"],
    [[(term, False), (definition, False)] for term, definition in glossary_terms]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# APPENDIX A — ARCHITECTURE COMPLIANCE MATRIX
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("Appendix A — Architecture Compliance Matrix", level=1)
section_break()

doc.add_paragraph(
    "This matrix maps every architectural concept to its verification method and governing authority."
)

create_table(
    ["Concept", "Domain", "Verification Method", "Governing Authority"],
    [
        [("Research Identity", False), ("Identity", False), ("Deterministic Replay", False), ("MAB + Volume I", False)],
        [("Research Agenda", False), ("Identity", False), ("Invariant Test", False), ("Volume I", False)],
        [("Research Area", False), ("Identity", False), ("Invariant Test", False), ("Volume I", False)],
        [("Research Question", False), ("Identity", False), ("Invariant Test", False), ("Volume I", False)],
        [("Evidence", False), ("Knowledge", False), ("Immutability Test", False), ("Volume II", False)],
        [("Knowledge", False), ("Knowledge", False), ("Evidence Link Verification", False), ("Volume II", False)],
        [("Narrative", False), ("Narrative", False), ("Semantic Contract Test", False), ("Volume III", False)],
        [("Publication", False), ("Publication", False), ("Entity Identity Test", False), ("Volume IV", False)],
        [("Visualization", False), ("Visualization", False), ("Cognitive Justification Test", False), ("Volume V", False)],
        [("Motion", False), ("Motion", False), ("Cognitive Service Test", False), ("Volume VI", False)],
        [("Engineering", False), ("Engineering", False), ("Cross-Cutting Compliance", False), ("Volume VII", False)],
        [("Evolution", False), ("Evolution", False), ("Version Compliance", False), ("Volume VIII", False)],
        [("Domain Events", False), ("All", False), ("Immutability Test", False), ("DDM + DMS", False)],
        [("Semantic Contracts", False), ("All", False), ("Contract Test", False), ("DMS + CTD", False)],
    ]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# APPENDIX B — REVIEW SCORECARD TEMPLATE
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("Appendix B — Review Scorecard Template", level=1)
section_break()

create_table(
    ["Review Dimension", "Weight", "Score (1-5)", "Weighted Score", "Comments"],
    [
        [("Architecture Compliance", True), ("25%", False), ("___", False), ("___", False), ("", False)],
        [("DDD Compliance", True), ("20%", False), ("___", False), ("___", False), ("", False)],
        [("Semantic Correctness", True), ("15%", False), ("___", False), ("___", False), ("", False)],
        [("Dependency Compliance", True), ("10%", False), ("___", False), ("___", False), ("", False)],
        [("Code Quality", True), ("10%", False), ("___", False), ("___", False), ("", False)],
        [("Test Coverage", True), ("10%", False), ("___", False), ("___", False), ("", False)],
        [("Documentation", True), ("5%", False), ("___", False), ("___", False), ("", False)],
        [("Security", True), ("5%", False), ("___", False), ("___", False), ("", False)],
        [("TOTAL", True), ("100%", False), ("", False), ("___", False), ("Pass threshold: ≥ 3.5", False)],
    ]
)

doc.add_paragraph()
doc.add_paragraph("Scoring Guide:")
create_table(
    ["Score", "Meaning", "Action"],
    [
        [("5", False), ("Exemplary — Sets a standard for the codebase", False), ("Accept", False)],
        [("4", False), ("Good — Meets all requirements with minor opportunities", False), ("Accept", False)],
        [("3", False), ("Acceptable — Meets requirements with some concerns", False), ("Accept with notes", False)],
        [("2", False), ("Below Standard — Significant improvements needed", False), ("Revise", False)],
        [("1", False), ("Unacceptable — Critical issues present", False), ("Reject", False)],
    ]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# APPENDIX C — AI SELF-AUDIT TEMPLATE
# ═════════════════════════════════════════════════════════════════════════════

doc.add_heading("Appendix C — AI Self-Audit Template", level=1)
section_break()

doc.add_paragraph("Task Description: _______________________________________________")
doc.add_paragraph("Date: ______________________ AI Agent: ______________________")
doc.add_paragraph()

audit_template = [
    ("Architecture Preservation", [
        "Preserved MAB compliance",
        "Preserved architectural invariants",
        "Preserved canonical dependency graph",
        "Preserved CQRS mandate",
        "Preserved append-only event stream",
        "Preserved read-only identity projection",
    ]),
    ("DDD Preservation", [
        "Preserved aggregate boundaries",
        "Preserved entity identity",
        "Preserved value object immutability",
        "Preserved domain event immutability",
        "Preserved repository patterns",
        "Kept domain logic in domain layer",
    ]),
    ("Ownership Preservation", [
        "Respected domain ownership",
        "Used published interfaces",
        "Verified concept owners",
        "Avoided cross-domain violations",
    ]),
    ("Semantic Preservation", [
        "Used canonical terminology",
        "Preserved semantic contracts",
        "No duplicate concept definitions",
        "Evidence precedes claims",
    ]),
    ("Quality Checks", [
        "No circular dependencies introduced",
        "No shared mutable state introduced",
        "No technology leakage into architecture",
        "No unnecessary complexity added",
        "No simplification that loses capability",
        "Traceability maintained",
    ]),
]

for category, items in audit_template:
    p = doc.add_paragraph()
    run = p.add_run(category)
    run.bold = True
    run.font.color.rgb = MEDIUM_BLUE
    for item in items:
        add_bullet(f"☐ {item} — PASS / WARNING / FAIL")

doc.add_paragraph()
doc.add_paragraph("Overall Audit Result: ☐ ALL PASS   ☐ WARNING   ☐ FAIL   ☐ CRITICAL FAIL")
doc.add_paragraph()
doc.add_paragraph("Notes: _______________________________________________")
doc.add_paragraph("Reviewer: _______________________________________________")
doc.add_paragraph("Date: _______________________________________________")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# FINAL PAGE
# ═════════════════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RIOS Claude Code Constitution v1.0")
run.font.size = Pt(24)
run.font.color.rgb = DARK_NAVY
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("End of Document")
run.font.size = Pt(16)
run.font.color.rgb = MEDIUM_BLUE

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "This Constitution is the permanent engineering authority for all AI coding agents\n"
    "and human engineers working on the Research Identity Operating System.\n\n"
    "Architecture owns semantics.\n"
    "Engineering owns realization.\n"
    "Implementation never defines architecture.\n\n"
    "If architecture and implementation conflict,\n"
    "architecture ALWAYS wins."
)
run.font.size = Pt(12)
run.font.color.rgb = DARK_GRAY
run.italic = True

# ═════════════════════════════════════════════════════════════════════════════
# SAVE
# ═════════════════════════════════════════════════════════════════════════════

output_path = "RIOS_Claude_Code_Constitution_v1.0.docx"
doc.save(output_path)
print(f"\n{'='*60}")
print(f"RIOS Claude Code Constitution v1.0 generated successfully!")
print(f"Output: {output_path}")
print(f"{'='*60}")