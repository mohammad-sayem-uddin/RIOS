#!/usr/bin/env python3
"""RIOS Atlas Generator — Produces the official RIOS Atlas as .docx"""
import os, math
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
try:
    from PIL import Image, ImageDraw, ImageFont
    PIL = True
except: PIL = False

DD = Path("atlas_diagrams"); DD.mkdir(exist_ok=True)
# Colors
NAVY="#1B2A4A"; STEEL="#2D5F8A"; AMBER="#E8913A"; GREEN="#2E7D32"
ORANGE="#F57F17"; RED="#C62828"; LGRAY="#F5F5F5"; MGRAY="#90A4AE"
DC={"Identity":"#1B2A4A","Knowledge":"#2D5F8A","KnowComm":"#1565C0",
    "ScholComm":"#00838F","Vis":"#2E7D32","Motion":"#6A1B9A",
    "Platform":"#E65100","Impl":"#37474F"}

def _f(s=14):
    if not PIL: return None
    for p in ["/System/Library/Fonts/Helvetica.ttc","/Library/Fonts/Arial.ttf"]:
        try: return ImageFont.truetype(p,s)
        except: pass
    return ImageFont.load_default()

def box(d,x,y,w,h,t,fc="#2D5F8A",tc="#FFF",fs=12,bc="#1B2A4A",r=8):
    d.rounded_rectangle([x,y,x+w,y+h],radius=r,fill=fc,outline=bc,width=2)
    f=_f(fs); lines=[t]
    bb=d.textbbox((0,0),t,font=f)
    if bb[2]-bb[0]>w-16:
        lines=[]; line=""
        for word in t.split():
            test=line+" "+word if line else word
            b2=d.textbbox((0,0),test,font=f)
            if b2[2]-b2[0]>w-16:
                if line: lines.append(line)
                line=word
            else: line=test
        if line: lines.append(line)
    th=len(lines)*(fs+4); sy=y+(h-th)//2
    for i,l in enumerate(lines):
        bb=d.textbbox((0,0),l,font=f); lw=bb[2]-bb[0]
        d.text((x+(w-lw)//2,sy+i*(fs+4)),l,fill=tc,font=f)

def arr(d,x1,y1,x2,y2,c="#1B2A4A",w=2):
    d.line([(x1,y1),(x2,y2)],fill=c,width=w)
    a=math.atan2(y2-y1,x2-x1); al=10
    for da in[0.4,-0.4]:
        d.line([(x2,y2),(x2-al*math.cos(a+da),y2-al*math.sin(a+da))],fill=c,width=w)

def title(d,w,t,st=None,y=15,fs=18):
    f=_f(fs); bb=d.textbbox((0,0),t,font=f); tw=bb[2]-bb[0]
    d.text(((w-tw)//2,y),t,fill=NAVY,font=f)
    if st:
        sf=_f(11); bb2=d.textbbox((0,0),st,font=sf); tw2=bb2[2]-bb2[0]
        d.text(((w-tw2)//2,y+fs+4),st,fill=MGRAY,font=sf)

def legend(d,x,y,items,fs=10):
    f=_f(fs)
    for i,(lb,co) in enumerate(items):
        d.rectangle([x,y+i*20,x+14,y+i*20+14],fill=co,outline="#333")
        d.text((x+20,y+i*20),lb,fill="#212121",font=f)

def newimg(fn,w=1200,h=800):
    if not PIL: return None,None
    img=Image.new("RGB",(w,h),"#FFF"); d=ImageDraw.Draw(img)
    d.rectangle([2,2,w-3,h-3],outline=NAVY,width=1); return img,d

def save(img,fn):
    p=DD/fn; img.save(str(p),"PNG",dpi=(150,150)); return str(p)

# ═══════════════════════════════════════════════════════════════════
# DIAGRAM GENERATORS
# ═══════════════════════════════════════════════════════════════════

def d_FOUND_001():
    img,d=newimg("ATL-FOUND-001.png",1200,800)
    title(d,1200,"ATL-FOUND-001: RIOS System Overview","Top-Level Architecture Vision")
    box(d,450,50,300,50,"Academic Evaluation Ecosystem",fc=LGRAY,bc=MGRAY,tc="#212121")
    d.rounded_rectangle([100,140,1100,700],radius=12,fill="#ECEFF1",outline=NAVY,width=2)
    d.text((480,148),"RIOS Platform",fill=NAVY,font=_f(14))
    doms=[("Identity",120,200),("Knowledge",360,200),("Know. Comm.",600,200),("Schol. Comm.",840,200),
          ("Scientific Vis.",120,340),("Cog. Motion",360,340),("Platform Eng.",600,340),("Implementation",840,340)]
    for n,x,y in doms:
        k=n.replace(". ","").replace(" ","")
        co=DC.get({"Identity":"Identity","Knowledge":"Knowledge","KnowComm":"KnowComm",
                    "ScholComm":"ScholComm","ScientificVis":"Vis","CogMotion":"Motion",
                    "PlatformEng":"Platform","Implementation":"Impl"}.get(k,"Platform"),STEEL)
        # Map names properly
        cmap={"Identity":"Identity","Knowledge":"Knowledge","Know.Comm.":"KnowComm",
              "Schol.Comm.":"ScholComm","ScientificVis":"Vis","Cog.Motion":"Motion",
              "PlatformEng.":"Platform","Implementation":"Impl"}
        co2=DC.get(cmap.get(n, n),"#2D5F8A")
        box(d,x,y,200,70,n,fc=co2,fs=13)
    box(d,200,480,800,45,"Foundation: Governance · Terminology · Domain Model · Dependency Matrix · Ownership · Requirements",fc=NAVY,fs=10)
    box(d,200,540,800,45,"Standards: Editorial Standard · Requirement Taxonomy · Architecture Governance",fc=STEEL,fs=10)
    legend(d,130,610,[("Identity Domain",DC["Identity"]),("Knowledge Domain",DC["Knowledge"]),
        ("Communication Domains",DC["KnowComm"]),("Visualization Domain",DC["Vis"]),
        ("Motion Domain",DC["Motion"]),("Engineering Domains",DC["Platform"])])
    return save(img,"ATL-FOUND-001.png")

def d_FOUND_002():
    img,d=newimg("ATL-FOUND-002.png",1200,700)
    title(d,1200,"ATL-FOUND-002: Architecture Vision","RIOS Guiding Philosophy")
    box(d,400,80,400,60,"Research Identity as a\nLiving Scientific System",fc=NAVY,fs=15)
    for i,(t,x) in enumerate([("Research Before Software",150),("Knowledge Before Publication",500),("Meaning Before Implementation",850)]):
        box(d,x,200,220,60,t,fc=STEEL,fs=12); arr(d,600,140,x+110,200)
    for i,(t,x) in enumerate([("Evidence Before Claims",150),("Architecture Before Technology",500),("Technology Serves Architecture",850)]):
        box(d,x,330,220,60,t,fc=AMBER,fs=12)
    box(d,300,450,600,50,"Architecture Serves Research",fc=NAVY,fs=14)
    for x in[260,600,940]: arr(d,x,390,600,450)
    attrs=["Coherent","Verifiable","Interpretable","Maintainable","Extensible","Technology-Independent"]
    for i,a in enumerate(attrs):
        box(d,100+(i%3)*350,540+(i//3)*50,300,40,a,fc=LGRAY,bc=MGRAY,tc="#212121",fs=11)
    return save(img,"ATL-FOUND-002.png")

def d_FOUND_003():
    img,d=newimg("ATL-FOUND-003.png",1200,750)
    title(d,1200,"ATL-FOUND-003: Architecture Principles")
    ps=[("AP-001","Research Primacy","Research direction is the primary driver of all architectural decisions"),
        ("AP-002","Semantic Integrity","Scientific meaning is canonical and preserved at every layer"),
        ("AP-003","Evidence Authority","All claims require supporting evidence; evidence does not define"),
        ("AP-004","Knowledge Independence","Knowledge exists before publication; it has its own lifecycle"),
        ("AP-005","Identity Persistence","Research Identity survives technological and career transitions"),
        ("AP-006","Architectural Purity","Architecture is independent of implementation technology"),
        ("AP-007","Domain Autonomy","Each domain owns its models, terminology, and business rules"),
        ("AP-008","Open Knowledge","Knowledge dissemination is a fundamental scientific obligation"),
        ("AP-009","Trust Through Evidence","Academic credibility is earned through verifiable evidence"),
        ("AP-010","Cognitive Accessibility","Architecture minimizes cognitive load for all stakeholders")]
    y=60
    for pid,pt,pd in ps:
        box(d,50,y,100,38,pid,fc=NAVY,fs=10)
        box(d,170,y,250,38,pt,fc=STEEL,fs=10)
        box(d,440,y,720,38,pd,fc=LGRAY,bc=MGRAY,tc="#212121",fs=9)
        y+=52
    return save(img,"ATL-FOUND-003.png")

def d_FOUND_004():
    img,d=newimg("ATL-FOUND-004.png",1000,850)
    title(d,1000,"ATL-FOUND-004: Architecture Layers")
    layers=[("Strategic Layer","Identity Architecture",NAVY),("Knowledge Layer","Knowledge Architecture","#1565C0"),
        ("Communication Layer","Knowledge + Scholarly Communication","#00838F"),("Presentation Layer","Scientific Visualization",GREEN),
        ("Motion Layer","Cognitive Motion","#6A1B9A"),("Platform Layer","Platform Engineering",ORANGE),
        ("Implementation Layer","Implementation Architecture","#37474F"),("Foundation Layer","Governance, Terminology, Specifications",MGRAY)]
    y=60
    for nm,desc,co in layers:
        d.rounded_rectangle([100,y,900,y+65],radius=10,fill=co,outline=NAVY,width=1)
        d.text((130,y+8),nm,fill="#FFF",font=_f(15))
        d.text((130,y+35),desc,fill="#E0E0E0",font=_f(11))
        if y>60: arr(d,500,y,500,y-10,c=NAVY)
        y+=80
    return save(img,"ATL-FOUND-004.png")

def d_FOUND_005():
    img,d=newimg("ATL-FOUND-005.png",1200,750)
    title(d,1200,"ATL-FOUND-005: Domain Overview","Eight Architectural Domains")
    dis=[("Identity","Who is the researcher?\nVision, agenda, philosophy, evidence",50,70,"Identity"),
         ("Knowledge","What does the researcher know?\nConcepts, claims, evidence, ontology",420,70,"Knowledge"),
         ("Know. Communication","How is knowledge communicated?\nNarrative, storytelling, research stories",790,70,"KnowComm"),
         ("Schol. Communication","How is scholarship shared?\nPublications, peer review, impact",50,280,"ScholComm"),
         ("Scientific Visualization","How is research visualized?\nDiagrams, knowledge graphs, dashboards",420,280,"Vis"),
         ("Cognitive Motion","How does identity evolve?\nDirection, momentum, trajectory",790,280,"Motion"),
         ("Platform Engineering","How is the platform built?\nServices, infrastructure, observability",50,490,"Platform"),
         ("Implementation","How is it realized?\nRepository, workflow, deployment, CI/CD",420,490,"Impl")]
    for nm,desc,x,y,k in dis:
        box(d,x,y,320,160,"",fc=DC[k],fs=12)
        d.text((x+15,y+10),nm,fill="#FFF",font=_f(15))
        d.text((x+15,y+60),desc,fill="#E0E0E0",font=_f(10))
    box(d,790,490,320,160,"",fc=NAVY)
    d.text((805,500),"Foundation\nArchitecture",fill="#FFF",font=_f(15))
    d.text((805,555),"Governance, terminology,\nspecifications, standards,\ndependency matrix",fill="#E0E0E0",font=_f(10))
    return save(img,"ATL-FOUND-005.png")

def d_FOUND_006():
    img,d=newimg("ATL-FOUND-006.png",1000,700)
    title(d,1000,"ATL-FOUND-006: Architecture Stack")
    stk=[("Strategic",["Identity"],NAVY),("Knowledge",["Knowledge"],"#1565C0"),
         ("Communication",["Knowledge Communication","Scholarly Communication"],"#00838F"),
         ("Presentation",["Scientific Visualization"],GREEN),("Evolution",["Cognitive Motion"],"#6A1B9A"),
         ("Platform",["Platform Engineering"],ORANGE),("Realization",["Implementation"],"#37474F"),
         ("Governance",["Foundation Architecture"],MGRAY)]
    y=60
    for ln,doms,co in stk:
        d.rounded_rectangle([80,y,920,y+55],radius=8,fill=co,outline=NAVY,width=1)
        d.text((100,y+5),ln,fill="#FFF",font=_f(13))
        d.text((100,y+28),"  ·  ".join(doms),fill="#E0E0E0",font=_f(10))
        if y>60: arr(d,500,y,500,y-8)
        y+=70
    return save(img,"ATL-FOUND-006.png")

def d_FOUND_007():
    img,d=newimg("ATL-FOUND-007.png",1200,750)
    title(d,1200,"ATL-FOUND-007: Dependency Architecture")
    pos={"Identity":(450,70,200,50),"Knowledge":(100,200,200,50),"Know. Comm.":(450,200,200,50),
         "Schol. Comm.":(800,200,200,50),"Scientific Vis.":(100,350,200,50),"Cog. Motion":(450,350,200,50),
         "Platform Eng.":(800,350,200,50),"Implementation":(450,500,200,50),"Foundation":(450,630,200,50)}
    cm={"Identity":"Identity","Knowledge":"Knowledge","Know. Comm.":"KnowComm","Schol. Comm.":"ScholComm",
        "Scientific Vis.":"Vis","Cog. Motion":"Motion","Platform Eng.":"Platform","Implementation":"Impl","Foundation":None}
    for nm,(x,y,w,h) in pos.items():
        k=cm.get(nm); co=DC.get(k,MGRAY) if k else MGRAY
        box(d,x,y,w,h,nm,fc=co,fs=12)
    deps=[("Identity","Knowledge"),("Identity","Know. Comm."),("Knowledge","Schol. Comm."),
          ("Knowledge","Scientific Vis."),("Identity","Cog. Motion"),("Know. Comm.","Schol. Comm."),
          ("Platform Eng.","Implementation")]
    for s,t in deps:
        sx,sy,sw,sh=pos[s]; dx,dy,dw,dh=pos[t]
        arr(d,sx+sw//2,sy+sh,dx+dw//2,dy,c="#455A64")
    return save(img,"ATL-FOUND-007.png")

def d_FOUND_008():
    img,d=newimg("ATL-FOUND-008.png",1200,800)
    title(d,1200,"ATL-FOUND-008: System Capability Map")
    caps={"Identity":["Research Vision","Research Agenda","Research Areas","Research Questions","Philosophy","Evidence"],
          "Knowledge":["Concepts","Claims","Evidence","Ontology","Reasoning","Knowledge Graph"],
          "Communication":["Narrative","Storytelling","Research Stories","Academic Profiles","Peer Review","Impact"],
          "Visualization":["Diagrams","Knowledge Graphs","Dashboards","Ontology Maps","Timelines","Metrics"],
          "Motion":["Direction","Momentum","Trajectory","Milestones","Evolution","Transitions"],
          "Platform":["Services","APIs","Database","Search","Auth","Observability"],
          "Implementation":["Repository","CI/CD","Testing","Deployment","Monitoring","Workflow"]}
    y=60
    for dom,items in caps.items():
        box(d,30,y,180,40,dom,fc=NAVY,fs=11)
        for i,it in enumerate(items):
            box(d,240+i*155,y,145,40,it,fc=LGRAY,bc=STEEL,tc="#212121",fs=9)
        y+=60
    box(d,30,y+10,180,40,"Foundation",fc=MGRAY,fs=11)
    for i,it in enumerate(["Governance","Terminology","Domain Model","Dep. Matrix","Ownership","Requirements"]):
        box(d,240+i*155,y+10,145,40,it,fc=LGRAY,bc=MGRAY,tc="#212121",fs=9)
    return save(img,"ATL-FOUND-008.png")

def d_FOUND_009():
    img,d=newimg("ATL-FOUND-009.png",1200,800)
    title(d,1200,"ATL-FOUND-009: System Context Diagram","C4 Level 0")
    for nm,x,y in [("Researcher",100,120),("Admissions Committee",100,300),("Faculty Reviewer",100,480),("AI Coding Agent",100,620)]:
        box(d,x,y,160,55,nm,fc=LGRAY,bc=AMBER,tc="#212121",fs=11)
    d.rounded_rectangle([350,100,900,680],radius=15,fill="#E3F2FD",outline=NAVY,width=3)
    d.text((520,108),"RIOS Platform",fill=NAVY,font=_f(18))
    for nm,x,y in [("Identity",380,170),("Knowledge",600,170),("Communication",380,300),("Visualization",600,300),("Motion",380,430),("Platform",600,430),("Implementation",490,560)]:
        k={"Communication":"KnowComm"}.get(nm,nm)
        box(d,x,y,180,50,nm,fc=DC.get(k,STEEL),fs=12)
    for nm,x,y in [("Academic Databases",980,200),("Publication Repos",980,350),("Citation Networks",980,500)]:
        box(d,x,y,180,50,nm,fc="#FFF3E0",bc=ORANGE,tc="#212121",fs=11)
    for sy,dy in[(150,195),(330,325),(510,455),(650,585)]: arr(d,260,sy,350,dy)
    for sy,dy in[(220,225),(370,375),(520,525)]: arr(d,900,sy,980,dy)
    return save(img,"ATL-FOUND-009.png")

def d_FOUND_010():
    img,d=newimg("ATL-FOUND-010.png",1000,650)
    title(d,1000,"ATL-FOUND-010: Architecture Hierarchy")
    lvls=[(400,60,200,42,"Architecture Vision",NAVY),(250,135,500,38,"Architecture Principles",STEEL),
          (100,210,300,38,"Domain Architecture","#1565C0"),(500,210,300,38,"Foundation Architecture","#00838F"),
          (50,290,180,38,"Identity",DC["Identity"]),(260,290,180,38,"Knowledge",DC["Knowledge"]),
          (470,290,180,38,"Platform Eng.",DC["Platform"]),(680,290,180,38,"Implementation",DC["Impl"]),
          (50,370,180,38,"Comm. Domains",DC["KnowComm"]),(260,370,180,38,"Visualization",DC["Vis"]),
          (470,370,180,38,"Motion",DC["Motion"]),(680,370,180,38,"Governance",MGRAY)]
    for x,y,w,h,t,co in lvls: box(d,x,y,w,h,t,fc=co,fs=10)
    return save(img,"ATL-FOUND-010.png")

def d_FOUND_011():
    img,d=newimg("ATL-FOUND-011.png",1200,450)
    title(d,1200,"ATL-FOUND-011: Architecture Flow")
    flow=[("Research Direction",NAVY),("Knowledge Construction","#1565C0"),("Communication","#00838F"),
          ("Visualization",GREEN),("Evolution & Motion","#6A1B9A"),("Platform & Infra",ORANGE),("Implementation","#37474F")]
    x=30
    for nm,co in flow:
        box(d,x,130,145,70,nm,fc=co,fs=10)
        if x>30: arr(d,x-10,165,x,165)
        x+=163
    box(d,50,270,1100,45,"Foundation Architecture: Governance · Terminology · Specifications · Standards",fc=NAVY,fs=10)
    d.text((400,350),"↑ Verification & Conformance Feedback Loop ↑",fill=MGRAY,font=_f(11))
    return save(img,"ATL-FOUND-011.png")

# ── IDENTITY DOMAIN ──

def d_IDN_001():
    img,d=newimg("ATL-IDN-001.png",1200,700)
    title(d,1200,"ATL-IDN-001: Identity Domain Overview","Volume I — Identity Architecture")
    box(d,400,70,400,50,"Research Identity",fc=NAVY,fs=16)
    comps=[("Vision Engine","10-20yr",50,180),("Agenda Engine","5-10yr",270,180),("Area Manager","3-10yr",490,180),
           ("Question Registry","Permanent",710,180),("Philosophy Layer","Slow",930,180),
           ("Evidence Layer","Fast",50,350),("Representation","Medium",270,350),("Interface Layer","Stable",490,350),("Evolution Layer","V. Stable",710,350)]
    for nm,hz,x,y in comps:
        box(d,x,y,200,55,nm,fc=STEEL,fs=11)
        d.text((x+50,y+60),hz,fill=MGRAY,font=_f(9))
        arr(d,600,120,x+100,180,c="#90A4AE")
    box(d,250,470,700,45,"Meaning flows DOWN  ·  Interpretation flows UP",fc=LGRAY,bc=NAVY,tc="#212121",fs=11)
    box(d,80,550,480,80,"SHALL Own:\nVision, Agenda, Philosophy,\nValues, Direction, Evolution",fc=GREEN,fs=10)
    box(d,640,550,480,80,"SHALL NOT Own:\nPublications, Datasets,\nVisualizations, Motion, Implementation",fc=RED,fs=10)
    return save(img,"ATL-IDN-001.png")

def d_IDN_002():
    img,d=newimg("ATL-IDN-002.png",1000,650)
    title(d,1000,"ATL-IDN-002: Identity Bounded Context")
    d.rounded_rectangle([50,55,950,600],radius=12,fill="#E8EAF6",outline=NAVY,width=2)
    d.text((380,62),"Identity Bounded Context",fill=NAVY,font=_f(14))
    box(d,350,100,300,45,"Research Identity (Aggregate Root)",fc=NAVY,fs=12)
    ents=[("Research Vision",80,200),("Research Agenda",300,200),("Research Area",520,200),("Research Question",740,200),
          ("Philosophy",80,330),("Evidence Item",300,330),("Identity Version",520,330),("Identity Milestone",740,330)]
    for n,x,y in ents: box(d,x,y,180,42,n,fc=STEEL,fs=10)
    vos=[("Method Principle",80,450),("Scientific Value",300,450),("Research Motivation",520,450),("Evidence Reference",740,450)]
    for n,x,y in vos: box(d,x,y,180,42,n,fc=AMBER,fs=9)
    legend(d,80,540,[("Aggregate Root",NAVY),("Entity",STEEL),("Value Object",AMBER)])
    return save(img,"ATL-IDN-002.png")

def d_IDN_003():
    img,d=newimg("ATL-IDN-003.png",1200,450)
    title(d,1200,"ATL-IDN-003: Identity Lifecycle")
    sts=[("Nascent Identity",MGRAY),("Emerging Identity","#1565C0"),("Established Identity",GREEN),("Mature Identity",NAVY),("Legacy Identity","#6A1B9A")]
    x=50
    for nm,co in sts:
        box(d,x,130,170,70,nm,fc=co,fs=12)
        if x>50: arr(d,x-15,165,x,165)
        x+=225
    evts=[("First Research Question",135,260),("First Publication + Areas",360,260),("Agenda Clarity + Evidence",585,260),("Sustained Impact",810,260)]
    for t,x,y in evts: box(d,x,y,170,45,t,fc=LGRAY,bc=AMBER,tc="#212121",fs=9)
    d.text((150,370),"Events: AgendaCreated · AreaAdded · PhilosophyRevised · MilestoneRecorded · CareerTransition",fill=MGRAY,font=_f(10))
    return save(img,"ATL-IDN-003.png")

def d_IDN_004():
    img,d=newimg("ATL-IDN-004.png",1200,700)
    title(d,1200,"ATL-IDN-004: Identity Entity-Relationship Diagram")
    ents={"ResearchIdentity":(450,70,200,45),"ResearchVision":(100,200,170,40),"ResearchAgenda":(340,200,170,40),
          "ResearchArea":(580,200,170,40),"ResearchQuestion":(820,200,180,40),
          "Philosophy":(100,340,170,40),"Evidence":(340,340,170,40),"IdentityVersion":(580,340,170,40),
          "IdentityMilestone":(820,340,180,40),"MethodPrinciple":(100,480,170,40),
          "ScientificValue":(340,480,170,40),"ResearchMotivation":(580,480,170,40),"Representation":(820,480,180,40)}
    for nm,(x,y,w,h) in ents.items():
        co=NAVY if nm=="ResearchIdentity" else STEEL
        box(d,x,y,w,h,nm,fc=co,fs=10)
    rels=[("ResearchIdentity","ResearchVision"),("ResearchIdentity","ResearchAgenda"),("ResearchIdentity","ResearchArea"),
          ("ResearchIdentity","ResearchQuestion"),("ResearchArea","ResearchQuestion"),("ResearchIdentity","Philosophy"),
          ("ResearchIdentity","Evidence"),("ResearchIdentity","IdentityVersion"),("ResearchIdentity","IdentityMilestone"),
          ("Philosophy","MethodPrinciple"),("Philosophy","ScientificValue"),("ResearchIdentity","Representation")]
    for s,t in rels:
        sx,sy,sw,sh=ents[s]; dx,dy,dw,dh=ents[t]
        arr(d,sx+sw//2,sy+sh,dx+dw//2,dy,c="#78909C")
    d.text((100,570),"All relationships carry explicit semantic meaning (investigates, extends, derives_from, supports, evolves_into)",fill="#212121",font=_f(10))
    d.text((100,600),"Circular semantic dependencies are prohibited. Cardinality: 1 ──── * (one-to-many)",fill=MGRAY,font=_f(10))
    return save(img,"ATL-IDN-004.png")

def d_IDN_005():
    img,d=newimg("ATL-IDN-005.png",800,850)
    title(d,800,"ATL-IDN-005: Identity Architecture Blueprint","9-Layer Canonical Structure")
    lys=[("Research Identity",NAVY),("Research Vision","#1A237E"),("Research Agenda","#1565C0"),
         ("Research Areas","#0277BD"),("Research Questions","#00838F"),("Research Philosophy","#2E7D32"),
         ("Research Evidence","#558B2F"),("Identity Representation",AMBER),("External Communication","#E65100")]
    y=55
    for nm,co in lys:
        box(d,120,y,460,50,nm,fc=co,fs=13)
        if y>55: arr(d,350,y,350,y-12,c=NAVY)
        y+=75
    stab=["Ext. High","High","Med-High","Medium","High","Low","Medium","High","High"]
    for i,s in enumerate(stab):
        d.text((600,68+i*75),f"Stability: {s}",fill=MGRAY,font=_f(9))
    return save(img,"ATL-IDN-005.png")

def d_IDN_006():
    img,d=newimg("ATL-IDN-006.png",1200,650)
    title(d,1200,"ATL-IDN-006: Identity Context Map")
    box(d,450,80,300,55,"Identity Domain",fc=NAVY,fs=16)
    ctxs=[("Knowledge Domain",100,230,"Downstream Consumer"),("Comm. Domains",450,230,"Downstream Consumer"),
          ("Visualization",800,230,"Downstream Consumer"),("Motion Domain",100,420,"Downstream Consumer"),
          ("Platform Eng.",450,420,"Infrastructure Provider"),("Implementation",800,420,"Realization Provider")]
    for nm,x,y,rl in ctxs:
        box(d,x,y,200,55,nm,fc=STEEL,fs=11)
        d.text((x+20,y+60),rl,fill=MGRAY,font=_f(8))
        arr(d,600,135,x+100,230,c="#78909C")
    box(d,300,550,600,45,"Governance: MAB · CTD · DDM · DMS · AGS · DOM · RTS",fc=LGRAY,bc=NAVY,tc="#212121",fs=10)
    return save(img,"ATL-IDN-006.png")

def d_IDN_007():
    img,d=newimg("ATL-IDN-007.png",1200,550)
    title(d,1200,"ATL-IDN-007: Identity Capability Map")
    caps=[("Vision Management",["Define vision","Update vision","Version vision"]),
          ("Agenda Management",["Define agenda","Prioritize areas","Link to vision"]),
          ("Area Management",["Create areas","Archive areas","Maintain hierarchy"]),
          ("Question Management",["Register questions","Track evolution","Link across projects"]),
          ("Philosophy Management",["Define principles","State values","Document motivations"]),
          ("Evidence Management",["Associate evidence","Validate claims","Track sources"]),
          ("Representation",["Generate homepage","Sync profiles","Maintain consistency"]),
          ("Evolution Tracking",["Record milestones","Track transitions","Preserve history"])]
    y=55
    for cn,subs in caps:
        box(d,30,y,200,32,cn,fc=NAVY,fs=9)
        for i,sc in enumerate(subs):
            box(d,250+i*180,y,170,32,sc,fc=LGRAY,bc=STEEL,tc="#212121",fs=8)
        y+=45
    return save(img,"ATL-IDN-007.png")

def d_IDN_008():
    img,d=newimg("ATL-IDN-008.png",1000,550)
    title(d,1000,"ATL-IDN-008: Identity Domain Dependencies")
    box(d,350,60,300,45,"Identity Domain",fc=NAVY,fs=14)
    d.text((80,140),"Upstream Dependencies: None",fill=GREEN,font=_f(12))
    d.text((80,165),"Identity is the root domain — no upstream architectural dependencies.",fill="#212121",font=_f(10))
    d.text((80,210),"Downstream Dependents:",fill=RED,font=_f(12))
    ds=[("Knowledge",80,260),("Know. Communication",310,260),("Schol. Communication",540,260),
        ("Scientific Visualization",80,360),("Cognitive Motion",310,360),("Platform Engineering",540,360)]
    for nm,x,y in ds:
        box(d,x,y,200,38,nm,fc=STEEL,fs=10)
        arr(d,500,105,x+100,y,c="#78909C")
    box(d,80,440,800,55,"Constraints: Questions precede Projects · Knowledge precedes Publications · Evidence precedes Claims\nVision precedes Strategy · All dependencies acyclic",fc=LGRAY,bc=NAVY,tc="#212121",fs=9)
    return save(img,"ATL-IDN-008.png")

def d_IDN_009():
    img,d=newimg("ATL-IDN-009.png",1200,650)
    title(d,1200,"ATL-IDN-009: Identity Domain-Driven Design")
    d.rounded_rectangle([60,50,1140,600],radius=12,fill="#E8EAF6",outline=NAVY,width=2)
    d.text((430,58),"Identity Aggregate",fill=NAVY,font=_f(14))
    box(d,400,100,300,42,"ResearchIdentity (Aggregate Root)",fc=NAVY,fs=12)
    for n,x,y in[("ResearchVision",70,190),("ResearchAgenda",290,190),("ResearchArea",510,190),("ResearchQuestion",730,190),("Philosophy",950,190)]:
        box(d,x,y,180,38,n,fc=STEEL,fs=10)
    for n,x,y in[("MethodPrinciple",70,310),("ScientificValue",290,310),("ResearchMotivation",510,310),("EvidenceReference",730,310),("QuestionType",950,310)]:
        box(d,x,y,180,38,n,fc=AMBER,fs=9)
    for n,x,y in[("VisionAlignmentService",130,420),("AgendaCoherenceService",480,420),("QuestionLinkingService",830,420)]:
        box(d,x,y,250,35,n,fc=LGRAY,bc=GREEN,tc="#212121",fs=9)
    for n,x,y in[("AgendaCreated",70,510),("AreaAdded",270,510),("PhilosophyRevised",470,510),("MilestoneRecorded",670,510),("VersionPublished",890,510)]:
        box(d,x,y,180,32,n,fc="#FFECB3",bc=ORANGE,tc="#212121",fs=8)
    legend(d,70,570,[("Aggregate Root",NAVY),("Entity",STEEL),("Value Object",AMBER),("Domain Service",GREEN),("Domain Event",ORANGE)],fs=9)
    return save(img,"ATL-IDN-009.png")

def d_IDN_010():
    img,d=newimg("ATL-IDN-010.png",1200,700)
    title(d,1200,"ATL-IDN-010: Identity Business Rules & Constraints")
    rules=[("IC-001","Identity exists independent of presentation",RED),("IC-002","Identity is never project-centric",RED),
           ("IC-003","Identity never depends upon technology",RED),("IC-004","Identity is always evidence-based",RED),
           ("IC-005","Identity preserves historical continuity",RED),("IC-006","Identity remains research-oriented",RED),
           ("SC-001","Research Questions precede Projects",ORANGE),("SC-002","Knowledge precedes Publications",ORANGE),
           ("SC-003","Evidence precedes Claims",ORANGE),("SC-004","Reasoning precedes Conclusions",ORANGE),
           ("SC-005","Vision precedes Strategy",ORANGE),("SC-006","Agenda precedes Implementation",ORANGE)]
    y=55
    for rid,rt,co in rules:
        box(d,50,y,90,32,rid,fc=co,fs=10)
        box(d,160,y,990,32,rt,fc=LGRAY,bc=co,tc="#212121",fs=9)
        y+=45
    d.text((50,y+15),"Structural: Single AR · Single Agenda · Unique Questions · Single Source of Truth",fill=NAVY,font=_f(10))
    d.text((50,y+40),"Evolution: No deletion · Transitions need evidence · Versions inspectable · Coherence strengthens",fill=NAVY,font=_f(10))
    return save(img,"ATL-IDN-010.png")

def d_IDN_011():
    img,d=newimg("ATL-IDN-011.png",1200,550)
    title(d,1200,"ATL-IDN-011: Identity Verification & Conformance Flow")
    lys=[("Philosophy",50),("Ontology",190),("Domain Model",330),("Interfaces",470),("Requirements",610),("Implementation",750)]
    for nm,x in lys:
        box(d,x,80,120,50,nm,fc=STEEL,fs=10)
        if x>50: arr(d,x-10,105,x,105)
    cats=[("Semantic",NAVY),("Structural","#1565C0"),("Behavioral",GREEN),("Representation",AMBER),("Quality","#6A1B9A"),("Governance",RED)]
    for i,(nm,co) in enumerate(cats):
        box(d,50+i*190,200,170,40,nm+" Verification",fc=co,fs=9)
    lvls=[("Level I: Foundational",MGRAY),("Level II: Operational","#1565C0"),("Level III: Excellence",GREEN),("Level IV: Reference",NAVY)]
    for i,(nm,co) in enumerate(lvls):
        box(d,80+i*270,310,240,38,nm,fc=co,fs=10)
    d.text((50,400),"Severity: Minor → Moderate → Major → Critical",fill="#212121",font=_f(11))
    d.text((50,430),"Critical non-conformance prevents release from being designated as conformant.",fill=RED,font=_f(10))
    return save(img,"ATL-IDN-011.png")

def d_IDN_012():
    img,d=newimg("ATL-IDN-012.png",1200,450)
    title(d,1200,"ATL-IDN-012: Identity Interaction Diagram")
    parts=["Researcher","Identity\nDomain","Knowledge\nDomain","Communication\nDomains","Visualization\nDomain"]
    for i,p in enumerate(parts):
        x=60+i*240
        box(d,x,70,160,45,p,fc=NAVY if i==1 else STEEL,fs=10)
        d.line([(x+80,115),(x+80,380)],fill=MGRAY,width=1)
    ints=[(1,0,"Authenticates researcher",160),(0,1,"Defines vision & agenda",195),(1,2,"Provides identity context",230),
          (1,3,"Provides narrative identity",265),(1,4,"Provides visual identity",300),(2,1,"Reports knowledge state",335),(3,1,"Reports comm. status",370)]
    for s,t,lb,y in ints:
        x1=140+s*240; x2=140+t*240
        arr(d,x1,y,x2,y+12,c="#455A64")
        d.text((min(x1,x2)+10,y-15),lb,fill="#212121",font=_f(8))
    return save(img,"ATL-IDN-012.png")

# ── KNOWLEDGE DOMAIN ──

def d_KNO_001():
    img,d=newimg("ATL-KNO-001.png",1200,650)
    title(d,1200,"ATL-KNO-001: Knowledge Domain Overview","Volume II — Knowledge Architecture")
    box(d,400,60,400,50,"Knowledge Domain",fc="#1565C0",fs=16)
    comps=[("Knowledge Claims",100,170),("Concept System",350,170),("Evidence Framework",600,170),("Ontology Engine",850,170),
           ("Reasoning Chain",100,330),("Knowledge Graph",350,330),("Knowledge Lifecycle",600,330),("Knowledge Interfaces",850,330)]
    for nm,x,y in comps:
        box(d,x,y,200,55,nm,fc="#1565C0",fs=11)
        arr(d,600,110,x+100,170,c="#90A4AE")
    box(d,200,450,800,45,"Upstream: Identity Domain  ·  Downstream: Communication, Visualization",fc=LGRAY,bc="#1565C0",tc="#212121",fs=10)
    box(d,80,530,500,70,"Core Principles:\nKnowledge exists before publication\nScientific meaning is canonical",fc="#1565C0",fs=10)
    box(d,620,530,500,70,"Key Invariants:\nConcepts have stable identifiers\nOntology remains consistent",fc=NAVY,fs=10)
    return save(img,"ATL-KNO-001.png")

def d_KNO_002():
    img,d=newimg("ATL-KNO-002.png",1000,550)
    title(d,1000,"ATL-KNO-002: Knowledge Bounded Context")
    d.rounded_rectangle([50,50,950,500],radius=12,fill="#E3F2FD",outline="#1565C0",width=2)
    d.text((370,57),"Knowledge Bounded Context",fill="#1565C0",font=_f(14))
    box(d,350,95,300,42,"Knowledge Base (Aggregate Root)",fc="#1565C0",fs=12)
    ents=[("Concept",80,185),("Claim",280,185),("EvidenceItem",480,185),("ReasoningChain",680,185),
          ("OntologyTerm",80,305),("KnowledgeRelation",280,305),("ConfidenceLevel",480,305),("KnowledgeVersion",680,305)]
    for n,x,y in ents: box(d,x,y,170,38,n,fc=STEEL,fs=10)
    legend(d,80,400,[("Aggregate Root","#1565C0"),("Entity",STEEL),("Value Object",AMBER)])
    return save(img,"ATL-KNO-002.png")

def d_KNO_003():
    img,d=newimg("ATL-KNO-003.png",1200,400)
    title(d,1200,"ATL-KNO-003: Knowledge Lifecycle")
    sts=[("Hypothesis",MGRAY),("Investigation","#1565C0"),("Evidence Gathering",GREEN),("Validated Knowledge",NAVY),("Disseminated Knowledge","#6A1B9A")]
    x=50
    for nm,co in sts:
        box(d,x,120,170,65,nm,fc=co,fs=11)
        if x>50: arr(d,x-15,152,x,152)
        x+=225
    d.text((150,250),"Events: ClaimProposed · EvidenceCollected · ClaimValidated · KnowledgePublished · KnowledgeRevised",fill=MGRAY,font=_f(10))
    return save(img,"ATL-KNO-003.png")

def d_KNO_004():
    img,d=newimg("ATL-KNO-004.png",1200,550)
    title(d,1200,"ATL-KNO-004: Knowledge Entity-Relationship Diagram")
    ents={"KnowledgeBase":(450,65,200,42),"Concept":(100,190,160,38),"Claim":(320,190,160,38),
          "Evidence":(540,190,160,38),"ReasoningChain":(760,190,180,38),
          "OntologyTerm":(100,330,160,38),"Relation":(320,330,160,38),"Confidence":(540,330,160,38),"Version":(760,330,180,38)}
    for nm,(x,y,w,h) in ents.items():
        box(d,x,y,w,h,nm,fc="#1565C0" if nm=="KnowledgeBase" else STEEL,fs=10)
    for s,t in[("KnowledgeBase","Concept"),("KnowledgeBase","Claim"),("KnowledgeBase","Evidence"),("Claim","Evidence"),
               ("Claim","ReasoningChain"),("Concept","OntologyTerm"),("Concept","Relation"),("Claim","Confidence"),("KnowledgeBase","Version")]:
        sx,sy,sw,sh=ents[s]; dx,dy,dw,dh=ents[t]
        arr(d,sx+sw//2,sy+sh,dx+dw//2,dy,c="#78909C")
    return save(img,"ATL-KNO-004.png")

def d_KNO_005():
    img,d=newimg("ATL-KNO-005.png",1200,550)
    title(d,1200,"ATL-KNO-005: Knowledge Domain-Driven Design")
    d.rounded_rectangle([50,50,1150,500],radius=12,fill="#E3F2FD",outline="#1565C0",width=2)
    box(d,420,65,300,42,"KnowledgeBase (Aggregate Root)",fc="#1565C0",fs=12)
    for n,x,y in[("Concept",80,160),("Claim",300,160),("EvidenceItem",520,160),("ReasoningChain",740,160),("OntologyTerm",960,160)]:
        box(d,x,y,180,38,n,fc=STEEL,fs=10)
    for n,x,y in[("ConfidenceLevel",80,280),("SemanticRelation",300,280),("EvidenceType",520,280),("ClaimStatus",740,280)]:
        box(d,x,y,180,38,n,fc=AMBER,fs=9)
    for n,x,y in[("KnowledgeConsistencyService",200,390),("OntologyValidationService",550,390)]:
        box(d,x,y,280,35,n,fc=LGRAY,bc=GREEN,tc="#212121",fs=9)
    legend(d,80,460,[("Aggregate Root","#1565C0"),("Entity",STEEL),("Value Object",AMBER),("Domain Service",GREEN)],fs=9)
    return save(img,"ATL-KNO-005.png")

def d_KNO_006():
    img,d=newimg("ATL-KNO-006.png",1200,450)
    title(d,1200,"ATL-KNO-006: Research Lifecycle")
    sts=[("Research Question",NAVY),("Hypothesis","#1565C0"),("Experiment Design",GREEN),("Data Collection",AMBER),
         ("Analysis","#6A1B9A"),("Knowledge Claim",ORANGE),("Publication","#37474F")]
    x=25
    for nm,co in sts:
        box(d,x,110,148,60,nm,fc=co,fs=10)
        if x>25: arr(d,x-10,140,x,140)
        x+=167
    box(d,100,280,1000,40,"Evidence Chain: Question → Hypothesis → Experiment → Data → Analysis → Claim → Publication",fc=LGRAY,bc=NAVY,tc="#212121",fs=9)
    d.text((350,370),"Continuous Feedback & Revision Cycle",fill=MGRAY,font=_f(11))
    return save(img,"ATL-KNO-006.png")

def d_KNO_007():
    img,d=newimg("ATL-KNO-007.png",1200,450)
    title(d,1200,"ATL-KNO-007: Knowledge Evolution Timeline")
    d.line([(100,230),(1100,230)],fill=NAVY,width=3)
    mls=[("Initial Question",150),("First Evidence",300),("Concept Formation",450),("Theory Development",600),("Validation",750),("Mature Knowledge",900),("Knowledge Transfer",1050)]
    for nm,x in mls:
        d.ellipse([x-8,222,x+8,238],fill=AMBER,outline=NAVY)
        box(d,x-60,120,120,45,nm,fc=STEEL,fs=9)
        arr(d,x,165,x,222,c="#78909C")
    d.text((100,280),"Knowledge Complexity →",fill=MGRAY,font=_f(11))
    pts=[(150,380),(300,360),(450,330),(600,290),(750,260),(900,245),(1050,235)]
    for i in range(len(pts)-1): d.line([pts[i],pts[i+1]],fill=GREEN,width=2)
    return save(img,"ATL-KNO-007.png")

def d_KNO_008():
    img,d=newimg("ATL-KNO-008.png",1200,550)
    title(d,1200,"ATL-KNO-008: Academic Trust Model")
    lvls=[("Recognition & Impact",200,390,800,45,AMBER),("Publication Record",250,325,700,45,ORANGE),
          ("Evidence Base",300,260,600,45,GREEN),("Methodological Rigor",350,195,500,45,"#1565C0"),
          ("Research Consistency",400,130,400,45,STEEL),("Identity Coherence",450,65,300,45,NAVY)]
    for t,x,y,w,h,co in lvls: box(d,x,y,w,h,t,fc=co,fs=12)
    d.text((50,470),"Trust builds from bottom to top — each level depends on the stability of lower levels",fill="#212121",font=_f(11))
    return save(img,"ATL-KNO-008.png")

# ── COMMUNICATION ──

def d_COM_001():
    img,d=newimg("ATL-COM-001.png",1200,550)
    title(d,1200,"ATL-COM-001: Knowledge Communication Domain","Volume III")
    box(d,400,55,400,45,"Knowledge Communication",fc="#00838F",fs=15)
    for nm,x,y in[("Narrative Engine",100,160),("Story Framework",350,160),("Research Narrative",600,160),("Audience Model",850,160),
                   ("Communication Strategy",100,310),("Knowledge Translation",350,310),("Engagement Design",600,310),("Impact Tracking",850,310)]:
        box(d,x,y,200,55,nm,fc="#00838F",fs=11)
    box(d,200,430,800,40,"Upstream: Identity + Knowledge  ·  Downstream: Scholarly Communication",fc=LGRAY,bc="#00838F",tc="#212121",fs=10)
    return save(img,"ATL-COM-001.png")

def d_COM_002():
    img,d=newimg("ATL-COM-002.png",1200,550)
    title(d,1200,"ATL-COM-002: Scholarly Communication Domain","Volume IV")
    box(d,400,55,400,45,"Scholarly Communication",fc="#00897B",fs=15)
    for nm,x,y in[("Publication Lifecycle",100,160),("Peer Review",350,160),("Citation Management",600,160),("Impact Measurement",850,160),
                   ("Academic Profile",100,310),("Venue Strategy",350,310),("Open Access",600,310),("Scholarly Reputation",850,310)]:
        box(d,x,y,200,55,nm,fc="#00897B",fs=11)
    box(d,200,430,800,40,"Upstream: Knowledge + Communication  ·  Downstream: Visualization, Motion",fc=LGRAY,bc="#00897B",tc="#212121",fs=10)
    return save(img,"ATL-COM-002.png")

# ── VISUALIZATION & MOTION ──

def d_VIS_001():
    img,d=newimg("ATL-VIS-001.png",1200,550)
    title(d,1200,"ATL-VIS-001: Scientific Visualization Domain","Volume V")
    box(d,400,55,400,45,"Scientific Visualization",fc=GREEN,fs=15)
    for nm,x,y in[("Knowledge Graphs",100,160),("Concept Maps",350,160),("Ontology Diagrams",600,160),("Research Dashboards",850,160),
                   ("Timeline Visualizations",100,310),("Citation Networks",350,310),("Impact Visualizations",600,310),("Architecture Diagrams",850,310)]:
        box(d,x,y,200,55,nm,fc=GREEN,fs=11)
    box(d,200,430,800,40,"Visualizations explain architecture — they NEVER redefine architecture",fc=LGRAY,bc=GREEN,tc="#212121",fs=10)
    return save(img,"ATL-VIS-001.png")

def d_MOT_001():
    img,d=newimg("ATL-MOT-001.png",1200,550)
    title(d,1200,"ATL-MOT-001: Cognitive Motion Domain","Volume VI")
    box(d,400,55,400,45,"Cognitive Motion",fc="#6A1B9A",fs=15)
    for nm,x,y in[("Research Direction",100,160),("Intellectual Momentum",350,160),("Trajectory Modeling",600,160),("Evolution Tracking",850,160),
                   ("Milestone Management",100,310),("Transition Reasoning",350,310),("Continuity Preservation",600,310),("Future Projection",850,310)]:
        box(d,x,y,200,55,nm,fc="#6A1B9A",fs=11)
    box(d,200,430,800,40,"Motion preserves identity through evolution — it is NOT change for change's sake",fc=LGRAY,bc="#6A1B9A",tc="#212121",fs=10)
    return save(img,"ATL-MOT-001.png")

# ── PLATFORM & IMPLEMENTATION ──

def d_ENG_001():
    img,d=newimg("ATL-ENG-001.png",1200,550)
    title(d,1200,"ATL-ENG-001: Platform Engineering Domain","Volume VII")
    box(d,400,55,400,45,"Platform Engineering",fc=ORANGE,fs=15)
    for nm,x,y in[("Service Architecture",100,160),("API Gateway",350,160),("Database Layer",600,160),("Search Engine",850,160),
                   ("Auth System",100,310),("Caching Layer",350,310),("Event Bus",600,310),("Observability Stack",850,310)]:
        box(d,x,y,200,55,nm,fc=ORANGE,fs=11)
    box(d,200,430,800,40,"Technology serves architecture · Architecture serves research",fc=LGRAY,bc=ORANGE,tc="#212121",fs=10)
    return save(img,"ATL-ENG-001.png")

def d_ENG_002():
    img,d=newimg("ATL-ENG-002.png",1200,650)
    title(d,1200,"ATL-ENG-002: Platform Service Architecture")
    lys=[("API Gateway",["Rate Limiting","Auth","Routing","Load Balancing"],60),
         ("Application Services",["Identity Svc","Knowledge Svc","Publication Svc","Visualization Svc"],170),
         ("Domain Services",["Vision Engine","Agenda Engine","Ontology Engine","Narrative Engine"],280),
         ("Data Layer",["PostgreSQL","Elasticsearch","Redis Cache","Object Storage"],390),
         ("Infrastructure",["Docker","Kubernetes","Message Queue","Monitoring"],500)]
    for ln,svcs,y in lys:
        box(d,50,y,200,45,ln,fc=NAVY,fs=10)
        for i,sv in enumerate(svcs): box(d,280+i*220,y,200,45,sv,fc=STEEL,fs=9)
    return save(img,"ATL-ENG-002.png")

def d_ENG_003():
    img,d=newimg("ATL-ENG-003.png",1200,550)
    title(d,1200,"ATL-ENG-003: Infrastructure Architecture")
    box(d,50,60,300,45,"Frontend (Next.js / Static)",fc=AMBER,fs=11)
    box(d,400,60,300,45,"API Gateway",fc=NAVY,fs=11)
    for nm,x,y in[("Identity Service",50,180),("Knowledge Service",280,180),("Publication Service",510,180),("Visualization Service",740,180)]:
        box(d,x,y,200,50,nm,fc=STEEL,fs=10)
    for nm,x,y in[("PostgreSQL",50,310),("Elasticsearch",280,310),("Redis",510,310),("S3 Storage",740,310)]:
        box(d,x,y,200,50,nm,fc=GREEN,fs=10)
    box(d,50,420,890,45,"Infrastructure: Docker · Kubernetes · Message Queue · CI/CD · Monitoring · Logging",fc="#37474F",fs=10)
    for i in range(4): arr(d,550,105,150+i*230,180)
    for i in range(4): arr(d,150+i*230,230,150+i*230,310)
    return save(img,"ATL-ENG-003.png")

def d_ENG_004():
    img,d=newimg("ATL-ENG-004.png",1200,500)
    title(d,1200,"ATL-ENG-004: Database Architecture")
    dbs=[("PostgreSQL","Tables: researchers, visions,\nagendas, areas, questions, evidence",50,80),
         ("Elasticsearch","Indices: concepts, claims,\npublications, full-text search",450,80),
         ("Redis Cache","Cache: sessions, computed views,\nrate limiting, hot queries",850,80),
         ("Object Storage","Blobs: diagrams, documents,\nPDFs, datasets, visualizations",50,300),
         ("Graph Database","Nodes: concepts, claims, evidence\nEdges: semantic relationships",450,300),
         ("Time Series DB","Metrics: API latency, error rates,\nuser engagement, system health",850,300)]
    for nm,desc,x,y in dbs:
        box(d,x,y,350,120,"",fc=LGRAY,bc=STEEL,tc="#212121")
        d.text((x+15,y+10),nm,fill=NAVY,font=_f(13))
        d.text((x+15,y+50),desc,fill="#212121",font=_f(10))
    return save(img,"ATL-ENG-004.png")

# ── IMPLEMENTATION ──

def d_IMP_001():
    img,d=newimg("ATL-IMP-001.png",1200,550)
    title(d,1200,"ATL-IMP-001: Implementation Domain","Volume VIII")
    box(d,400,55,400,45,"Implementation",fc="#37474F",fs=15)
    for nm,x,y in[("Repository Structure",100,160),("Development Workflow",350,160),("Claude Code Integration",600,160),("Testing Framework",850,160),
                   ("CI/CD Pipeline",100,310),("Deployment Strategy",350,310),("Monitoring Pipeline",600,310),("Production Architecture",850,310)]:
        box(d,x,y,200,55,nm,fc="#37474F",fs=11)
    box(d,200,430,800,40,"AI-first development · Architecture-driven implementation · Planning before coding",fc=LGRAY,bc="#37474F",tc="#212121",fs=10)
    return save(img,"ATL-IMP-001.png")

def d_IMP_002():
    img,d=newimg("ATL-IMP-002.png",1000,650)
    title(d,1000,"ATL-IMP-002: Repository Layout")
    folders=[("rios/",60,65,NAVY),("├── src/",120,110,STEEL),
             ("│   ├── identity/",180,155,"#1565C0"),("│   ├── knowledge/",180,195,"#1565C0"),
             ("│   ├── communication/",180,235,"#00838F"),("│   ├── visualization/",180,275,GREEN),
             ("│   ├── motion/",180,315,"#6A1B9A"),("│   └── platform/",180,355,ORANGE),
             ("├── tests/",120,400,GREEN),("├── docs/",120,440,AMBER),
             ("├── config/",120,480,MGRAY),("├── migrations/",120,520,"#37474F"),
             ("└── infrastructure/",120,560,"#37474F")]
    for nm,x,y,co in folders:
        d.text((x,y),nm,fill=co,font=_f(13))
    return save(img,"ATL-IMP-002.png")

def d_IMP_003():
    img,d=newimg("ATL-IMP-003.png",1200,450)
    title(d,1200,"ATL-IMP-003: Development Workflow")
    sts=[("Architecture Review",NAVY),("Planning & Design","#1565C0"),("Claude Code Implementation",GREEN),
         ("Testing & Validation",AMBER),("Conformance Check","#6A1B9A"),("Deployment",ORANGE)]
    x=40
    for nm,co in sts:
        box(d,x,130,160,65,nm,fc=co,fs=10)
        if x>40: arr(d,x-10,162,x,162)
        x+=190
    d.text((400,300),"Continuous Feedback & Iteration",fill=MGRAY,font=_f(11))
    return save(img,"ATL-IMP-003.png")

def d_IMP_004():
    img,d=newimg("ATL-IMP-004.png",1200,500)
    title(d,1200,"ATL-IMP-004: Deployment Architecture")
    box(d,50,70,250,110,"",fc=LGRAY,bc=STEEL,tc="#212121")
    d.text((70,80),"Development",fill=NAVY,font=_f(14))
    d.text((70,108),"Local Docker Compose\nHot Reload\nDebug Mode",fill="#212121",font=_f(10))
    box(d,400,70,250,110,"",fc=LGRAY,bc=AMBER,tc="#212121")
    d.text((420,80),"Staging",fill=NAVY,font=_f(14))
    d.text((420,108),"Kubernetes Cluster\nIntegration Tests\nPerformance Tests",fill="#212121",font=_f(10))
    box(d,750,70,350,110,"",fc=LGRAY,bc=GREEN,tc="#212121")
    d.text((770,80),"Production",fill=NAVY,font=_f(14))
    d.text((770,108),"Kubernetes (Multi-Region)\nCDN · Load Balancer\nAuto-Scaling · Monitoring",fill="#212121",font=_f(10))
    arr(d,300,125,400,125); arr(d,650,125,750,125)
    box(d,150,270,900,40,"CI/CD: Build → Test → Lint → Security Scan → Staging → Integration → Production",fc=NAVY,fs=9)
    box(d,150,340,900,40,"Observability: Prometheus · Grafana · ELK · Distributed Tracing · Alert Manager",fc="#37474F",fs=10)
    return save(img,"ATL-IMP-004.png")

# ── SYSTEM RELATIONSHIPS ──

def d_REF_001():
    img,d=newimg("ATL-REF-001.png",1200,700)
    title(d,1200,"ATL-REF-001: Cross-Domain Dependency Map")
    pos={"Identity":(500,60),"Knowledge":(200,200),"Know. Comm.":(500,200),"Schol. Comm.":(800,200),
         "Visualization":(200,380),"Motion":(500,380),"Platform Eng.":(800,380),"Implementation":(500,540)}
    cm={"Identity":"Identity","Knowledge":"Knowledge","Know. Comm.":"KnowComm","Schol. Comm.":"ScholComm",
        "Visualization":"Vis","Motion":"Motion","Platform Eng.":"Platform","Implementation":"Impl"}
    for nm,(x,y) in pos.items():
        box(d,x,y,180,45,nm,fc=DC.get(cm[nm],STEEL),fs=11)
    deps=[("Identity","Knowledge"),("Identity","Know. Comm."),("Knowledge","Schol. Comm."),
          ("Knowledge","Visualization"),("Know. Comm.","Schol. Comm."),("Identity","Motion"),("Platform Eng.","Implementation")]
    for s,t in deps:
        sx,sy=pos[s]; dx,dy=pos[t]
        arr(d,sx+90,sy+45,dx+90,dy,c="#78909C")
    d.text((50,630),"Arrows indicate architectural dependency direction (upstream → downstream)",fill="#212121",font=_f(11))
    return save(img,"ATL-REF-001.png")

def d_REF_002():
    img,d=newimg("ATL-REF-002.png",1200,700)
    title(d,1200,"ATL-REF-002: Strategic Context Map","DDD Context Map Pattern")
    ctxs=[("Identity Context",450,55,250,55,NAVY,"Upstream (Core)"),
          ("Knowledge Context",80,200,220,55,"#1565C0","Core"),
          ("Communication Context",450,200,250,55,"#00838F","Downstream"),
          ("Visualization Context",800,200,220,55,GREEN,"Supporting"),
          ("Motion Context",80,380,220,55,"#6A1B9A","Supporting"),
          ("Platform Context",450,380,250,55,ORANGE,"Supporting"),
          ("Implementation Context",800,380,220,55,"#37474F","Generic")]
    for nm,x,y,w,h,co,rl in ctxs:
        box(d,x,y,w,h,nm,fc=co,fs=12)
        d.text((x+w//2-30,y+h+5),rl,fill=MGRAY,font=_f(8))
    pos={c[0]:(c[1],c[2]) for c in ctxs}
    for s,t in[("Identity Context","Knowledge Context"),("Identity Context","Communication Context"),
               ("Knowledge Context","Communication Context"),("Knowledge Context","Visualization Context"),
               ("Identity Context","Motion Context"),("Platform Context","Implementation Context")]:
        sx,sy=pos[s]; dx,dy=pos[t]
        arr(d,sx+100,sy+55,dx+100,dy,c="#78909C")
    return save(img,"ATL-REF-002.png")

def d_REF_003():
    img,d=newimg("ATL-REF-003.png",1200,550)
    title(d,1200,"ATL-REF-003: Consumer-Provider Relationship Graph")
    provs=[("Identity",80,80),("Knowledge",80,220),("Platform Eng.",80,360)]
    cons=[("Knowledge",380,60),("Know. Comm.",380,160),("Schol. Comm.",380,260),
          ("Visualization",680,60),("Motion",680,160),("Implementation",680,260)]
    for nm,x,y in provs: box(d,x,y,180,40,f"Provider: {nm}",fc=GREEN,fs=9)
    for nm,x,y in cons: box(d,x,y,180,40,f"Consumer: {nm}",fc=AMBER,fs=9)
    for dy in[82,182,282]: arr(d,260,100,380,dy,c="#78909C")
    arr(d,260,240,680,82,c="#78909C"); arr(d,260,240,680,182,c="#78909C")
    arr(d,260,380,680,282,c="#78909C")
    return save(img,"ATL-REF-003.png")

# ═══════════════════════════════════════════════════════════════════
# DOCUMENT ASSEMBLY
# ═══════════════════════════════════════════════════════════════════

def build_doc():
    doc=Document()
    st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(11); st.font.color.rgb=RGBColor(0x21,0x21,0x21)
    for i in range(1,5):
        hs=doc.styles[f'Heading {i}']; hs.font.color.rgb=RGBColor(0x1B,0x2A,0x4A); hs.font.name='Calibri'
    doc.styles['Heading 1'].font.size=Pt(24); doc.styles['Heading 2'].font.size=Pt(18)
    doc.styles['Heading 3'].font.size=Pt(14); doc.styles['Heading 4'].font.size=Pt(12)

    def cover():
        for _ in range(6): doc.add_paragraph("")
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run("RIOS ATLAS"); r.font.size=Pt(42); r.font.color.rgb=RGBColor(0x1B,0x2A,0x4A); r.bold=True
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run("Visual Architecture Reference"); r.font.size=Pt(22); r.font.color.rgb=RGBColor(0x2D,0x5F,0x8A)
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run("Research Identity Operating System"); r.font.size=Pt(16); r.font.color.rgb=RGBColor(0x90,0xA4,0xAE)
        for _ in range(4): doc.add_paragraph("")
        t=doc.add_table(rows=8,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER
        for i,(k,v) in enumerate([("Document","RIOS Atlas — Visual Architecture Reference"),("Version","1.0.0"),
            ("Date","July 2026"),("Classification","Official Architecture Document"),("Status","Publication Release"),
            ("Standard","IEEE 1471 / ISO 42010"),("Domains","8 Architectural Domains"),("Diagrams","60+ Architecture Diagrams")]):
            t.rows[i].cells[0].text=k; t.rows[i].cells[1].text=v
        doc.add_page_break()

    def revision():
        doc.add_heading("Revision History",level=2)
        t=doc.add_table(rows=2,cols=4); t.style='Table Grid'
        for i,h in enumerate(["Version","Date","Author","Description"]):
            t.rows[0].cells[i].text=h
            for p in t.rows[0].cells[i].paragraphs:
                for r in p.runs: r.bold=True
        for i,v in enumerate(["1.0.0","July 2026","RIOS Architecture Team","Initial release — Complete Atlas"]):
            t.rows[1].cells[i].text=v
        doc.add_page_break()

    def toc():
        doc.add_heading("Table of Contents",level=2)
        items=[("PART I: Architecture Foundation",1),("  System Overview, Vision, Principles, Layers, Domains",2),
               ("PART II: Domain Atlas",1),("  Identity · Knowledge · Communication · Visualization · Motion · Platform · Implementation",2),
               ("PART III: System Relationships",1),("  Cross-Domain Dependencies, Context Maps, Consumer-Provider Graphs",2),
               ("PART IV: Knowledge & Research Maps",1),("  Research Lifecycle, Knowledge Evolution, Trust Model",2),
               ("PART V: Engineering Atlas",1),("  Platform, Infrastructure, Database Architecture",2),
               ("PART VI: Implementation Atlas",1),("  Repository Layout, Workflow, Deployment",2),
               ("PART VII: Reference Atlas",1),("  Glossary, Notation Guide, Diagram Legend, Index",2)]
        for it,lv in items:
            p=doc.add_paragraph(); r=p.add_run(it)
            r.font.size=Pt(11 if lv==2 else 13); r.bold=(lv==1)
            r.font.color.rgb=RGBColor(0x1B,0x2A,0x4A) if lv==1 else RGBColor(0x42,0x42,0x42)
        doc.add_page_break()

    def add_d(did,title,purpose,vols,layer,desc,insights,img_path,xrefs=None):
        doc.add_heading(f"{did}: {title}",level=3)
        t=doc.add_table(rows=5,cols=2); t.style='Table Grid'
        for i,(k,v) in enumerate([("Diagram ID",did),("Purpose",purpose),("Referenced Volume(s)",vols),
                                   ("Architecture Layer",layer),("Cross References",xrefs or "—")]):
            t.rows[i].cells[0].text=k; t.rows[i].cells[1].text=v
            for p in t.rows[i].cells[0].paragraphs:
                for r in p.runs: r.bold=True; r.font.size=Pt(9)
            for p in t.rows[i].cells[1].paragraphs:
                for r in p.runs: r.font.size=Pt(9)
        doc.add_paragraph("")
        if img_path and os.path.exists(img_path):
            doc.add_picture(img_path,width=Inches(6.0))
            doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")
        doc.add_heading("Description",level=4); doc.add_paragraph(desc)
        doc.add_heading("Key Architectural Insights",level=4)
        for ins in insights: doc.add_paragraph(ins,style='List Bullet')
        doc.add_page_break()

    # ── BUILD ──
    cover(); revision(); toc()

    # PART I
    doc.add_heading("PART I: Architecture Foundation",level=1)
    doc.add_paragraph("This part establishes the complete foundational view of RIOS — its vision, principles, layers, domains, dependencies, and landscape.")
    doc.add_page_break()

    add_d("ATL-FOUND-001","System Overview","Top-level view of the entire RIOS platform","All Volumes","System",
        "RIOS consists of eight architectural domains organized into strategic, knowledge, communication, presentation, evolution, platform, and implementation layers. All domains rest upon a shared foundation architecture.",
        ["RIOS is a composition of eight specialized architectural domains","Foundation Architecture is shared across all domains","Each domain maps to a specific architectural volume"],
        d_FOUND_001(),"ATL-FOUND-004, ATL-FOUND-005")

    add_d("ATL-FOUND-002","Architecture Vision","Guiding philosophy and vision","Volume 0","Strategic",
        "RIOS is built upon Research Identity as a Living Scientific System with six pillars: Research Before Software, Knowledge Before Publication, Meaning Before Implementation, Evidence Before Claims, Architecture Before Technology, Technology Serves Architecture.",
        ["Every architectural decision must trace back to research value","Technology is a means, never an end"],
        d_FOUND_002(),"ATL-FOUND-003")

    add_d("ATL-FOUND-003","Architecture Principles","Ten normative architectural principles","Volume 0","Governance",
        "Ten principles (AP-001 through AP-010) define the inviolable design philosophy: Research Primacy, Semantic Integrity, Evidence Authority, Knowledge Independence, Identity Persistence, Architectural Purity, Domain Autonomy, Open Knowledge, Trust Through Evidence, Cognitive Accessibility.",
        ["Principles are ordered by importance — Research Primacy is supreme","Domain Autonomy prevents cross-contamination"],
        d_FOUND_003(),"ATL-FOUND-002, ATL-IDN-010")

    add_d("ATL-FOUND-004","Architecture Layers","Layered architecture from strategic to foundation","Volume 0","System",
        "Eight layers: Strategic (Identity), Knowledge, Communication, Presentation (Visualization), Evolution (Motion), Platform, Implementation, and Foundation. Each layer depends only on layers above it.",
        ["The Strategic Layer has no upstream dependencies — it is the root","Layer separation prevents architectural drift"],
        d_FOUND_004(),"ATL-FOUND-006, ATL-FOUND-007")

    add_d("ATL-FOUND-005","Domain Overview","All eight RIOS domains with their core questions","All Volumes","System",
        "Each domain addresses a fundamental question: Identity (Who?), Knowledge (What?), Communication (How shared?), Visualization (How shown?), Motion (How evolved?), Platform (How built?), Implementation (How realized?).",
        ["Every domain has a clear, distinct responsibility","Domains are organized by architectural concern, not technology"],
        d_FOUND_005(),"ATL-FOUND-001, ATL-FOUND-008")

    add_d("ATL-FOUND-006","Architecture Stack","Technology-independent architecture stack","Volume 0","System",
        "The stack progresses from Strategic through Knowledge, Communication, Presentation, Evolution, Platform, Realization, to shared Governance foundation.",
        ["The stack is technology-independent","Each layer encapsulates a specific category of concern"],
        d_FOUND_006(),"ATL-FOUND-004, ATL-FOUND-007")

    add_d("ATL-FOUND-007","Dependency Architecture","Complete inter-domain dependency graph","All Volumes","System",
        "Identity sits at the top with no upstream dependencies. All other domains depend on Identity or Knowledge. Dependencies are strictly acyclic. Foundation Architecture is a shared dependency for all.",
        ["All dependencies flow in one direction","Circular dependencies are prohibited","Identity is the architectural root"],
        d_FOUND_007(),"ATL-FOUND-004, ATL-REF-001")

    add_d("ATL-FOUND-008","Capability Map","Complete capability set for each domain","All Volumes","System",
        "Each domain provides six primary capabilities. Foundation capabilities (governance, terminology, domain model, dependency matrix, ownership, requirements) are cross-cutting.",
        ["Capabilities are non-overlapping between domains","Each capability maps to specific functional requirements"],
        d_FOUND_008(),"ATL-FOUND-005, ATL-IDN-007")

    add_d("ATL-FOUND-009","System Context Diagram","C4 Level 0 — RIOS and external actors","All Volumes","System",
        "RIOS interacts with Researchers, Admissions Committees, Faculty Reviewers, and AI Coding Agents. External integrations include Academic Databases, Publication Repositories, and Citation Networks.",
        ["AI Coding Agents are first-class actors in the RIOS ecosystem","All external interactions pass through defined interfaces"],
        d_FOUND_009(),"ATL-FOUND-001, ATL-ENG-002")

    add_d("ATL-FOUND-010","Architecture Hierarchy","Hierarchical decomposition from vision to domains","Volume 0","System",
        "Architecture Vision → Architecture Principles → Domain Architecture + Foundation Architecture → Individual Domains.",
        ["Architecture Vision is the single root of the entire hierarchy"],
        d_FOUND_010(),"ATL-FOUND-003, ATL-FOUND-004")

    add_d("ATL-FOUND-011","Architecture Flow","Primary flow from research direction through implementation","All Volumes","System",
        "Research Direction → Knowledge Construction → Communication → Visualization → Evolution → Platform → Implementation. Foundation underpins all. Verification provides continuous feedback.",
        ["The flow is unidirectional","No architectural shortcut shall bypass this flow"],
        d_FOUND_011(),"ATL-FOUND-007, ATL-IDN-011")

    # PART II
    doc.add_heading("PART II: Domain Atlas",level=1)
    doc.add_paragraph("Complete visual chapter for each of the eight architectural domains.")
    doc.add_page_break()

    doc.add_heading("Chapter 1: Identity Architecture",level=2)
    doc.add_paragraph("The Identity Domain is the strategic root of RIOS. Volume I provides the complete specification.")
    doc.add_page_break()

    add_d("ATL-IDN-001","Identity Domain Overview","Nine architectural components with stability characteristics","Volume I","Strategic",
        "Nine components from Vision Engine (10-20yr, extremely stable) through Evolution Layer (very stable). Meaning flows downward; interpretation flows upward. Identity owns vision, agenda, philosophy, values, direction, evolution. It does NOT own publications, datasets, visualizations, motion, or implementation.",
        ["Higher-stability components SHALL NOT depend upon lower-stability components","The dependency chain SHALL remain acyclic"],
        d_IDN_001(),"ATL-IDN-005, ATL-IDN-008")

    add_d("ATL-IDN-002","Identity Bounded Context","DDD bounded context with aggregate root, entities, value objects","Volume I","Domain Model",
        "Single aggregate root (Research Identity) with eight entities (Vision, Agenda, Area, Question, Philosophy, Evidence, Version, Milestone) and four value objects (Method Principle, Scientific Value, Research Motivation, Evidence Reference).",
        ["Single Aggregate Root (ST-001)","Every entity possesses at least one semantic relationship (RC-001)"],
        d_IDN_002(),"ATL-IDN-004, ATL-IDN-009")

    add_d("ATL-IDN-003","Identity Lifecycle","Lifecycle states from Nascent to Legacy","Volume I","Domain Model",
        "Five states: Nascent → Emerging → Established → Mature → Legacy. Domain events mark transitions. Evolution SHALL strengthen coherence.",
        ["No historical deletion (EC-001)","Every transition requires evidence (EC-002)"],
        d_IDN_003(),"ATL-IDN-001, ATL-MOT-001")

    add_d("ATL-IDN-004","Identity ER Diagram","Complete entity-relationship diagram","Volume I","Domain Model",
        "Research Identity central entity connected to Vision, Agenda, Area, Question, Philosophy, Evidence, Version, Milestone. Secondary entities: Method Principle, Scientific Value, Research Motivation, Representation.",
        ["Every relationship is semantically typed","Circular semantic dependencies prohibited (RC-003)"],
        d_IDN_004(),"ATL-IDN-002, ATL-IDN-009")

    add_d("ATL-IDN-005","Identity Architecture Blueprint","9-layer canonical structure with stability annotations","Volume I","Architecture Blueprint",
        "Nine layers from Research Identity through External Communication. Each has stability rating. This blueprint is normative — all implementations SHALL conform.",
        ["Vision exists before Agenda; Questions before Publications","Evidence validates but does not define higher layers"],
        d_IDN_005(),"ATL-IDN-001, ATL-IDN-008")

    add_d("ATL-IDN-006","Identity Context Map","DDD context map with all bounded contexts","Volume I","Domain Model",
        "Identity provides context to Knowledge, Communication, Visualization, and Motion domains. All downstream consumers depend on Identity as upstream provider.",
        ["Downstream consumers SHALL NOT modify canonical identity (IDN-INT-003)"],
        d_IDN_006(),"ATL-REF-002, ATL-IDN-008")

    add_d("ATL-IDN-007","Identity Capability Map","Detailed capability breakdown","Volume I","Capabilities",
        "Eight capability groups: Vision, Agenda, Area, Question, Philosophy, Evidence, Representation, and Evolution Management.",
        ["Each capability maps to specific functional requirements","Evidence Management bridges Identity and Knowledge"],
        d_IDN_007(),"ATL-IDN-001, ATL-FOUND-008")

    add_d("ATL-IDN-008","Identity Domain Dependencies","Upstream and downstream dependency mapping","Volume I","Dependencies",
        "Zero upstream dependencies. Six downstream dependents. Semantic constraints enforce ordering: Questions precede Projects, Knowledge precedes Publications, etc.",
        ["Identity is the architectural root","All dependencies are acyclic"],
        d_IDN_008(),"ATL-FOUND-007, ATL-REF-001")

    add_d("ATL-IDN-009","Identity DDD Diagram","Complete tactical DDD diagram","Volume I","Domain Model",
        "Aggregate root, 5 entities, 5 value objects, 3 domain services, 5 domain events. Services handle cross-entity operations.",
        ["Domain events enable event-driven integration","Value objects enforce immutability"],
        d_IDN_009(),"ATL-IDN-002, ATL-IDN-004")

    add_d("ATL-IDN-010","Identity Business Rules","All invariants and constraints","Volume I","Constraints",
        "6 foundational invariants (IC), 6 semantic constraints (SC), 4 structural constraints (ST). All are immutable rules — violation is architectural non-conformance.",
        ["IC-001 is most fundamental: presentation communicates identity but never creates it","IC-003 ensures identity survives technological evolution"],
        d_IDN_010(),"ATL-IDN-008, ATL-IDN-011")

    add_d("ATL-IDN-011","Identity Verification Flow","Verification process across 6 layers and 6 categories","Volume I","Verification",
        "Verification: Philosophy → Ontology → Domain Model → Interfaces → Requirements → Implementation. Categories: Semantic, Structural, Behavioral, Representation, Quality, Governance. 4 conformance levels.",
        ["Critical non-conformance prevents conformant designation","Verification should be continuous"],
        d_IDN_011(),"ATL-IDN-010, ATL-FOUND-011")

    add_d("ATL-IDN-012","Identity Interaction Diagram","Sequence diagram of typical interactions","Volume I","Interactions",
        "Researcher authenticates → defines vision → Identity provides context to Knowledge, Communication, Visualization domains. All interactions through standardized interfaces.",
        ["All interactions pass through defined interfaces","Interface consumers SHALL NOT modify identity"],
        d_IDN_012(),"ATL-IDN-006, ATL-REF-003")

    # Knowledge Domain
    doc.add_heading("Chapter 2: Knowledge Architecture",level=2)
    doc.add_paragraph("The Knowledge Domain manages what the researcher knows. Volume II provides the complete specification.")
    doc.add_page_break()

    add_d("ATL-KNO-001","Knowledge Domain Overview","Eight architectural components","Volume II","Knowledge",
        "Components: Knowledge Claims, Concept System, Evidence Framework, Ontology Engine, Reasoning Chain, Knowledge Graph, Knowledge Lifecycle, Knowledge Interfaces.",
        ["Knowledge exists independently of publication","Scientific meaning is canonical"],
        d_KNO_001(),"ATL-KNO-002, ATL-KNO-006")

    add_d("ATL-KNO-002","Knowledge Bounded Context","DDD bounded context","Volume II","Domain Model",
        "Knowledge Base as aggregate root with 8 entities: Concept, Claim, EvidenceItem, ReasoningChain, OntologyTerm, KnowledgeRelation, ConfidenceLevel, KnowledgeVersion.",
        ["Claims are primary knowledge assertions — each requires evidence","Ontology terms provide formal semantic grounding"],
        d_KNO_002(),"ATL-KNO-004, ATL-KNO-005")

    add_d("ATL-KNO-003","Knowledge Lifecycle","Hypothesis to dissemination","Volume II","Domain Model",
        "Five states: Hypothesis → Investigation → Evidence Gathering → Validated Knowledge → Disseminated Knowledge.",
        ["Evidence precedes claims","The lifecycle is circular"],
        d_KNO_003(),"ATL-KNO-006, ATL-IDN-003")

    add_d("ATL-KNO-004","Knowledge ER Diagram","Complete entity-relationship diagram","Volume II","Domain Model",
        "Knowledge Base connected to Concept, Claim, Evidence, ReasoningChain, OntologyTerm, Relation, Confidence, Version. Claims connect to Evidence and ReasoningChain.",
        ["Every claim MUST reference supporting evidence","Confidence is a value object"],
        d_KNO_004(),"ATL-KNO-002, ATL-IDN-004")

    add_d("ATL-KNO-005","Knowledge DDD Diagram","Complete tactical DDD diagram","Volume II","Domain Model",
        "Aggregate root, 5 entities, 4 value objects, 2 domain services (KnowledgeConsistencyService, OntologyValidationService).",
        ["Services ensure cross-entity consistency and ontology compliance"],
        d_KNO_005(),"ATL-KNO-002, ATL-KNO-004")

    # Communication
    doc.add_heading("Chapter 3: Knowledge Communication",level=2)
    doc.add_paragraph("How knowledge is communicated through narrative. Volume III.")
    doc.add_page_break()

    add_d("ATL-COM-001","Knowledge Communication Domain","Eight components","Volume III","Communication",
        "Components: Narrative Engine, Story Framework, Research Narrative, Audience Model, Communication Strategy, Knowledge Translation, Engagement Design, Impact Tracking.",
        ["Translates expert knowledge for broader audiences"],
        d_COM_001(),"ATL-COM-002, ATL-KNO-001")

    doc.add_heading("Chapter 4: Scholarly Communication",level=2)
    doc.add_paragraph("Formal academic communication. Volume IV.")
    doc.add_page_break()

    add_d("ATL-COM-002","Scholarly Communication Domain","Eight components","Volume IV","Communication",
        "Components: Publication Lifecycle, Peer Review Framework, Citation Management, Impact Measurement, Academic Profile, Venue Strategy, Open Access Management, Scholarly Reputation.",
        ["Publications are evidence, not identity — they belong to this domain"],
        d_COM_002(),"ATL-COM-001, ATL-VIS-001")

    # Visualization & Motion
    doc.add_heading("Chapter 5: Scientific Visualization",level=2)
    doc.add_paragraph("Visual representations of research. Volume V.")
    doc.add_page_break()

    add_d("ATL-VIS-001","Scientific Visualization Domain","Eight components","Volume V","Visualization",
        "Components: Knowledge Graphs, Concept Maps, Ontology Diagrams, Research Dashboards, Timeline Visualizations, Citation Networks, Impact Visualizations, Architecture Diagrams.",
        ["Visualizations explain architecture — they NEVER redefine it","Every diagram preserves semantic integrity"],
        d_VIS_001(),"ATL-KNO-001, ATL-MOT-001")

    doc.add_heading("Chapter 6: Cognitive Motion",level=2)
    doc.add_paragraph("How research identity evolves. Volume VI.")
    doc.add_page_break()

    add_d("ATL-MOT-001","Cognitive Motion Domain","Eight components","Volume VI","Evolution",
        "Components: Research Direction, Intellectual Momentum, Trajectory Modeling, Evolution Tracking, Milestone Management, Transition Reasoning, Continuity Preservation, Future Projection.",
        ["Motion preserves identity through evolution","Evolution SHALL strengthen coherence"],
        d_MOT_001(),"ATL-IDN-003, ATL-VIS-001")

    # Platform & Implementation
    doc.add_heading("Chapter 7: Platform Engineering",level=2)
    doc.add_paragraph("Technical infrastructure. Volume VII.")
    doc.add_page_break()

    add_d("ATL-ENG-001","Platform Engineering Domain","Eight components","Volume VII","Platform",
        "Components: Service Architecture, API Gateway, Database Layer, Search Engine, Auth System, Caching Layer, Event Bus, Observability Stack.",
        ["Technology serves architecture — architecture serves research"],
        d_ENG_001(),"ATL-ENG-002, ATL-ENG-003")

    add_d("ATL-ENG-002","Platform Service Architecture","Five-layer service architecture","Volume VII","Platform",
        "Layers: API Gateway (rate limiting, auth, routing, load balancing), Application Services (identity, knowledge, publication, visualization), Domain Services (vision, agenda, ontology, narrative engines), Data Layer (PostgreSQL, Elasticsearch, Redis, S3), Infrastructure (Docker, K8s, message queue, monitoring).",
        ["Each service encapsulates a single domain concern"],
        d_ENG_002(),"ATL-ENG-001, ATL-ENG-003")

    add_d("ATL-ENG-003","Infrastructure Architecture","Complete infrastructure layout","Volume VII","Platform",
        "Frontend → API Gateway → Application Services → Data Stores → Infrastructure layer. Services communicate through defined interfaces and event bus.",
        ["Infrastructure choices SHALL NOT influence domain semantics"],
        d_ENG_003(),"ATL-ENG-002, ATL-ENG-004")

    add_d("ATL-ENG-004","Database Architecture","Six data storage technologies","Volume VII","Platform",
        "PostgreSQL (core data), Elasticsearch (search), Redis (caching), Object Storage (media/blobs), Graph Database (semantic relations), Time Series DB (metrics).",
        ["Each storage technology serves a specific architectural purpose"],
        d_ENG_004(),"ATL-ENG-003, ATL-IMP-004")

    doc.add_heading("Chapter 8: Implementation Architecture",level=2)
    doc.add_paragraph("Software realization. Volume VIII.")
    doc.add_page_break()

    add_d("ATL-IMP-001","Implementation Domain","Eight components","Volume VIII","Implementation",
        "Components: Repository Structure, Development Workflow, Claude Code Integration, Testing Framework, CI/CD Pipeline, Deployment Strategy, Monitoring Pipeline, Production Architecture.",
        ["AI-first development","Architecture-driven implementation","Planning before coding"],
        d_IMP_001(),"ATL-IMP-002, ATL-IMP-003")

    add_d("ATL-IMP-002","Repository Layout","Project folder structure","Volume VIII","Implementation",
        "Root: rios/ with src/ (identity, knowledge, communication, visualization, motion, platform), tests/, docs/, config/, migrations/, infrastructure/.",
        ["Each domain has its own source directory","Tests mirror source structure"],
        d_IMP_002(),"ATL-IMP-001, ATL-IMP-003")

    add_d("ATL-IMP-003","Development Workflow","Six-stage development process","Volume VIII","Implementation",
        "Architecture Review → Planning & Design → Claude Code Implementation → Testing & Validation → Conformance Check → Deployment. Continuous feedback loop.",
        ["Every implementation begins with architecture review","Claude Code is a first-class development tool"],
        d_IMP_003(),"ATL-IMP-001, ATL-IMP-004")

    add_d("ATL-IMP-004","Deployment Architecture","Three-environment deployment strategy","Volume VIII","Implementation",
        "Development (Docker Compose, hot reload) → Staging (Kubernetes, integration tests) → Production (Kubernetes multi-region, CDN, auto-scaling, monitoring).",
        ["CI/CD pipeline: Build → Test → Lint → Security → Staging → Integration → Production"],
        d_IMP_004(),"ATL-IMP-003, ATL-ENG-003")

    # PART III
    doc.add_heading("PART III: System Relationships",level=1)
    doc.add_paragraph("Cross-domain visualizations showing how all RIOS domains relate to each other.")
    doc.add_page_break()

    add_d("ATL-REF-001","Cross-Domain Dependency Map","Complete inter-domain dependency visualization","All Volumes","System",
        "Identity at the top with no upstream dependencies. Knowledge, Communication downstream. Visualization depends on Knowledge. Motion depends on Identity. Platform supports Implementation.",
        ["All dependencies flow in one direction","Identity is the architectural root"],
        d_REF_001(),"ATL-FOUND-007, ATL-IDN-008")

    add_d("ATL-REF-002","Strategic Context Map","DDD context map pattern for all domains","All Volumes","System",
        "Seven bounded contexts: Identity (upstream core), Knowledge (core), Communication (downstream), Visualization (supporting), Motion (supporting), Platform (supporting), Implementation (generic).",
        ["Context relationships follow DDD upstream/downstream patterns"],
        d_REF_002(),"ATL-IDN-006, ATL-REF-001")

    add_d("ATL-REF-003","Consumer-Provider Graph","Which domains consume from which providers","All Volumes","System",
        "Three providers (Identity, Knowledge, Platform) serve six consumers. Identity provides to Knowledge, Communication, Visualization, Motion. Knowledge provides to Communication, Visualization. Platform provides to Implementation.",
        ["Provider-consumer relationships define integration contracts"],
        d_REF_003(),"ATL-REF-001, ATL-IDN-012")

    # PART IV
    doc.add_heading("PART IV: Knowledge & Research Maps",level=1)
    doc.add_paragraph("Research knowledge visualizations showing lifecycles, evolution, and trust models.")
    doc.add_page_break()

    add_d("ATL-KNO-006","Research Lifecycle","Seven-stage research lifecycle","All Volumes","Knowledge",
        "Research Question → Hypothesis → Experiment Design → Data Collection → Analysis & Reasoning → Knowledge Claim → Publication & Dissemination. Continuous feedback loop.",
        ["The lifecycle is iterative, not linear","Evidence chain connects all stages"],
        d_KNO_006(),"ATL-KNO-003, ATL-IDN-003")

    add_d("ATL-KNO-007","Knowledge Evolution Timeline","Temporal view of knowledge development","Volume II","Knowledge",
        "Seven milestones from Initial Question through Knowledge Transfer. Knowledge complexity increases over time. Growth curve shows accelerating understanding.",
        ["Knowledge evolves through accumulation and refinement"],
        d_KNO_007(),"ATL-KNO-003, ATL-MOT-001")

    add_d("ATL-KNO-008","Academic Trust Model","Trust pyramid for academic credibility","All Volumes","Knowledge",
        "Six-level trust pyramid: Identity Coherence (base) → Research Consistency → Methodological Rigor → Evidence Base → Publication Record → Recognition & Impact (top).",
        ["Trust builds from bottom to top","Each level depends on stability of lower levels","Identity coherence is the foundation of academic trust"],
        d_KNO_008(),"ATL-IDN-001, ATL-COM-002")

    # PART V
    doc.add_heading("PART V: Engineering Atlas",level=1)
    doc.add_paragraph("Engineering visualizations for platform, infrastructure, and database architecture.")
    doc.add_page_break()

    doc.add_paragraph("Engineering diagrams are included in Part II (Chapters 7-8) under Platform Engineering and Implementation domains. See ATL-ENG-002, ATL-ENG-003, ATL-ENG-004 for platform architecture, service architecture, infrastructure, and database architecture.")

    # PART VI
    doc.add_heading("PART VI: Implementation Atlas",level=1)
    doc.add_paragraph("Implementation visualizations for repository, workflow, and deployment.")
    doc.add_page_break()

    doc.add_paragraph("Implementation diagrams are included in Part II (Chapter 8) under Implementation Architecture. See ATL-IMP-002, ATL-IMP-003, ATL-IMP-004 for repository layout, development workflow, and deployment architecture.")

    # PART VII — REFERENCE
    doc.add_heading("PART VII: Reference Atlas",level=1)
    doc.add_paragraph("Glossary, notation guide, diagram legend, and cross-references.")
    doc.add_page_break()

    # Glossary
    doc.add_heading("Architecture Glossary",level=2)
    glossary=[("Aggregate Root","The primary entity in a DDD aggregate that maintains identity and consistency boundary"),
              ("Bounded Context","A DDD pattern defining a clear boundary within which a domain model applies"),
              ("Canonical","The single authoritative source of truth for a concept or entity"),
              ("Conformance","Evaluation of whether an implementation satisfies architectural requirements"),
              ("Domain Event","A record of something significant that happened in the domain"),
              ("Domain Service","A DDD pattern for operations that don't naturally belong to an entity"),
              ("Entity","A domain object with a distinct identity that persists across state changes"),
              ("Invariant","A condition that must always be true for the system to be correct"),
              ("MAB","Master Architecture Blueprint — the governing architecture document"),
              ("Normative","A requirement or specification that must be followed for conformance"),
              ("Research Identity","The complete, canonical representation of who a researcher is"),
              ("Semantic Integrity","Preservation of meaning across all architectural layers and representations"),
              ("Value Object","An immutable domain object defined by its attributes rather than identity"),
              ("Verification","Evaluation of whether the architecture has been implemented correctly"),
              ("RIOS","Research Identity Operating System")]
    for term,defn in glossary:
        p=doc.add_paragraph()
        r=p.add_run(f"{term}: "); r.bold=True; r.font.size=Pt(10)
        r2=p.add_run(defn); r2.font.size=Pt(10)
    doc.add_page_break()

    # Notation Guide
    doc.add_heading("Notation Guide",level=2)
    doc.add_paragraph("All diagrams in this Atlas use the following consistent visual language:")
    notations=[("Box (Navy)","Aggregate root, primary domain entity, or system-level concept"),
               ("Box (Steel Blue)","Entity or domain component"),
               ("Box (Amber/Gold)","Value object or event"),
               ("Box (Green)","Success state, positive constraint, or evidence"),
               ("Box (Red)","Negative constraint, prohibition, or error state"),
               ("Box (Light Gray)","External actor, neutral concept, or passive element"),
               ("Arrow (Solid)","Architectural dependency or information flow"),
               ("Rounded Rectangle","Bounded context or aggregate boundary"),
               ("Dashed Line","Loose coupling or optional relationship")]
    for notation,desc in notations:
        p=doc.add_paragraph()
        r=p.add_run(f"{notation}: "); r.bold=True; r.font.size=Pt(10)
        r2=p.add_run(desc); r2.font.size=Pt(10)
    doc.add_page_break()

    # Diagram Index
    doc.add_heading("Diagram Index",level=2)
    diagrams_idx=[
        ("ATL-FOUND-001","System Overview"),("ATL-FOUND-002","Architecture Vision"),("ATL-FOUND-003","Architecture Principles"),
        ("ATL-FOUND-004","Architecture Layers"),("ATL-FOUND-005","Domain Overview"),("ATL-FOUND-006","Architecture Stack"),
        ("ATL-FOUND-007","Dependency Architecture"),("ATL-FOUND-008","Capability Map"),("ATL-FOUND-009","System Context Diagram"),
        ("ATL-FOUND-010","Architecture Hierarchy"),("ATL-FOUND-011","Architecture Flow"),
        ("ATL-IDN-001","Identity Domain Overview"),("ATL-IDN-002","Identity Bounded Context"),
        ("ATL-IDN-003","Identity Lifecycle"),("ATL-IDN-004","Identity ER Diagram"),
        ("ATL-IDN-005","Identity Architecture Blueprint"),("ATL-IDN-006","Identity Context Map"),
        ("ATL-IDN-007","Identity Capability Map"),("ATL-IDN-008","Identity Domain Dependencies"),
        ("ATL-IDN-009","Identity DDD Diagram"),("ATL-IDN-010","Identity Business Rules"),
        ("ATL-IDN-011","Identity Verification Flow"),("ATL-IDN-012","Identity Interaction Diagram"),
        ("ATL-KNO-001","Knowledge Domain Overview"),("ATL-KNO-002","Knowledge Bounded Context"),
        ("ATL-KNO-003","Knowledge Lifecycle"),("ATL-KNO-004","Knowledge ER Diagram"),
        ("ATL-KNO-005","Knowledge DDD Diagram"),("ATL-KNO-006","Research Lifecycle"),
        ("ATL-KNO-007","Knowledge Evolution Timeline"),("ATL-KNO-008","Academic Trust Model"),
        ("ATL-COM-001","Knowledge Communication Domain"),("ATL-COM-002","Scholarly Communication Domain"),
        ("ATL-VIS-001","Scientific Visualization Domain"),("ATL-MOT-001","Cognitive Motion Domain"),
        ("ATL-ENG-001","Platform Engineering Domain"),("ATL-ENG-002","Platform Service Architecture"),
        ("ATL-ENG-003","Infrastructure Architecture"),("ATL-ENG-004","Database Architecture"),
        ("ATL-IMP-001","Implementation Domain"),("ATL-IMP-002","Repository Layout"),
        ("ATL-IMP-003","Development Workflow"),("ATL-IMP-004","Deployment Architecture"),
        ("ATL-REF-001","Cross-Domain Dependency Map"),("ATL-REF-002","Strategic Context Map"),
        ("ATL-REF-003","Consumer-Provider Graph"),
    ]
    t=doc.add_table(rows=len(diagrams_idx)+1,cols=2); t.style='Table Grid'
    t.rows[0].cells[0].text="Diagram ID"; t.rows[0].cells[1].text="Title"
    for p in t.rows[0].cells[0].paragraphs:
        for r in p.runs: r.bold=True
    for p in t.rows[0].cells[1].paragraphs:
        for r in p.runs: r.bold=True
    for i,(did,dtitle) in enumerate(diagrams_idx):
        t.rows[i+1].cells[0].text=did; t.rows[i+1].cells[1].text=dtitle
    doc.add_page_break()

    # Quality Review Checklist
    doc.add_heading("Quality Review Checklist",level=2)
    checks=[
        "All diagrams use consistent visual language (typography, spacing, arrow styles, colors)",
        "No duplicated concepts across the Atlas",
        "No contradictory diagrams",
        "Consistent terminology throughout (matches Canonical Terminology Dictionary)",
        "All diagrams accurately reflect the uploaded architecture volumes",
        "No architectural drift from source documents",
        "Every diagram has unique ID, title, purpose, and description",
        "Cross-references are accurate and complete",
        "Glossary terms match Canonical Terminology Dictionary",
        "All eight domains are represented",
        "Foundation Architecture is referenced correctly",
        "Dependency directions are consistent across all diagrams",
        "DDD patterns (aggregate root, entity, value object, domain service, domain event) are used correctly",
        "No decorative graphics — all visual elements communicate meaning",
        "Professional formatting suitable for enterprise architecture documentation",
    ]
    for chk in checks:
        doc.add_paragraph(f"☐ {chk}")
    doc.add_page_break()

    # Architecture Validation Checklist
    doc.add_heading("Architecture Validation Checklist",level=2)
    val_checks=[
        "RIOS vision: Research Identity as a Living Scientific System — preserved",
        "Six architectural pillars — all present and correctly represented",
        "Ten architecture principles (AP-001 to AP-010) — all enumerated",
        "Identity Domain: 9-layer blueprint — correctly depicted",
        "Identity invariants (IC-001 to IC-006) — all listed",
        "Semantic constraints (SC-001 to SC-006) — all enforced in dependency diagrams",
        "Structural constraints (ST-001 to ST-004) — all represented",
        "Knowledge Domain: independent lifecycle — correctly shown",
        "Communication domains: narrative + scholarly — both covered",
        "Visualization: explains, never redefines — principle preserved",
        "Motion: preserves identity through evolution — principle preserved",
        "Platform: technology serves architecture — principle preserved",
        "Implementation: AI-first, architecture-driven — correctly represented",
        "Foundation Architecture: all 7 governing documents referenced",
        "Cross-domain dependencies: acyclic — verified across all diagrams",
    ]
    for chk in val_checks:
        doc.add_paragraph(f"☐ {chk}")

    # Save
    out="RIOS_ATLAS.docx"
    doc.save(out)
    print(f"Atlas saved: {out}")
    return out

if __name__=="__main__":
    print("Generating RIOS Atlas diagrams...")
    print("Building document...")
    build_doc()
    print("Done!")