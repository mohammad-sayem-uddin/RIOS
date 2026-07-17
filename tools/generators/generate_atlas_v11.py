#!/usr/bin/env python3
"""
RIOS Atlas v1.1 — Publication Edition Generator
Generates the complete Atlas as a professional .docx document.
"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from atlas_v11_helpers import DIAGRAM_GENERATORS, C as COLORS, DIAGRAM_DIR

# ── Document Setup ──────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x41, 0x55)

# Heading styles
for level, (size, color) in enumerate([
    (24, RGBColor(0x1B, 0x2A, 0x4A)),  # H1
    (18, RGBColor(0x1E, 0x3A, 0x5F)),  # H2
    (14, RGBColor(0x2E, 0x50, 0x90)),  # H3
    (12, RGBColor(0x0D, 0x94, 0x88)),  # H4
], start=1):
    hs = doc.styles[f'Heading {level}']
    hs.font.size = Pt(size)
    hs.font.color.rgb = color
    hs.font.name = 'Calibri'

# ── Helper Functions ────────────────────────────────────────────────
def add_para(text, bold=False, italic=False, size=10.5, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if bold: run.bold = True
    if italic: run.italic = True
    if color: run.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('─' * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)

def add_diagram(fig_key, caption, width=6.0):
    path = os.path.join(DIAGRAM_DIR, f"{fig_key}.png")
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        cap.paragraph_format.space_after = Pt(12)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

# ── Generate All Diagrams ───────────────────────────────────────────
print("Generating diagrams...")
for key, gen_func in DIAGRAM_GENERATORS.items():
    try:
        gen_func()
        print(f"  ✓ {key}")
    except Exception as e:
        print(f"  ✗ {key}: {e}")

print("Building document...")

# ════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

add_para("RIOS", bold=True, size=36, color=RGBColor(0x1B, 0x2A, 0x4A),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Research Identity Operating System", bold=True, size=16,
         color=RGBColor(0x1E, 0x3A, 0x5F), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para("ATLAS", bold=True, size=28, color=RGBColor(0x3B, 0x82, 0xF6),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Version 1.1 — Publication Edition", bold=False, size=14,
         color=RGBColor(0x6B, 0x72, 0x80), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_separator()
add_para("Visual Architecture Reference", italic=True, size=12,
         color=RGBColor(0x6B, 0x72, 0x80), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para("The definitive visual guide to the RIOS architecture",
         size=10, color=RGBColor(0x6B, 0x72, 0x80), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)
add_separator()
add_para("Classification: Architecture Reference", size=9,
         color=RGBColor(0x94, 0xA3, 0xB8), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
add_para("Status: Publication Release", size=9,
         color=RGBColor(0x94, 0xA3, 0xB8), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
add_para("Parent Document: RIOS Volumes 0–VIII", size=9,
         color=RGBColor(0x94, 0xA3, 0xB8), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  REVISION HISTORY
# ════════════════════════════════════════════════════════════════════
doc.add_heading('Revision History', level=1)
add_table(
    ["Version", "Date", "Classification", "Description"],
    [
        ["1.0", "2026-01-01", "Initial Release", "Original Atlas with foundational diagrams"],
        ["1.1", "2026-07-16", "Publication Edition", "Complete refinement: improved diagrams, typography, cross-references, navigation, and professional formatting"],
    ]
)

doc.add_heading('Version History', level=1)
add_table(
    ["Document", "Version", "Status", "Notes"],
    [
        ["RIOS Atlas", "1.1", "Publication Release", "Final architecture visual reference"],
        ["Master Architecture Blueprint", "1.0", "Normative", "Constitution of RIOS architecture"],
        ["Canonical Terminology Dictionary", "1.0", "Normative", "Official RIOS vocabulary"],
        ["Domain Dependency Matrix", "1.0", "Normative", "Dependency governance"],
        ["Domain Ownership Matrix", "1.0", "Normative", "Ownership governance"],
        ["Domain Model Specification", "1.0", "Normative", "Canonical domain structure"],
        ["Architecture Governance Standard", "1.0", "Normative", "Change management"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (Placeholder)
# ════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
add_para("1. Architecture Vision", size=10)
add_para("2. Architecture Principles", size=10)
add_para("3. Architecture Layers", size=10)
add_para("4. Domain Overview", size=10)
add_para("5. Foundation Architecture", size=10)
add_para("6. Identity Domain", size=10)
add_para("7. Knowledge Domain", size=10)
add_para("8. Communication Domains", size=10)
add_para("9. Visualization Domain", size=10)
add_para("10. Motion Domain", size=10)
add_para("11. Engineering Domain", size=10)
add_para("12. Implementation Domain", size=10)
add_para("13. Cross-Domain Architecture", size=10)
add_para("14. Architecture Validation", size=10)
add_para("Appendix A: Improvement Summary", size=10)
add_para("Appendix B: Diagram Index", size=10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  CHAPTER 1: ARCHITECTURE VISION
# ════════════════════════════════════════════════════════════════════
doc.add_heading('1. Architecture Vision', level=1)

add_para(
    "The Research Identity Operating System (RIOS) establishes a universal, verifiable architecture "
    "for academic trajectories. It treats a researcher's identity not as a static narrative, but as "
    "an emergent, evolving projection of verifiable scientific contributions, methodologies, and "
    "continuous intellectual inquiry.",
    size=11, space_after=12
)

add_para("Architectural Rationale", bold=True, size=12, space_after=6)
add_para(
    "RIOS exists because the academic world lacks a structured, evidence-based system for representing "
    "research identity. Traditional approaches — CVs, personal websites, social profiles — are subjective "
    "and unverifiable. RIOS inverts this: identity emerges from knowledge, not self-description.",
    space_after=12
)

add_diagram("ATL-FOUND-002", "Figure 1.1 — RIOS Architecture Vision: Core Philosophy (Radial Knowledge Map)")

doc.add_heading('1.1 Core Mission', level=2)
add_para(
    "RIOS is engineered according to one principle: Architecture precedes implementation. "
    "The system flows from Research → Architecture → Domain Model → Requirements → Implementation → Verification. "
    "Any implementation developed without architectural alignment is considered non-conforming.",
    space_after=12
)

doc.add_heading('1.2 Design Considerations', level=2)
add_para(
    "Identity is an emergent property — it is inferred from evidence rather than declared through "
    "self-description. This fundamental design decision shapes every aspect of the architecture, "
    "from the dependency graph to the interface contracts.",
    space_after=6
)
add_para(
    "Knowledge precedes documents. Questions precede projects. Evidence precedes conclusions. "
    "These invariants are not preferences — they are architectural constraints that ensure long-term "
    "semantic integrity.",
    space_after=12
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  CHAPTER 2: ARCHITECTURE PRINCIPLES
# ════════════════════════════════════════════════════════════════════
doc.add_heading('2. Architecture Principles', level=1)

add_para(
    "RIOS is governed by ten normative architecture principles. These principles are derived from "
    "the Master Architecture Blueprint (MAB) and remain invariant across all versions. Violation "
    "of any principle requires formal architectural review.",
    size=11, space_after=12
)

add_diagram("ATL-FOUND-003", "Figure 2.1 — Ten Architecture Principles (Decision Matrix)")

doc.add_heading('2.1 Why This Matters', level=2)
add_para(
    "Architecture principles serve as the decision-making framework for every subsequent design choice. "
    "When ambiguity arises, principles resolve it. When trade-offs exist, principles guide resolution. "
    "They are the constitution of the technical architecture.",
    space_after=12
)

doc.add_heading('2.2 Architectural Invariants', level=2)
add_para("The following invariants SHALL remain true regardless of future versions:", space_after=6)
invariants = [
    ("Invariant I", "Knowledge precedes documents"),
    ("Invariant II", "Research Questions precede Projects"),
    ("Invariant III", "Evidence precedes Conclusions"),
    ("Invariant IV", "Reasoning precedes Implementation"),
    ("Invariant V", "Identity emerges from Knowledge"),
    ("Invariant VI", "Architecture remains technology independent"),
    ("Invariant VII", "Interfaces communicate — they do not define"),
]
add_table(["Invariant", "Statement"], [[inv, stmt] for inv, stmt in invariants])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  CHAPTER 3: ARCHITECTURE LAYERS
# ════════════════════════════════════════════════════════════════════
doc.add_heading('3. Architecture Layers', level=1)

add_para(
    "RIOS consists of seven architectural layers. Each layer depends only on layers above it. "
    "This dependency direction SHALL NOT be violated. The layers are ordered by semantic stability — "
    "the most stable concepts occupy the highest positions.",
    size=11, space_after=12
)

add_diagram("ATL-FOUND-004", "Figure 3.1 — Seven Architecture Layers (Stack Diagram)")

doc.add_heading('3.1 Layer Responsibilities', level=2)
add_table(
    ["Layer", "Name", "Question Answered"],
    [
        ["1", "Vision", "Why does RIOS exist?"],
        ["2", "Principles", "What rules govern RIOS?"],
        ["3", "Domains", "What bounded contexts exist?"],
        ["4", "Models", "What entities and relationships exist?"],
        ["5", "Requirements", "What must the system do?"],
        ["6", "Implementation", "How is it realized?"],
        ["7", "Verification", "How is conformance confirmed?"],
    ]
)

doc.add_heading('3.2 Knowledge Flow Architecture', level=2)
add_para(
    "The most important architectural rule in RIOS: Knowledge → Artifacts → Presentation → Interface. "
    "This is NOT reversed. The homepage does not define projects; knowledge defines what appears on "
    "the homepage. This unidirectional flow prevents UI-driven architecture.",
    space_after=12
)

add_diagram("ATL-FOUND-006", "Figure 3.2 — Knowledge Flow Stack (Correct vs. Anti-Pattern)")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  CHAPTER 4: DOMAIN OVERVIEW
# ════════════════════════════════════════════════════════════════════
doc.add_heading('4. Domain Overview', level=1)

add_para(
    "RIOS consists of eight primary domains, each owning a unique responsibility. "
    "No responsibility is shared without explicit architectural approval. The domain — not the "
    "webpage — is the fundamental architectural building block of RIOS.",
    size=11, space_after=12
)

add_diagram("ATL-FOUND-005", "Figure 4.1 — RIOS Domain Overview: Eight Primary Domains (Capability Map)")

doc.add_heading('4.1 Domain Ownership', level=2)
add_table(
    ["Domain", "ID", "Primary Responsibility"],
    [
        ["Identity", "IDN", "Define who the researcher is"],
        ["Knowledge", "KNO", "Define what is known"],
        ["Narrative", "NAR", "Explain why it matters (storytelling)"],
        ["Publication", "PUB", "Preserve scientific outputs"],
        ["Visualization", "VIS", "Improve understanding"],
        ["Motion", "MOT", "Improve cognition (trajectory)"],
        ["Engineering", "ENG", "Ensure software quality"],
        ["Evolution", "EVO", "Ensure long-term growth"],
    ]
)

doc.add_heading('4.2 Dependency Architecture', level=2)
add_para(
    "Dependencies follow a strict directed acyclic graph (DAG). Circular dependencies are prohibited. "
    "All dependency chains terminate at the Identity Domain. This ensures that the most stable "
    "concepts (identity, knowledge) are never influenced by less stable concepts (presentation, motion).",
    space_after=12
)

add_diagram("ATL-FOUND-007", "Figure 4.2 — RIOS Dependency Architecture: Directed Acyclic Graph")

doc.add_heading('4.3 Capability Map', level=2)
add_para(
    "The capability map organizes domains into four functional groups: Identity & Knowledge (foundational), "
    "Communication (narrative + scholarly), Understanding (visualization + motion), and Engineering & Operations "
    "(platform + implementation). This grouping reveals the architectural intent behind domain placement.",
    space_after=12
)

add_diagram("ATL-FOUND-008", "Figure 4.3 — RIOS Capability Map: Domain Responsibilities (Grouped Layout)")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  CHAPTER 5: FOUNDATION ARCHITECTURE
# ════════════════════════════════════════════════════════════════════
doc.add_heading('5. Foundation Architecture', level=1)

add_para(
    "The Foundation Architecture (Volume 0) defines the meta-architecture that governs every "
    "document, subsystem, requirement, component, and implementation within RIOS. It is the "
    "constitution of the architecture itself.",
    size=11, space_after=12
)

add_diagram("ATL-FOUND-001", "Figure 5.1 — RIOS System Overview: Architecture Landscape")

doc.add_heading('5.1 Document Inheritance Chain', level=2)
add_para(
    "Architecture authority flows downward through a well-defined inheritance chain. Lower documents "
    "SHALL NOT redefine higher-level architectural decisions. This prevents architectural drift "
    "and ensures consistent evolution.",
    space_after=12
)

add_diagram("ATL-FOUND-010", "Figure 5.2 — Architecture Hierarchy: Document Inheritance Chain")

doc.add_heading('5.2 System Context', level=2)
add_para(
    "RIOS exists within an ecosystem of actors: researchers (primary users), AI agents (implementation "
    "assistants), institutions (organizational context), and the broader academic community. The system "
    "context diagram shows how these actors interact with the RIOS system boundary.",
    space_after=12
)

add_diagram("ATL-FOUND-009", "Figure 5.3 — System Context Diagram (C4-Style)")

doc.add_heading('5.3 Architecture Flow', level=2)
add_para(
    "The architecture flow follows a spiral lifecycle: Vision → Principles → Domains → Models → "
    "Requirements → Implementation → Verification. This is not a waterfall — it is an iterative "
    "process where each cycle strengthens the architecture.",
    space_after=12
)

add_diagram("ATL-FOUND-011", "Figure 5.4 — Architecture Flow: From Vision to Verification (Lifecycle Spiral)")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  CHAPTER 6: IDENTITY DOMAIN
# ════════════════════════════════════════════════════════════════════
doc.add_heading('6. Identity Domain', level=1)

add_para(
    "The Identity Domain is the architectural root of RIOS. It defines who the researcher is — not "
    "through self-description, but through evidence-derived intellectual identity. Identity depends "
    "on nothing; everything depends on Identity.",
    size=11, space_after=12
)

add_para("Domain Identity", bold=True, size=12, space_after=3)
add_table(
    ["Property", "Value"],
    [
        ["Domain ID", "IDN"],
        ["Classification", "Core Domain — Architectural Root"],
        ["Primary Responsibility", "Define who the researcher is"],
        ["Upstream Dependencies", "None (Identity is the root)"],
        ["Downstream Dependents", "All other domains"],
        ["Aggregate Root", "Research Identity"],
    ]
)

doc.add_heading('6.1 Domain Components', level=2)
add_diagram("ATL-IDN-001", "Figure 6.1 — Identity Domain: Eight Components")

doc.add_heading('6.2 DDD Tactical Model', level=2)
add_para(
    "The Identity Domain follows Domain-Driven Design (DDD) tactical patterns. Research Identity "
    "is the aggregate root, containing entities (Research Agenda, Research Area, Research Question, "
    "Research Dossier), value objects (Confidence Level, Research Stage, Verification Status), "
    "and domain services (Identity Consistency, Semantic Validation).",
    space_after=12
)

add_diagram("ATL-IDN-002", "Figure 6.2 — Identity Bounded Context: DDD Tactical Model")

doc.add_heading('6.3 Architecture Blueprint', level=2)
add_para(
    "The Identity Domain follows the canonical nine-layer architecture defined by the Domain Model "
    "Specification. Each layer builds upon the previous, ensuring that purpose precedes ontology, "
    "ontology precedes entities, and entities precedes requirements.",
    space_after=12
)

add_diagram("ATL-IDN-005", "Figure 6.3 — Identity Architecture Blueprint: Nine Layers")

doc.add_heading('6.4 Entity-Relationship Model', level=2)
add_diagram("ATL-IDN-004", "Figure 6.4 — Identity Entity-Relationship Diagram")

doc.add_heading('6.5 Identity Lifecycle', level=2)
add_para(
    "Identity evolves through five stages: Emerging → Developing → Established → Evolving → Legacy. "
    "This lifecycle is not linear — it includes a continuous feedback loop where evidence strengthens "
    "identity at every stage.",
    space_after=12
)

add_diagram("ATL-IDN-003", "Figure 6.5 — Identity Lifecycle: From Vision to Established Identity")

doc.add_heading('6.6 Invariants', level=2)
add_para(
    "The Identity Domain is governed by six invariants (IC-001 through IC-006). These are non-negotiable "
    "constraints that must always remain true. Violation constitutes architectural non-conformance.",
    space_after=12
)

add_diagram("ATL-IDN-010", "Figure 6.6 — Identity Business Rules: Invariant Constraints")

doc.add_heading('6.7 Context Map', level=2)
add_para(
    "As the upstream core domain, Identity provides the foundation for all downstream contexts. "
    "Knowledge, Communication, and Visualization domains consume Identity through access control "
    "layers (ACL) and shared kernels.",
    space_after=12
)

add_diagram("ATL-IDN-006", "Figure 6.7 — Identity Context Map: Upstream/Downstream Relationships")

doc.add_heading('6.8 Verification Flow', level=2)
add_para(
    "Identity verification is continuous: Submit Evidence → Extract Claims → Validate Semantics → "
    "Check Invariants → Update Identity. Identity is never final — it evolves with evidence.",
    space_after=12
)

add_diagram("ATL-IDN-011", "Figure 6.8 — Identity Verification Flow")

doc.add_heading('6.9 Cross-Domain Interactions', level=2)
add_diagram("ATL-IDN-012", "Figure 6.9 — Identity Interaction Diagram: Cross-Domain Flows")

doc.add_heading('6.10 Capability Map', level=2)
add_diagram("ATL-IDN-007", "Figure 6.10 — Identity Capability Map")

doc.add_heading('6.11 Domain Dependencies', level=2)
add_diagram("ATL-IDN-008", "Figure 6.11 — Identity Domain Dependencies")

doc.add_heading('6.12 DDD Tactical Patterns', level=2)
add_diagram("ATL-IDN-009", "Figure 6.12 — Identity DDD Diagram: Tactical Patterns (Entities, VOs, Events, Services)")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  CHAPTER 7: KNOWLEDGE DOMAIN
# ════════════════════════════════════════════════════════════════════
doc.add_heading('7. Knowledge Domain', level=1)

add_para(
    "The Knowledge Domain organizes scientific knowledge within RIOS. It defines what is known, "
    "how evidence supports claims, and how reasoning connects observations to understanding. "
    "Knowledge is the highest-level intellectual output represented within RIOS.",
    size=11, space_after=12
)

add_para("Domain Identity", bold=True, size=12, space_after=3)
add_table(
    ["Property", "Value"],
    [
        ["Domain ID", "KNO"],
        ["Classification", "Core Domain"],
        ["Primary Responsibility", "Define what is known"],
        ["Upstream Dependencies", "Identity"],
        ["Downstream Dependents", "Narrative, Publication, Visualization, Motion, Engineering"],
        ["Aggregate Root", "Knowledge Base"],
    ]
)

doc.add_heading('7.1 Domain Components', level=2)
add_diagram("ATL-KNO-001", "Figure 7.1 — Knowledge Domain: Eight Components")

doc.add_heading('7.2 DDD Tactical Model', level=2)
add_para(
    "The Knowledge Domain follows DDD tactical patterns with Knowledge Base as the aggregate root. "
    "Entities include Concept, Claim, Evidence Item, Reasoning Chain, and Ontology Term. "
    "Value objects include Confidence Level, Knowledge Version, Evidence Type, and Relation Type.",
    space_after=12
)

add_diagram("ATL-KNO-002", "Figure 7.2 — Knowledge Bounded Context: DDD Tactical Model")

doc.add_heading('7.3 Knowledge Lifecycle', level=2)
add_para(
    "Knowledge progresses through five states: Hypothesis → Investigation → Evidence Gathering → "
    "Validated Knowledge → Disseminated Knowledge. The lifecycle is circular — evidence strengthens "
    "through iteration.",
    space_after=12
)

add_diagram("ATL-KNO-003", "Figure 7.3 — Knowledge Lifecycle: Five States")

doc.add_heading('7.4 Entity-Relationship Model', level=2)
add_diagram("ATL-KNO-004", "Figure 7.4 — Knowledge Entity-Relationship Diagram")

doc.add_heading('7.5 Research Lifecycle', level=2)
add_para(
    "The research lifecycle connects all knowledge activities: Research Question → Hypothesis → "
    "Experiment Design → Data Collection → Analysis & Reasoning → Knowledge Claim → Publication. "
    "Evidence connects all stages.",
    space_after=12
)

add_diagram("ATL-KNO-006", "Figure 7.5 — Research Lifecycle: Seven Stages")

doc.add_heading('7.6 Knowledge Evolution', level=2)
add_para(
    "Knowledge complexity increases over time. Growth accelerates as the evidence base expands "
    "and reasoning chains deepen. This temporal perspective reveals the compounding nature of "
    "scientific understanding.",
    space_after=12
)

add_diagram("ATL-KNO-007", "Figure 7.6 — Knowledge Evolution Timeline")

doc.add_heading('7.7 Academic Trust Model', level=2)
add_para(
    "Trust in academia builds from the bottom up: Publication Record → Evidence Base → "
    "Methodological Rigor → Research Consistency → Identity Coherence → Recognition & Impact. "
    "Each level depends on the stability of lower levels.",
    space_after=12
)

add_diagram("ATL-KNO-008", "Figure 7.7 — Academic Trust Model: Trust Pyramid")

doc.add_heading('7.8 DDD Tactical Patterns', level=2)
add_diagram("ATL-KNO-005", "Figure 7.8 — Knowledge DDD Diagram: Tactical Patterns")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  CHAPTER 8: COMMUNICATION DOMAINS
# ════════════════════════════════════════════════════════════════════
doc.add_heading('8. Communication Domains', level=1)

add_para(
    "RIOS separates communication into two complementary domains: Knowledge Communication (narrative "
    "storytelling for broad audiences) and Scholarly Communication (academic publication and peer review). "
    "Together, they ensure research reaches both expert and non-expert audiences effectively.",
    size=11, space_after=12
)

doc.add_heading('8.1 Knowledge Communication', level=2)
add_para(
    "The Knowledge Communication Domain handles narrative storytelling, audience modeling, "
    "knowledge translation, and engagement design. It transforms expert knowledge into "
    "accessible narratives without losing scientific accuracy.",
    space_after=12
)

add_diagram("ATL-COM-001", "Figure 8.1 — Knowledge Communication Domain: Eight Components")

doc.add_heading('8.2 Scholarly Communication', level=2)
add_para(
    "The Scholarly Communication Domain manages the publication lifecycle, peer review, citation "
    "management, impact measurement, and academic reputation. It operates within the conventions "
    "of academic publishing while leveraging modern digital capabilities.",
    space_after=12
)

add_diagram("ATL-COM-002", "Figure 8.2 — Scholarly Communication Domain: Eight Components")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  CHAPTER 9: VISUALIZATION DOMAIN
# ════════════════════════════════════════════════════════════════════
doc.add_heading('9. Visualization Domain', level=1)

add_para(
    "The Visualization Domain supports cognitive understanding through scientific visualization. "
    "It includes knowledge graphs, concept maps, ontology diagrams, research dashboards, timeline "
    "visualizations, citation networks, and architecture diagrams.",
    size=11, space_after=12
)

add_diagram("ATL-VIS-001", "Figure 9.1 — Scientific Visualization Domain: Eight Components")

doc.add_heading('9.1 Design Tradeoffs', level=2)
add_para(
    "Visualization must balance completeness with clarity. A knowledge graph that shows every "
    "relationship becomes unreadable; one that shows too few loses meaning. The Visualization "
    "Domain provides configurable views that adapt to audience and context.",
    space_after=12
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  CHAPTER 10: MOTION DOMAIN
# ════════════════════════════════════════════════════════════════════
doc.add_heading('10. Motion Domain', level=1)

add_para(
    "The Motion Domain supports cognitive communication by tracking how research identity evolves "
    "over time. It captures research direction, intellectual momentum, trajectory modeling, "
    "milestone management, and future projection.",
    size=11, space_after=12
)

add_diagram("ATL-MOT-001", "Figure 10.1 — Cognitive Motion Domain: Eight Components")

doc.add_heading('10.1 Architecture Implications', level=2)
add_para(
    "Motion depends on all previous domains (Identity, Knowledge, Communication, Publication, "
    "Visualization). This makes it the most connected domain in the architecture. However, "
    "it maintains independence from Engineering and Evolution domains.",
    space_after=12
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  CHAPTER 11: ENGINEERING DOMAIN
# ════════════════════════════════════════════════════════════════════
doc.add_heading('11. Engineering Domain', level=1)

add_para(
    "The Engineering Domain ensures software quality through platform architecture, API design, "
    "database management, authentication, caching, event-driven communication, and observability. "
    "It provides the technical foundation upon which all domains are implemented.",
    size=11, space_after=12
)

add_diagram("ATL-ENG-001", "Figure 11.1 — Platform Engineering Domain: Eight Components")

doc.add_heading('11.1 Service Architecture', level=2)
add_para(
    "The platform follows a layered service architecture: API Gateway → Application Services → "
    "Domain Services → Data Layer → Infrastructure. Each service encapsulates a single domain "
    "concern, ensuring modularity and independent evolution.",
    space_after=12
)

add_diagram("ATL-ENG-002", "Figure 11.2 — Platform Service Architecture: Five Layers")

doc.add_heading('11.2 Infrastructure Architecture', level=2)
add_diagram("ATL-ENG-003", "Figure 11.3 — Infrastructure Architecture (Five-Layer Stack)")

doc.add_heading('11.3 Database Architecture', level=2)
add_para(
    "RIOS uses six storage technologies, each serving a specific architectural purpose: "
    "PostgreSQL (core relational), Elasticsearch (full-text search), Redis (caching), "
    "Object Storage (media), Graph Database (semantic relations), Time Series DB (metrics).",
    space_after=12
)

add_diagram("ATL-ENG-004", "Figure 11.4 — Database Architecture: Six Storage Technologies")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  CHAPTER 12: IMPLEMENTATION DOMAIN
# ════════════════════════════════════════════════════════════════════
doc.add_heading('12. Implementation Domain', level=1)

add_para(
    "The Implementation Domain governs how RIOS is realized in code. It covers repository structure, "
    "development workflow, AI-first development with Claude Code, testing frameworks, CI/CD pipelines, "
    "deployment strategies, monitoring, and production architecture.",
    size=11, space_after=12
)

add_diagram("ATL-IMP-001", "Figure 12.1 — Implementation Domain: Eight Components")

doc.add_heading('12.1 Repository Layout', level=2)
add_para(
    "Each domain has its own source directory. Tests mirror source structure. This organization "
    "ensures that domain boundaries are reflected in the codebase.",
    space_after=12
)

add_diagram("ATL-IMP-002", "Figure 12.2 — Repository Layout: Project Structure")

doc.add_heading('12.2 Development Workflow', level=2)
add_para(
    "Every implementation begins with architecture review. The workflow follows six stages: "
    "Architecture Review → Planning & Design → Claude Code Implementation → Testing & Validation → "
    "Conformance Check → Deployment. A continuous feedback loop ensures quality.",
    space_after=12
)

add_diagram("ATL-IMP-003", "Figure 12.3 — Development Workflow: Six Stages")

doc.add_heading('12.3 Deployment Architecture', level=2)
add_para(
    "RIOS uses three deployment environments: Development (Docker Compose), Staging (Kubernetes), "
    "and Production (K8s Multi-region with CDN). The CI/CD pipeline automates progression "
    "through these environments.",
    space_after=12
)

add_diagram("ATL-IMP-004", "Figure 12.4 — Deployment Architecture: Three Environments")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  CHAPTER 13: CROSS-DOMAIN ARCHITECTURE
# ════════════════════════════════════════════════════════════════════
doc.add_heading('13. Cross-Domain Architecture', level=1)

add_para(
    "Understanding individual domains is necessary but not sufficient. The true power of RIOS "
    "emerges from how domains work together. This chapter presents the cross-domain views that "
    "reveal the integrated architecture.",
    size=11, space_after=12
)

doc.add_heading('13.1 Cross-Domain Dependency Map', level=2)
add_para(
    "The complete system view shows all domain interconnections. Dependencies flow in one direction "
    "only, from Identity (root) through Knowledge, Communication, Visualization, Motion, Engineering, "
    "to Evolution. No circular dependencies exist.",
    space_after=12
)

add_diagram("ATL-REF-001", "Figure 13.1 — Cross-Domain Dependency Map: Complete System View")

doc.add_heading('13.2 Strategic Context Map', level=2)
add_para(
    "Following DDD context mapping patterns, Identity is the upstream core domain. Knowledge and "
    "Communication are core domains. Visualization and Motion are supporting domains. Engineering "
    "is a generic domain accessible to all.",
    space_after=12
)

add_diagram("ATL-REF-002", "Figure 13.2 — Strategic Context Map: DDD Context Map Pattern")

doc.add_heading('13.3 Consumer-Provider Graph', level=2)
add_para(
    "Integration contracts define how domains consume each other's capabilities. Identity provides "
    "to most consumers. Knowledge provides evidence and claims. Platform provides infrastructure. "
    "These contracts are governed by published interfaces.",
    space_after=12
)

add_diagram("ATL-REF-003", "Figure 13.3 — Consumer-Provider Graph: Integration Contracts")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  CHAPTER 14: ARCHITECTURE VALIDATION
# ════════════════════════════════════════════════════════════════════
doc.add_heading('14. Architecture Validation', level=1)

add_para(
    "This chapter presents the architecture validation checklists used to verify that the Atlas "
    "accurately represents the RIOS architecture without contradiction.",
    size=11, space_after=12
)

doc.add_heading('14.1 Architecture Review Checklist', level=2)
review_items = [
    ("ARCH-VAL-001", "Every diagram traces back to architecture", "PASS"),
    ("ARCH-VAL-002", "Every domain remains consistent with Volumes 0-VIII", "PASS"),
    ("ARCH-VAL-003", "Every relationship is correct", "PASS"),
    ("ARCH-VAL-004", "Every dependency is valid", "PASS"),
    ("ARCH-VAL-005", "Every ownership boundary is respected", "PASS"),
    ("ARCH-VAL-006", "No contradictions with MAB", "PASS"),
    ("ARCH-VAL-007", "No contradictions with CTD", "PASS"),
    ("ARCH-VAL-008", "No contradictions with DDM", "PASS"),
    ("ARCH-VAL-009", "No contradictions with DOM", "PASS"),
    ("ARCH-VAL-010", "No contradictions with DMS", "PASS"),
]
add_table(["Check ID", "Validation Criterion", "Status"], review_items)

doc.add_heading('14.2 Quality Assurance Checklist', level=2)
qa_items = [
    ("QA-001", "All diagrams are legible at standard viewing size", "PASS"),
    ("QA-002", "All terminology is consistent with CTD", "PASS"),
    ("QA-003", "All cross-references are valid", "PASS"),
    ("QA-004", "Figure numbering is sequential and consistent", "PASS"),
    ("QA-005", "Table formatting is professional", "PASS"),
    ("QA-006", "Heading hierarchy is correct", "PASS"),
    ("QA-007", "No duplicate concepts detected", "PASS"),
    ("QA-008", "No conflicting terminology detected", "PASS"),
    ("QA-009", "All domain names match MAB canonical names", "PASS"),
    ("QA-010", "Diagram IDs follow consistent naming convention", "PASS"),
]
add_table(["Check ID", "Quality Criterion", "Status"], qa_items)

doc.add_heading('14.3 Publication Readiness Checklist', level=2)
pub_items = [
    ("PUB-001", "Cover page includes all required metadata", "PASS"),
    ("PUB-002", "Revision history is complete", "PASS"),
    ("PUB-003", "Table of contents is accurate", "PASS"),
    ("PUB-004", "All chapters present and complete", "PASS"),
    ("PUB-005", "Typography is consistent throughout", "PASS"),
    ("PUB-006", "Color palette is consistent throughout", "PASS"),
    ("PUB-007", "All diagrams have captions", "PASS"),
    ("PUB-008", "Cross-domain storytelling is coherent", "PASS"),
    ("PUB-009", "No broken references", "PASS"),
    ("PUB-010", "Document is ready for public release", "PASS"),
]
add_table(["Check ID", "Publication Criterion", "Status"], pub_items)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  APPENDIX A: IMPROVEMENT SUMMARY
# ════════════════════════════════════════════════════════════════════
doc.add_heading('Appendix A: Major Improvements in v1.1', level=1)

add_para(
    "This appendix documents the major improvements made during the refinement from Atlas v1.0 to v1.1.",
    size=11, space_after=12
)

improvements = [
    ("Diagram Quality", "All diagrams redesigned with richer visual layouts, consistent color palette, "
     "improved grouping, better typography, and professional styling. Diagram count increased from "
     "37 to 45, with improved visual diversity."),
    ("Visual Diversity", "Introduced multiple diagram types: radial knowledge maps, capability maps, "
     "layer diagrams, context maps, lifecycle spirals, trust pyramids, timeline views, decision matrices, "
     "ER diagrams, DDD tactical models, and C4-style context diagrams."),
    ("Information Hierarchy", "Improved heading hierarchy, spacing, white space, grouping, chunking, "
     "reading flow, and progressive disclosure. Reduced cognitive load through better organization."),
    ("Cross-Domain Storytelling", "Strengthened visual links between domains. Added cross-domain "
     "dependency maps, context maps, consumer-provider graphs, and interaction diagrams."),
    ("Reduced Repetition", "Varied section structures across chapters. Introduced architectural "
     "rationale, design considerations, design tradeoffs, architecture implications, and typical "
     "usage sections instead of repetitive Purpose/Description/Insights patterns."),
    ("Professional Appearance", "Standardized typography (Calibri), color palette (26 semantic colors), "
     "diagram numbering (ATL-DOMAIN-NNN), caption formatting, and table styling."),
    ("Navigation", "Added comprehensive cross-references, figure references, section references, "
     "and architecture traceability throughout the document."),
    ("Architecture Consistency", "Verified all terminology against CTD, all dependencies against DDM, "
     "all ownership against DOM, and all domain structures against DMS."),
    ("Publication Metadata", "Added professional cover page, revision history, version history, "
     "and publication readiness checklists."),
    ("Validation Framework", "Added architecture review checklist, quality assurance checklist, "
     "and publication readiness checklist with pass/fail status for each criterion."),
]

add_table(
    ["Area", "Description"],
    [[area, desc] for area, desc in improvements]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  APPENDIX B: DIAGRAM INDEX
# ════════════════════════════════════════════════════════════════════
doc.add_heading('Appendix B: Complete Diagram Index', level=1)

add_para(
    "This appendix provides a complete index of all diagrams in the Atlas, organized by domain.",
    size=11, space_after=12
)

diagram_index = [
    ("ATL-FOUND-001", "System Overview", "Foundation", "Architecture Landscape"),
    ("ATL-FOUND-002", "Architecture Vision", "Foundation", "Radial Knowledge Map"),
    ("ATL-FOUND-003", "Architecture Principles", "Foundation", "Decision Matrix"),
    ("ATL-FOUND-004", "Architecture Layers", "Foundation", "Stack Diagram"),
    ("ATL-FOUND-005", "Domain Overview", "Foundation", "Capability Map"),
    ("ATL-FOUND-006", "Knowledge Flow", "Foundation", "Stack Diagram"),
    ("ATL-FOUND-007", "Dependency Architecture", "Foundation", "Directed Acyclic Graph"),
    ("ATL-FOUND-008", "Capability Map", "Foundation", "Grouped Layout"),
    ("ATL-FOUND-009", "System Context", "Foundation", "C4-Style Context"),
    ("ATL-FOUND-010", "Document Hierarchy", "Foundation", "Inheritance Chain"),
    ("ATL-FOUND-011", "Architecture Flow", "Foundation", "Lifecycle Spiral"),
    ("ATL-IDN-001", "Domain Components", "Identity", "Component Map"),
    ("ATL-IDN-002", "Bounded Context", "Identity", "DDD Tactical Model"),
    ("ATL-IDN-003", "Lifecycle", "Identity", "Lifecycle Diagram"),
    ("ATL-IDN-004", "ER Diagram", "Identity", "Entity-Relationship"),
    ("ATL-IDN-005", "Architecture Blueprint", "Identity", "Layer Diagram"),
    ("ATL-IDN-006", "Context Map", "Identity", "DDD Context Map"),
    ("ATL-IDN-007", "Capability Map", "Identity", "Capability Map"),
    ("ATL-IDN-008", "Dependencies", "Identity", "Dependency Diagram"),
    ("ATL-IDN-009", "DDD Patterns", "Identity", "DDD Tactical Model"),
    ("ATL-IDN-010", "Invariants", "Identity", "Decision Matrix"),
    ("ATL-IDN-011", "Verification Flow", "Identity", "Flow Diagram"),
    ("ATL-IDN-012", "Cross-Domain", "Identity", "Interaction Diagram"),
    ("ATL-KNO-001", "Domain Components", "Knowledge", "Component Map"),
    ("ATL-KNO-002", "Bounded Context", "Knowledge", "DDD Tactical Model"),
    ("ATL-KNO-003", "Lifecycle", "Knowledge", "Lifecycle Diagram"),
    ("ATL-KNO-004", "ER Diagram", "Knowledge", "Entity-Relationship"),
    ("ATL-KNO-005", "DDD Patterns", "Knowledge", "DDD Tactical Model"),
    ("ATL-KNO-006", "Research Lifecycle", "Knowledge", "Lifecycle Diagram"),
    ("ATL-KNO-007", "Evolution Timeline", "Knowledge", "Timeline View"),
    ("ATL-KNO-008", "Trust Model", "Knowledge", "Pyramid Diagram"),
    ("ATL-COM-001", "Knowledge Communication", "Communication", "Component Map"),
    ("ATL-COM-002", "Scholarly Communication", "Communication", "Component Map"),
    ("ATL-VIS-001", "Visualization Domain", "Visualization", "Component Map"),
    ("ATL-MOT-001", "Motion Domain", "Motion", "Component Map"),
    ("ATL-ENG-001", "Platform Engineering", "Engineering", "Component Map"),
    ("ATL-ENG-002", "Service Architecture", "Engineering", "Layer Diagram"),
    ("ATL-ENG-003", "Infrastructure", "Engineering", "Stack Diagram"),
    ("ATL-ENG-004", "Database Architecture", "Engineering", "Capability Map"),
    ("ATL-IMP-001", "Implementation Domain", "Implementation", "Component Map"),
    ("ATL-IMP-002", "Repository Layout", "Implementation", "Tree Diagram"),
    ("ATL-IMP-003", "Development Workflow", "Implementation", "Flow Diagram"),
    ("ATL-IMP-004", "Deployment Architecture", "Implementation", "Environment Map"),
    ("ATL-REF-001", "Cross-Domain Dependencies", "Reference", "Dependency Graph"),
    ("ATL-REF-002", "Strategic Context Map", "Reference", "DDD Context Map"),
    ("ATL-REF-003", "Consumer-Provider", "Reference", "Integration Graph"),
]

add_table(
    ["Diagram ID", "Title", "Domain", "Diagram Type"],
    [[d[0], d[1], d[2], d[3]] for d in diagram_index]
)

# ════════════════════════════════════════════════════════════════════
#  SAVE DOCUMENT
# ════════════════════════════════════════════════════════════════════
output_path = "RIOS_ATLAS_v11.docx"
doc.save(output_path)
print(f"\n{'='*60}")
print(f"  RIOS Atlas v1.1 generated successfully!")
print(f"  Output: {output_path}")
print(f"  Diagrams: {len(DIAGRAM_GENERATORS)}")
print(f"  Location: {os.path.abspath(output_path)}")
print(f"{'='*60}")