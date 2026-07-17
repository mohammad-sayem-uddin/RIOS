#!/usr/bin/env python3
"""RIOS Repository Transformation Script
Converts DOCX/TXT to Markdown, reorganizes structure, archives old files.
"""
import os
import shutil
import subprocess
import zipfile
from pathlib import Path
from docx import Document

ROOT = Path("/Users/sayemuddin/Desktop/RIOS")

# ============================================================
# PHASE 1: Convert DOCX files to Markdown
# ============================================================

def docx_to_markdown(docx_path):
    """Convert a DOCX file to Markdown string."""
    doc = Document(str(docx_path))
    md_lines = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            md_lines.append("")
            continue
            
        style = para.style.name if para.style else ""
        
        if "Heading 1" in style:
            md_lines.append(f"# {text}")
        elif "Heading 2" in style:
            md_lines.append(f"## {text}")
        elif "Heading 3" in style:
            md_lines.append(f"### {text}")
        elif "Heading 4" in style:
            md_lines.append(f"#### {text}")
        elif "List" in style or "Bullet" in style:
            # Check for bold runs
            bold_text = ""
            for run in para.runs:
                if run.bold:
                    bold_text += f"**{run.text}**"
                else:
                    bold_text += run.text
            md_lines.append(f"- {bold_text}")
        else:
            # Handle inline formatting
            formatted = ""
            for run in para.runs:
                t = run.text
                if run.bold and run.italic:
                    t = f"***{t}***"
                elif run.bold:
                    t = f"**{t}**"
                elif run.italic:
                    t = f"*{t}*"
                formatted += t
            md_lines.append(formatted)
    
    # Handle tables
    for table in doc.tables:
        md_lines.append("")
        headers = []
        for cell in table.rows[0].cells:
            headers.append(cell.text.strip())
        md_lines.append("| " + " | ".join(headers) + " |")
        md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            md_lines.append("| " + " | ".join(cells) + " |")
        md_lines.append("")
    
    return "\n".join(md_lines)


def convert_all_docx():
    """Convert all DOCX files in the repo to Markdown."""
    conversions = []
    
    # Root-level DOCX files
    root_docx = [
        "RIOS_ATLAS.docx",
        "RIOS_ATLAS_v11.docx",
        "RIOS_Architecture_Decision_Records_v1.0.docx",
        "RIOS_Architecture_Traceability_Matrix_v1.0.docx",
        "RIOS_Architecture_Validation_Report_v1.0.docx",
        "RIOS_Claude_Code_Constitution_v1.0.docx",
        "Volume 0.docx",
    ]
    
    for docx_name in root_docx:
        docx_path = ROOT / docx_name
        if docx_path.exists():
            md_name = docx_name.replace(".docx", ".md").replace(" ", "-")
            md_path = ROOT / md_name
            print(f"Converting: {docx_name} -> {md_name}")
            try:
                md_content = docx_to_markdown(docx_path)
                md_path.write_text(md_content, encoding="utf-8")
                conversions.append((str(docx_path), str(md_path)))
            except Exception as e:
                print(f"  ERROR: {e}")
    
    # Foundation DOCX files
    foundation_dir = ROOT / "Foundation"
    if foundation_dir.exists():
        for docx_file in foundation_dir.glob("*.docx"):
            md_name = docx_file.stem.replace(" ", "-") + ".md"
            md_path = ROOT / "docs" / "architecture" / "foundation" / md_name
            md_path.parent.mkdir(parents=True, exist_ok=True)
            print(f"Converting: Foundation/{docx_file.name} -> docs/architecture/foundation/{md_name}")
            try:
                md_content = docx_to_markdown(docx_file)
                md_path.write_text(md_content, encoding="utf-8")
                conversions.append((str(docx_file), str(md_path)))
            except Exception as e:
                print(f"  ERROR: {e}")
    
    # Volumes DOCX files
    volumes_dir = ROOT / "Volumes"
    if volumes_dir.exists():
        for docx_file in volumes_dir.glob("*.docx"):
            md_name = docx_file.stem.replace(" ", "-") + ".md"
            md_path = ROOT / "docs" / "architecture" / "volumes" / md_name
            md_path.parent.mkdir(parents=True, exist_ok=True)
            print(f"Converting: Volumes/{docx_file.name} -> docs/architecture/volumes/{md_name}")
            try:
                md_content = docx_to_markdown(docx_file)
                md_path.write_text(md_content, encoding="utf-8")
                conversions.append((str(docx_file), str(md_path)))
            except Exception as e:
                print(f"  ERROR: {e}")
        
        # Volume-1 chapters
        vol1_dir = volumes_dir / "Volume-1"
        if vol1_dir.exists():
            for docx_file in vol1_dir.glob("*.docx"):
                # Normalize name: "Volume I-Chapter-2" -> "Volume-I-Chapter-02"
                stem = docx_file.stem
                stem = stem.replace(" ", "-")
                # Pad chapter numbers
                if "Chapter" in stem or "chapter" in stem:
                    parts = stem.split("-")
                    for i, p in enumerate(parts):
                        if p.lower() == "chapter" and i + 1 < len(parts):
                            try:
                                num = int(parts[i + 1])
                                parts[i + 1] = f"{num:02d}"
                            except ValueError:
                                pass
                    stem = "-".join(parts)
                stem = stem.lower().replace("chapter", "Chapter")  # Keep Chapter capitalized
                
                md_name = stem + ".md"
                md_path = ROOT / "docs" / "architecture" / "volumes" / md_name
                md_path.parent.mkdir(parents=True, exist_ok=True)
                print(f"Converting: Volumes/Volume-1/{docx_file.name} -> docs/architecture/volumes/{md_name}")
                try:
                    md_content = docx_to_markdown(docx_file)
                    md_path.write_text(md_content, encoding="utf-8")
                    conversions.append((str(docx_file), str(md_path)))
                except Exception as e:
                    print(f"  ERROR: {e}")
    
    # Standards DOCX files (keep latest version only)
    standards_dir = ROOT / "Standards"
    if standards_dir.exists():
        # Find the highest numbered editorial standard
        editorial_files = sorted(standards_dir.glob("RIOS Editorial Standard-*.docx"), 
                                  key=lambda x: int(x.stem.split("-")[-1]))
        if editorial_files:
            latest = editorial_files[-1]
            md_name = "Editorial-Standard.md"
            md_path = ROOT / "docs" / "architecture" / "standards" / md_name
            md_path.parent.mkdir(parents=True, exist_ok=True)
            print(f"Converting: Standards/{latest.name} -> docs/architecture/standards/{md_name}")
            try:
                md_content = docx_to_markdown(latest)
                md_path.write_text(md_content, encoding="utf-8")
                conversions.append((str(latest), str(md_path)))
            except Exception as e:
                print(f"  ERROR: {e}")
    
    return conversions


# ============================================================
# PHASE 2: Move Architecture_v2 files to docs/
# ============================================================

def move_architecture_v2():
    """Move Architecture_v2 markdown files to docs/architecture/foundation/"""
    src = ROOT / "Architecture_v2"
    if not src.exists():
        return
    
    dest = ROOT / "docs" / "architecture" / "foundation"
    dest.mkdir(parents=True, exist_ok=True)
    
    for f in src.glob("*.md"):
        dest_file = dest / f.name
        if not dest_file.exists():
            print(f"Moving: Architecture_v2/{f.name} -> docs/architecture/foundation/{f.name}")
            shutil.copy2(str(f), str(dest_file))


# ============================================================
# PHASE 3: Move Engineering docs to docs/engineering/
# ============================================================

def move_engineering():
    """Move Engineering/ markdown files to docs/engineering/"""
    src = ROOT / "Engineering"
    if not src.exists():
        return
    
    dest = ROOT / "docs" / "engineering"
    dest.mkdir(parents=True, exist_ok=True)
    
    for f in src.glob("*.md"):
        dest_file = dest / f.name
        if not dest_file.exists():
            print(f"Moving: Engineering/{f.name} -> docs/engineering/{f.name}")
            shutil.copy2(str(f), str(dest_file))


# ============================================================
# PHASE 4: Move diagrams to docs/architecture/atlas/diagrams/
# ============================================================

def move_diagrams():
    """Move atlas diagrams to docs structure."""
    # v11 diagrams are canonical (latest)
    src = ROOT / "atlas_diagrams_v11"
    dest = ROOT / "docs" / "architecture" / "atlas" / "diagrams"
    
    if src.exists():
        dest.mkdir(parents=True, exist_ok=True)
        for f in src.glob("*.png"):
            dest_file = dest / f.name
            if not dest_file.exists():
                shutil.copy2(str(f), str(dest_file))
        print(f"Copied atlas_diagrams_v11/ -> docs/architecture/atlas/diagrams/")


# ============================================================
# PHASE 5: Move Python generators to tools/generators/
# ============================================================

def move_generators():
    """Move Python generator scripts to tools/generators/"""
    dest = ROOT / "tools" / "generators"
    dest.mkdir(parents=True, exist_ok=True)
    
    generators = [
        "generate_atlas.py",
        "generate_atlas_v11.py",
        "atlas_v11_helpers.py",
        "generate_adr.py",
        "generate_atm.py",
        "generate_constitution.py",
        "generate_validation_report.py",
    ]
    
    for gen in generators:
        src = ROOT / gen
        if src.exists():
            dest_file = dest / gen
            if not dest_file.exists():
                print(f"Moving: {gen} -> tools/generators/{gen}")
                shutil.copy2(str(src), str(dest_file))


# ============================================================
# PHASE 6: Archive old files
# ============================================================

def archive_files():
    """Move old/obsolete files to archive/"""
    archive = ROOT / "archive"
    
    # Archive old DOCX files (now have MD equivalents)
    docx_archive = archive / "docx"
    docx_archive.mkdir(parents=True, exist_ok=True)
    
    for docx_file in ROOT.glob("*.docx"):
        dest = docx_archive / docx_file.name
        if not dest.exists():
            print(f"Archiving: {docx_file.name} -> archive/docx/")
            shutil.move(str(docx_file), str(dest))
    
    # Archive Foundation (DOCX + TXT)
    foundation_archive = archive / "foundation"
    foundation_archive.mkdir(parents=True, exist_ok=True)
    foundation_dir = ROOT / "Foundation"
    if foundation_dir.exists():
        for f in foundation_dir.iterdir():
            if f.suffix in ('.docx', '.txt'):
                dest = foundation_archive / f.name
                if not dest.exists():
                    shutil.move(str(f), str(dest))
        print("Archived Foundation/ -> archive/foundation/")
    
    # Archive Standards (all DOCX + TXT versions)
    standards_archive = archive / "standards"
    standards_archive.mkdir(parents=True, exist_ok=True)
    standards_dir = ROOT / "Standards"
    if standards_dir.exists():
        for f in standards_dir.iterdir():
            if f.suffix in ('.docx', '.txt', '.DS_Store'):
                dest = standards_archive / f.name
                if not dest.exists():
                    shutil.move(str(f), str(dest))
        print("Archived Standards/ -> archive/standards/")
    
    # Archive Volumes (DOCX)
    volumes_archive = archive / "volumes"
    volumes_archive.mkdir(parents=True, exist_ok=True)
    volumes_dir = ROOT / "Volumes"
    if volumes_dir.exists():
        for f in volumes_dir.glob("*.docx"):
            dest = volumes_archive / f.name
            if not dest.exists():
                shutil.move(str(f), str(dest))
        vol1 = volumes_dir / "Volume-1"
        if vol1.exists():
            for f in vol1.iterdir():
                if f.suffix in ('.docx', '.txt', '.DS_Store'):
                    dest = volumes_archive / f.name
                    if not dest.exists():
                        shutil.move(str(f), str(dest))
        print("Archived Volumes/ -> archive/volumes/")
    
    # Archive old atlas diagrams (v1 is superseded by v11)
    old_atlas = ROOT / "atlas_diagrams"
    if old_atlas.exists():
        atlas_archive = archive / "atlas_diagrams_v1"
        atlas_archive.mkdir(parents=True, exist_ok=True)
        for f in old_atlas.glob("*.png"):
            dest = atlas_archive / f.name
            if not dest.exists():
                shutil.move(str(f), str(dest))
        print("Archived atlas_diagrams/ -> archive/atlas_diagrams_v1/")


# ============================================================
# PHASE 7: Clean up empty directories and temp files
# ============================================================

def cleanup():
    """Remove empty dirs, .DS_Store, __pycache__, etc."""
    import glob
    
    # Remove .DS_Store files
    for ds in ROOT.rglob(".DS_Store"):
        ds.unlink()
        print(f"Deleted: {ds.relative_to(ROOT)}")
    
    # Remove __pycache__
    for cache in ROOT.rglob("__pycache__"):
        if cache.is_dir():
            shutil.rmtree(str(cache))
            print(f"Deleted: {cache.relative_to(ROOT)}/")
    
    # Remove rios_atlas_env (empty venv-like dir)
    atlas_env = ROOT / "rios_atlas_env"
    if atlas_env.exists():
        shutil.rmtree(str(atlas_env))
        print("Deleted: rios_atlas_env/")
    
    # Remove empty placeholder directories that had .gitkeep
    for d in ["apps", "scripts", "infrastructure", "packages"]:
        p = ROOT / d
        if p.exists():
            contents = list(p.iterdir())
            gitkeeps = [f for f in contents if f.name == ".gitkeep"]
            if len(contents) == len(gitkeeps) and gitkeeps:
                # Only .gitkeep, move to keep structure but note it
                pass  # Keep these for now as they're part of the monorepo structure
    
    # Clean up root-level converted MD files (move to proper locations)
    for md in ROOT.glob("*.md"):
        if md.name == "README.md":
            continue
        # These are newly converted from DOCX, they're already in the right place
        # as reference. We'll handle them in phase 8.


# ============================================================
# PHASE 8: Organize converted root-level MD files
# ============================================================

def organize_root_mds():
    """Move converted MD files from root to docs/ structure."""
    atlas = ROOT / "docs" / "architecture" / "atlas"
    atlas.mkdir(parents=True, exist_ok=True)
    
    traceability = ROOT / "docs" / "architecture" / "traceability"
    traceability.mkdir(parents=True, exist_ok=True)
    
    validation = ROOT / "docs" / "architecture" / "validation"
    validation.mkdir(parents=True, exist_ok=True)
    
    adr = ROOT / "docs" / "architecture" / "adr"
    adr.mkdir(parents=True, exist_ok=True)
    
    # Map root MD files to destinations
    moves = {
        "RIOS-ATLAS.md": atlas / "ATLAS.md",
        "RIOS-ATLAS_v11.md": atlas / "ATLAS-v11.md",
        "RIOS-Architecture-Decision-Records_v1.0.md": adr / "ADR.md",
        "RIOS-Architecture-Traceability-Matrix_v1.0.md": traceability / "ATM.md",
        "RIOS-Architecture-Validation-Report_v1.0.md": validation / "Validation-Report.md",
        "RIOS-Claude-Code-Constitution_v1.0.md": ROOT / "docs" / "architecture" / "foundation" / "Constitution.md",
        "Volume-0.md": ROOT / "docs" / "architecture" / "volumes" / "Volume-0.md",
    }
    
    for src_name, dest_path in moves.items():
        src = ROOT / src_name
        if src.exists():
            dest_path.parent.mkdir(parents=True, exist_ok=True)
            if not dest_path.exists():
                print(f"Moving: {src_name} -> {dest_path.relative_to(ROOT)}")
                shutil.move(str(src), str(dest_path))


# ============================================================
# PHASE 9: Create ZIP archives
# ============================================================

def create_zips():
    """Compress archive subdirectories into ZIP files."""
    archive = ROOT / "archive"
    
    for subdir in ["docx", "foundation", "standards", "volumes", "atlas_diagrams_v1"]:
        dir_path = archive / subdir
        if dir_path.exists() and any(dir_path.iterdir()):
            zip_path = archive / f"{subdir}.zip"
            print(f"Creating: archive/{subdir}.zip")
            with zipfile.ZipFile(str(zip_path), 'w', zipfile.ZIP_DEFLATED) as zf:
                for f in dir_path.rglob("*"):
                    if f.is_file():
                        zf.write(str(f), f.relative_to(str(dir_path)))
            # Remove the uncompressed directory after zipping
            shutil.rmtree(str(dir_path))
            print(f"  Compressed and removed: archive/{subdir}/")


# ============================================================
# PHASE 10: Remove empty directories
# ============================================================

def remove_empty_dirs():
    """Remove now-empty directories."""
    dirs_to_check = [
        "Foundation",
        "Standards", 
        "Volumes",
        "Volumes/Volume-1",
        "atlas_diagrams",
        "Architecture_v2",
    ]
    
    for d in dirs_to_check:
        p = ROOT / d
        if p.exists():
            try:
                remaining = list(p.rglob("*"))
                files = [f for f in remaining if f.is_file()]
                if not files:
                    shutil.rmtree(str(p))
                    print(f"Removed empty directory: {d}/")
            except Exception as e:
                print(f"  Could not remove {d}: {e}")


# ============================================================
# MAIN
# ============================================================

def main():
    print("=" * 60)
    print("RIOS Repository Transformation")
    print("=" * 60)
    
    print("\n--- Phase 1: Converting DOCX to Markdown ---")
    conversions = convert_all_docx()
    print(f"Converted {len(conversions)} files")
    
    print("\n--- Phase 2: Moving Architecture_v2 ---")
    move_architecture_v2()
    
    print("\n--- Phase 3: Moving Engineering docs ---")
    move_engineering()
    
    print("\n--- Phase 4: Moving diagrams ---")
    move_diagrams()
    
    print("\n--- Phase 5: Moving generators ---")
    move_generators()
    
    print("\n--- Phase 6: Archiving old files ---")
    archive_files()
    
    print("\n--- Phase 7: Cleanup ---")
    cleanup()
    
    print("\n--- Phase 8: Organizing root MD files ---")
    organize_root_mds()
    
    print("\n--- Phase 9: Creating ZIP archives ---")
    create_zips()
    
    print("\n--- Phase 10: Removing empty directories ---")
    remove_empty_dirs()
    
    print("\n" + "=" * 60)
    print("Transformation complete!")
    print("=" * 60)

if __name__ == "__main__":
    main()